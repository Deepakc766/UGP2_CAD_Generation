# app_streamlit_cad.py
import traceback, sys
try:
    import transformers
    print("transformers:", getattr(transformers,'__file__',None))
    print("transformers version:", getattr(transformers,'__version__',None))
    from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
    print("SUCCESS: pipeline imported")
except Exception:
    print("---- FULL TRACEBACK (BEGIN) ----")
    traceback.print_exc()
    print("---- FULL TRACEBACK (END) ----")
    sys.exit(1)
import os
import io
import json
import textwrap
import tempfile
from pathlib import Path
from typing import List, Dict, Any
import logging
import re
from collections import defaultdict
from torch.nn import parallel
import streamlit as st
import pdfplumber                    # pip install pdfplumber
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas  # pip install reportlab
from dateutil import parser as dateparser  # pip install python-dateutil

from pdf2image import convert_from_bytes
from PIL import Image
import pytesseract
import shutil
import torch
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
from docx import Document as DocxDocument
from docx.shared import Pt
import pandas as pd
from torch.nn import DataParallel, DistributedDataParallel

# from langchain import LLMChain, PromptTemplate
# from langchain.memory import ConversationBufferMemory
# from langchain_huggingface import HuggingFacePipeline

from langchain.chains import LLMChain
from langchain_core.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory
from langchain_huggingface import HuggingFacePipeline

from huggingface_hub import login

login(token="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")
try:
    GEN_MODEL = None
    try:
        GEN_MODEL = st.secrets.get("GEN_MODEL")
    except Exception:
        GEN_MODEL = None
    if not GEN_MODEL:
        GEN_MODEL = os.environ.get("GEN_MODEL", "google/flan-t5-large")
except Exception:
    GEN_MODEL = os.environ.get("GEN_MODEL", "google/flan-t5-large")

DEFAULT_WINDOW_TOKENS = 2048
DEFAULT_OVERLAP_TOKENS = 256
DEFAULT_RESERVED_QA_TOKENS = 512
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200

# Ensure tesseract path
tesseract_path = shutil.which("tesseract")
if tesseract_path:
    pytesseract.pytesseract.tesseract_cmd = tesseract_path

class Node:
    def __init__(self, name: str, fn):
        self.name = name
        self.fn = fn
        self.outputs = None

    def run(self, *args, **kwargs):
        self.outputs = self.fn(*args, **kwargs)
        return self.outputs

class SimpleGraph:
    def __init__(self):
        self.nodes = []

    def add_node(self, node: Node):
        self.nodes.append(node)

    def run(self, start_node_index: int = 0, input=None, **kwargs):
        data = None
        for idx, node in enumerate(self.nodes[start_node_index:], start=start_node_index):
            if idx == start_node_index:
                data = node.run(input, **kwargs)
            else:
                data = node.run(data)
        return data


def pdf_bytes_to_images(pdf_bytes: bytes, dpi: int = 200) -> List[Image.Image]:
    images = convert_from_bytes(pdf_bytes, dpi=dpi)
    return images

def ocr_image_to_text(image: Image.Image) -> str:
    text = pytesseract.image_to_string(image)
    return text

def extract_text_from_pdf_bytes(pdf_bytes: bytes, dpi: int = 200) -> str:
    # fallback full-OCR (kept for compatibility)
    pages = pdf_bytes_to_images(pdf_bytes, dpi=dpi)
    texts = []
    for i, img in enumerate(pages):
        page_text = ocr_image_to_text(img)
        header = f"\n\n--- PAGE {i+1} ---\n\n"
        texts.append(header + page_text)
    return "\n".join(texts)

def extract_text_prefer_pdfplumber(pdf_bytes: bytes, dpi: int = 200):
    """
    Use pdfplumber to extract selectable text per page; fallback to OCR for pages with no text.
    Returns: (pages_list, full_text)
      pages_list: [{'page': n, 'text': str, 'is_ocr': bool}, ...]
      full_text: a single string with --- PAGE n --- markers
    """
    pages = []
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                try:
                    txt = page.extract_text() or ""
                except Exception:
                    txt = ""
                if txt and txt.strip():
                    pages.append({"page": i, "text": txt, "is_ocr": False})
                else:
                    # fallback OCR for blank/unextractable pages
                    try:
                        pil_img = page.to_image(resolution=dpi).original
                        ocr_txt = pytesseract.image_to_string(pil_img)
                    except Exception:
                        # final fallback: empty string
                        ocr_txt = ""
                    pages.append({"page": i, "text": ocr_txt, "is_ocr": True})
    except Exception as e:
        logging.warning("pdfplumber failed: %s. Falling back to full OCR.", e)
        # fallback to original full-OCR path
        ocr_full = extract_text_from_pdf_bytes(pdf_bytes, dpi=dpi)
        # attempt to split by page markers if present
        page_texts = []
        # Split by form-feed or by our marker
        if "\f" in ocr_full:
            page_texts = ocr_full.split("\f")
        else:
            # try to split with our PAGE marker
            splits = re.split(r"\n\s*--- PAGE\s+(\d+)\s+---\n", ocr_full)
            # crude fallback: treat as single page
            page_texts = [ocr_full]
        for idx, p_text in enumerate(page_texts, start=1):
            pages.append({"page": idx, "text": p_text, "is_ocr": True})

    full = []
    for p in pages:
        full.append(f"\n\n--- PAGE {p['page']} ---\n{p['text']}")
    return pages, "\n".join(full)

def simple_chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    if not text:
        return []
    joined = text.replace("\r\n", "\n")
    chunks = []
    start = 0
    while start < len(joined):
        end = min(len(joined), start + size)
        chunks.append(joined[start:end])
        if end == len(joined):
            break
        start = max(0, end - overlap)
    return chunks

@st.cache_resource(show_spinner=False)
def build_generation_pipeline_and_wrapper(model_name: str = GEN_MODEL):
    tokenizer = AutoTokenizer.from_pretrained(model_name, use_fast=True)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
    device = 0 if torch.cuda.is_available() else -1
    trans_pipe = pipeline("text2text-generation", model=model, tokenizer=tokenizer, device=device if device>=0 else -1)
    llm_wrapper = HuggingFacePipeline(pipeline=trans_pipe)
    return trans_pipe, tokenizer, llm_wrapper

def count_tokens(text: str, tokenizer) -> int:
    return len(tokenizer.encode(text, add_special_tokens=False))

def chunk_text_by_tokens(text: str, tokenizer, window_tokens: int, overlap_tokens: int) -> List[str]:
    tokens = tokenizer.encode(text, add_special_tokens=False)
    if len(tokens) <= window_tokens:
        return [tokenizer.decode(tokens, skip_special_tokens=True)]
    windows = []
    start = 0
    step = window_tokens - overlap_tokens
    while start < len(tokens):
        end = min(start + window_tokens, len(tokens))
        slice_ids = tokens[start:end]
        windows.append(tokenizer.decode(slice_ids, skip_special_tokens=True))
        if end == len(tokens):
            break
        start += step
    return windows

def make_full_contract_text(chunks: List[str], include_chunk_headers: bool = True) -> str:
    parts = []
    for i, c in enumerate(chunks):
        if include_chunk_headers:
            parts.append(f"\n\n--- CHUNK {i+1} ---\n{c}")
        else:
            parts.append(c)
    return "\n\n".join(parts)

CONTRACT_Q_PROMPT = PromptTemplate(
    input_variables=["contract", "question"],
    template=textwrap.dedent("""\
        You are a contract analyst. Use ONLY the CONTRACT text below to answer the question.
        If the contract does not contain the answer, reply exactly: "Not found in contract."

        CONTRACT:
        {contract}

        QUESTION:
        {question}

        ANSWER:""")
)

WINDOW_Q_PROMPT = PromptTemplate(
    input_variables=["window", "question", "idx", "total"],
    template=textwrap.dedent("""\
        You are a contract analyst. Use ONLY the CONTRACT WINDOW below to answer the question.
        If the window does not contain the answer, reply exactly: "Not found in contract."

        CONTRACT WINDOW {idx}/{total}:
        {window}

        QUESTION:
        {question}

        ANSWER:""")
)

# CAD JSON prompt (strict JSON schema output)
CAD_JSON_PROMPT = PromptTemplate(
    input_variables=["window_text"],
    template=textwrap.dedent("""\
        You are a contract analyst. Use ONLY the CONTRACT TEXT provided below. Produce a single valid JSON object EXACTLY matching the schema described below. DO NOT output any extra commentary, explanations, or markdown — output must be strict JSON.

        CONTRACT WINDOW:
        {window_text}

        SCHEMA:
        {
          "salient_features": {
            "project_name": "string | null",
            "project_location": "string | null",
            "contract_ref": "string | null",
            "employer": {"name":"string | null","address":"string | null","contact":"string | null","sources":["page X"]},
            "contractor": {"name":"string | null","address":"string | null","contact":"string | null","sources":["page X"]},
            "contract_type": "string | null",
            "contract_value": {"amount":number | null,"currency":"string | null","source":["page X"]},
            "price_escalation": {"applicable": true|false, "index":"string|null", "formula":"string|null", "source":["page X"]},
            "award_date": "YYYY-MM-DD | null",
            "commencement_date": "YYYY-MM-DD | null",
            "scheduled_completion_date": "YYYY-MM-DD | null",
            "defect_liability_period": "string | null",
            "scope_overview": "string | null",
            "sources": ["page X"]
          },
          "important_submittals": [
            {"stage":"before|during|after|null","document":"string","due":"string","notes":"string","sources":["page X"]}
          ],
          "notice_clauses": [
            {"event":"string","notifier":"string","recipient":"string","timeline":"string","method":"string","sources":["page X"]}
          ],
          "project_progress_clauses": [
            {"topic":"string","summary":"string","sources":["page X","clause Y"]}
          ],
          "payment": {
            "contract_sum": {"amount": number | null, "currency":"string | null", "source":["page X"]},
            "mobilization_advance": {"applicable": true|false,"amount_pct": number | null,"security":"string|null","recovery":"string|null"},
            "interim_payments": [{"frequency":"string","certifier":"string","payment_days_after_certification":number | null}],
            "retention": {"percent": number | null,"release_condition":"string","source":["page X"]},
            "final_payment": {"required_docs":["string"],"release_days_after_submission": number | null},
            "escalation_clause": {"applicable": true|false,"details":"string","sources":["page X"]},
            "sources": ["page X"]
          },
          "risks_and_allocation": [
            {"risk":"string","severity":"Major|High|Medium|Low","probability":"Likely|Possible|Rare","responsibility":"Employer|Contractor|Shared","notes":"string","sources":["page X"]}
          ],
          "claims_disputes_arbitration": {
            "arbitration": {"applicable": true|false,"notes":"string","sources":["page X"]},
            "dispute_forum":"string|null",
            "claims_summary":[{"topic":"string","process":"string","sources":["page X"]}],
            "sources":["page X"]
          }
        }

        INSTRUCTIONS:
        - Use ONLY information present in the provided text window.
        - If a field is not present, use null, empty list [], or false as appropriate.
        - For fields you fill, add 'sources' with page/clause if determinable. If not, use [].
        - Output MUST be strict JSON and parsable by json.loads().
        - Keep individual text fields concise (<= 300 words).

        OUTPUT:
        A single JSON object following the SCHEMA above.
        """)
)

# Compliance JSON prompt
COMPLIANCE_JSON_PROMPT = PromptTemplate(
    input_variables=["window_text", "rule"],
    template=textwrap.dedent("""\
        You are a contract analyst. Use ONLY the CONTRACT WINDOW below to answer the compliance rule.

        CONTRACT WINDOW:
        {window_text}

        RULE:
        {rule}

        INSTRUCTIONS:
        - If the window contains evidence answering the rule, return JSON exactly like:
          {{"rule":"{rule}", "present": true, "summary":"short summary", "quote":"verbatim quote (<=200 chars)", "sources":["page X","clause Y"], "confidence":0.90}}
        - If not present, return:
          {{"rule":"{rule}", "present": false, "summary": null, "quote": null, "sources": [], "confidence": 0.99}}
        - Output MUST be strict JSON only (no explanatory text).
        """)
)

def ask_with_concatenate_chain(llm_chain: LLMChain, contract_text: str, question: str, reserved_tokens: int = DEFAULT_RESERVED_QA_TOKENS) -> str:
    return llm_chain.predict(contract=contract_text, question=question)

def ask_with_sliding_chain(llm_chain: LLMChain, tokenizer, contract_text: str, question: str,
                           window_tokens: int = DEFAULT_WINDOW_TOKENS, overlap_tokens: int = DEFAULT_OVERLAP_TOKENS) -> Dict[str, Any]:
    windows = chunk_text_by_tokens(contract_text, tokenizer, window_tokens, overlap_tokens)
    answers = []
    for i, w in enumerate(windows, start=1):
        out = llm_chain.predict(window=w, question=question, idx=i, total=len(windows))
        answers.append({"window_index": i-1, "answer": out.strip()})
    non_empty = [a for a in answers if a["answer"] and "not found in contract" not in a["answer"].lower()]
    if not non_empty:
        longest = max(answers, key=lambda x: len(x["answer"] or ""))
        return {"answer": longest["answer"].strip() or "Not found in contract.", "used_windows": [longest["window_index"]], "method": "sliding", "all_window_answers": answers}
    freq = {}
    for a in non_empty:
        key = a["answer"].strip()
        freq[key] = freq.get(key, 0) + 1
    best_text = max(freq.items(), key=lambda kv: (kv[1], len(kv[0])))[0]
    used_windows = [a["window_index"] for a in non_empty if a["answer"].strip() == best_text]
    unique_answers = list({a["answer"].strip() for a in non_empty})
    final_ans = best_text
    if len(unique_answers) > 1:
        final_ans = "\n\n--- AGGREGATED ANSWERS FROM MULTIPLE WINDOWS ---\n\n" + "\n\n---\n\n".join(unique_answers)
    return {"answer": final_ans.strip(), "used_windows": used_windows, "method": "sliding", "all_window_answers": answers}

def ask_direct_langchain(llm_wrapper: HuggingFacePipeline, tokenizer, chunks: List[str], question: str, strategy: str = "sliding",
                         window_tokens: int = DEFAULT_WINDOW_TOKENS, overlap_tokens: int = DEFAULT_OVERLAP_TOKENS) -> Dict[str, Any]:
    contract_text = make_full_contract_text(chunks, include_chunk_headers=True)
    if strategy == "concatenate":
        chain = LLMChain(llm=llm_wrapper, prompt=CONTRACT_Q_PROMPT)
        ans = chain.predict(contract=contract_text, question=question)
        return {"answer": ans.strip(), "method": "concatenate", "used_windows": [0]}
    else:
        chain = LLMChain(llm=llm_wrapper, prompt=WINDOW_Q_PROMPT)
        return ask_with_sliding_chain(chain, tokenizer, contract_text, question, window_tokens=window_tokens, overlap_tokens=overlap_tokens)

def compliance_check_json(chunks, rules, llm_wrapper, tokenizer, strategy="sliding", window_tokens: int = DEFAULT_WINDOW_TOKENS, overlap_tokens: int = DEFAULT_OVERLAP_TOKENS):
    chain = LLMChain(llm=llm_wrapper, prompt=COMPLIANCE_JSON_PROMPT)
    windows = chunk_text_by_tokens("".join(chunks), tokenizer, window_tokens, overlap_tokens)
    results = []
    for rule in rules:
        aggregated = []
        for w in windows:
            raw = chain.predict(window_text=w, rule=rule)
            raw = raw.strip()
            parsed = None
            try:
                parsed = json.loads(raw)
            except:
                s = raw.find('{'); e = raw.rfind('}')
                if s!=-1 and e!=-1:
                    try:
                        parsed = json.loads(raw[s:e+1])
                    except:
                        parsed = None
            if parsed:
                aggregated.append(parsed)
        # merge
        final = {"rule": rule, "present": False, "summary": None, "quote": None, "sources": [], "confidence": 0.0}
        for a in aggregated:
            if a.get("present"):
                final["present"] = True
                final["summary"] = a.get("summary") or final["summary"]
                final["quote"] = a.get("quote") or final["quote"]
                final["sources"].extend([s for s in a.get("sources", []) if s not in final["sources"]])
                try:
                    final["confidence"] = max(final["confidence"], float(a.get("confidence", 0.0)))
                except:
                    pass
        if not aggregated:
            final["confidence"] = 0.0
        results.append(final)
    return results


def merge_json_objects(objs):
    result = {}
    for o in objs:
        if not isinstance(o, dict):
            continue
        for k, v in o.items():
            if v in (None, [], "", {}):
                continue
            if k not in result:
                result[k] = v
            else:
                if isinstance(v, list) and isinstance(result[k], list):
                    for it in v:
                        if it not in result[k]:
                            result[k].append(it)
                elif isinstance(v, dict) and isinstance(result[k], dict):
                    for kk, vv in v.items():
                        if kk not in result[k] and vv not in (None, "", [], {}):
                            result[k][kk] = vv
                else:
                    # keep existing
                    pass
    return result

def find_quote_sources(quote: str, pages) -> list:
    """Search for exact or fuzzy occurrences of quote across pages; return page list."""
    if not quote or not quote.strip():
        return []
    quote_norm = re.sub(r'\s+', ' ', quote.strip())[:400]
    sources = []
    for p in pages:
        page_text_norm = re.sub(r'\s+', ' ', (p.get('text') or "").strip())
        if quote_norm in page_text_norm:
            sources.append(f"page {p['page']}")
    # fuzzy fallback: substring of quote
    if not sources and len(quote_norm) > 40:
        sub = quote_norm[:80]
        for p in pages:
            if sub in (p.get('text') or ""):
                sources.append(f"page {p['page']}")
    return sources

def node_generate_cad_json_docx_pdf(llm_wrapper, tokenizer, chunks, pages, output_basename="generated_CAD", strategy="sliding", window_tokens: int = DEFAULT_WINDOW_TOKENS, overlap_tokens: int = DEFAULT_OVERLAP_TOKENS):
    """
    - Ask LLM per window to return strict JSON CAD as per schema (CAD_JSON_PROMPT)
    - Merge results, save JSON
    - Create DOCX and PDF summaries
    """
    chain = LLMChain(llm=llm_wrapper, prompt=CAD_JSON_PROMPT)
    windows = chunk_text_by_tokens("".join(chunks), tokenizer, window_tokens, overlap_tokens)
    json_responses = []
    for w in windows:
        raw = chain.predict(window_text=w).strip()
        parsed = None
        try:
            parsed = json.loads(raw)
        except:
            s = raw.find('{'); e = raw.rfind('}')
            if s != -1 and e != -1:
                try:
                    parsed = json.loads(raw[s:e+1])
                except:
                    parsed = None
        if parsed:
            json_responses.append(parsed)
    if not json_responses:
        # fallback: single concatenate attempt
        concat_chain = LLMChain(llm=llm_wrapper, prompt=CONTRACT_Q_PROMPT)
        fallback_raw = concat_chain.predict(contract=make_full_contract_text(chunks), question="Generate CAD JSON as per schema.")
        try:
            parsed = json.loads(fallback_raw)
            json_responses.append(parsed)
        except:
            raise RuntimeError("No valid JSON responses from LLM for CAD generation. Try concatenate strategy or adjust prompts.")
    merged = merge_json_objects(json_responses)

    # attach sources for quotes/summaries
    def attach_sources_for_quotes(container):
        if isinstance(container, dict):
            for k, v in container.items():
                if isinstance(v, str) and len(v) > 20:
                    if any(keyword in k.lower() for keyword in ['quote','scope','summary','overview','text']):
                        srcs = find_quote_sources(v, pages)
                        if srcs:
                            container.setdefault('sources', []).extend([s for s in srcs if s not in container.get('sources', [])])
                elif isinstance(v, (dict, list)):
                    attach_sources_for_quotes(v)
        elif isinstance(container, list):
            for item in container:
                attach_sources_for_quotes(item)

    attach_sources_for_quotes(merged)

    # Save JSON
    json_path = Path(tempfile.gettempdir()) / f"{output_basename}.json"
    with open(json_path, "w", encoding="utf8") as f:
        json.dump(merged, f, indent=2, ensure_ascii=False)

    # Create DOCX (structured)
    docx = DocxDocument()
    style = docx.styles['Normal']; font = style.font; font.name='Arial'; font.size = Pt(11)
    docx.add_heading("Contract Appreciation Document (Generated)", level=1)
    docx.add_paragraph("This document is generated automatically. Verify against contract before finalization.")

    # Salient features
    sf = merged.get("salient_features", {})
    docx.add_heading("1. Salient Features", level=2)
    if sf:
        for k, v in sf.items():
            if isinstance(v, dict):
                display_val = v.get('name') or v.get('amount') or json.dumps(v)
                src = ", ".join(v.get('sources', [])) if v.get('sources') else ""
                docx.add_paragraph(f"{k}: {display_val} {'(sources: '+src+')' if src else ''}")
            else:
                docx.add_paragraph(f"{k}: {v}")
    else:
        docx.add_paragraph("Not found")

    # Important submittals
    submittals = merged.get("important_submittals", [])
    if submittals:
        docx.add_heading("2. Important Submittals", level=2)
        cols = ["stage", "document", "due", "notes", "sources"]
        table = docx.add_table(rows=1, cols=len(cols))
        for i, c in enumerate(cols):
            table.rows[0].cells[i].text = c
        for r in submittals:
            rc = table.add_row().cells
            rc[0].text = r.get("stage","")
            rc[1].text = r.get("document","")
            rc[2].text = r.get("due","")
            rc[3].text = r.get("notes","")
            rc[4].text = ", ".join(r.get("sources", []))
    else:
        docx.add_paragraph("No submittals found")

    # Notices
    notices = merged.get("notice_clauses", [])
    if notices:
        docx.add_heading("3. Notice Clauses", level=2)
        for n in notices:
            docx.add_paragraph(f"{n.get('event')} | {n.get('notifier')} -> {n.get('recipient')} | timeline: {n.get('timeline')} | method: {n.get('method')} | sources: {', '.join(n.get('sources',[]))}")

    # Payment summary
    docx.add_heading("4. Payment Summary", level=2)
    pay = merged.get("payment", {})
    if pay:
        for k, v in pay.items():
            docx.add_paragraph(f"{k}: {json.dumps(v) if isinstance(v,(dict,list)) else str(v)}")
    else:
        docx.add_paragraph("No payment information found")

    # Risks
    docx.add_heading("5. Risks and Allocation", level=2)
    risks = merged.get("risks_and_allocation", [])
    if risks:
        for r in risks:
            docx.add_paragraph(json.dumps(r))
    else:
        docx.add_paragraph("No risks found")

    # Claims / Disputes
    docx.add_heading("6. Claims, Disputes & Arbitration", level=2)
    docx.add_paragraph(json.dumps(merged.get("claims_disputes_arbitration", {}), indent=2))

    docx_path = Path(tempfile.gettempdir()) / f"{output_basename}.docx"
    docx.save(docx_path)

    # Create a simple PDF summary with reportlab
    pdf_path = Path(tempfile.gettempdir()) / f"{output_basename}.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=A4)
    width, height = A4
    margin = 50
    y = height - margin



    def write_wrapped_text(cobj, text, x, y, max_width, leading=12):
        from reportlab.lib.utils import simpleSplit
        lines = simpleSplit(text, "Helvetica", 10, max_width)
        for ln in lines:
            if y < margin + 40:
                cobj.showPage()
                y = height - margin
            cobj.drawString(x, y, ln)
            y -= leading
        return y

    c.setFont("Helvetica-Bold", 14)
    c.drawString(margin, y, "Contract Appreciation Document (Generated)")
    y -= 22
    c.setFont("Helvetica", 10)
    # Project summary short
    proj = sf.get("project_name") if isinstance(sf, dict) else None
    c.drawString(margin, y, f"Project: {proj or 'N/A'}")
    y -= 14
    employer_name = sf.get("employer", {}).get("name") if isinstance(sf.get("employer"), dict) else sf.get("employer")
    contractor_name = sf.get("contractor", {}).get("name") if isinstance(sf.get("contractor"), dict) else sf.get("contractor")
    c.drawString(margin, y, f"Employer: {employer_name or 'N/A'}")
    y -= 14
    c.drawString(margin, y, f"Contractor: {contractor_name or 'N/A'}")
    y -= 18
    # scope overview
    scope = sf.get("scope_overview") if isinstance(sf, dict) else sf.get("scope_overview") if isinstance(sf, dict) else sf.get("scope_overview", None)
    if isinstance(scope, str):
        y = write_wrapped_text(c, "Scope Overview: " + (scope[:1500]), margin, y, width - 2*margin, leading=12)
    c.showPage()
    c.save()

    return str(json_path), str(docx_path), str(pdf_path)

def split_text_into_clauses_with_positions(pages):
    """
    Heuristic splitting into clause-like blocks per page.
    Returns list of clause dicts with page info.
    """
    clauses = []
    for p in pages:
        page_num = p['page']
        text = p['text'] or ""
        pattern = re.compile(r'(^|\n)\s*(\d{1,2}(?:\.\d+){0,3})\s+([A-Z][^\n]{1,120})', re.M)
        matches = list(pattern.finditer(text))
        if not matches:
            clauses.append({
                "clause_id": f"p{page_num}_blk1",
                "heading": None,
                "text": text,
                "page": page_num,
                "char_start": 0,
                "char_end": len(text)
            })
        else:
            for idx, m in enumerate(matches):
                start = m.start()
                end = matches[idx+1].start() if idx+1 < len(matches) else len(text)
                heading_num = m.group(2).strip()
                heading_text = m.group(3).strip()
                body = text[start:end].strip()
                clauses.append({
                    "clause_id": heading_num,
                    "heading": heading_text,
                    "text": body,
                    "page": page_num,
                    "char_start": start,
                    "char_end": end
                })
    return clauses

# -------------------------
# Conflict detection engine (priority checks)
# -------------------------
def check_commencement_vs_site_possession(pages, full_text):
    date_pattern = r'((?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4})|(?:\d{1,2}\s+[A-Za-z]+\s+\d{4})|(?:[A-Za-z]+\s+\d{1,2},\s*\d{4}))'
    site_pos_pattern = re.compile(r'(site possession|possession of the site|hand over (?:the )?site)[^\.\n]{0,250}(' + date_pattern + r')', re.I)
    commence_pattern = re.compile(r'(commenc(?:e|ement) date|start date|date of commencement|date of commencement of works)[^\.\n]{0,250}(' + date_pattern + r')', re.I)

    site_matches = []
    comm_matches = []
    for p in pages:
        text = p.get('text') or ""
        for m in site_pos_pattern.finditer(text):
            raw = m.group(2)
            try:
                dt = dateparser.parse(raw, dayfirst=True)
            except:
                dt = None
            site_matches.append({"page": p['page'], "raw": raw, "date": dt, "snippet": m.group(0)})
        for m in commence_pattern.finditer(text):
            raw = m.group(2)
            try:
                dt = dateparser.parse(raw, dayfirst=True)
            except:
                dt = None
            comm_matches.append({"page": p['page'], "raw": raw, "date": dt, "snippet": m.group(0)})

    conflicts = []
    if site_matches and comm_matches:
        comm_dates = [c['date'] for c in comm_matches if c['date']]
        site_dates = [s['date'] for s in site_matches if s['date']]
        if comm_dates and site_dates:
            min_comm = min(comm_dates)
            max_site = max(site_dates)
            if max_site > min_comm:
                conflicts.append({
                    "type": "commencement_vs_site_possession",
                    "severity": "Critical",
                    "message": f"Site possession ({max_site.date()}) occurs after commencement ({min_comm.date()}).",
                    "evidence": {"commencement": comm_matches, "site_possession": site_matches}
                })
    return conflicts

def check_payment_term_mismatch(pages):
    pay_pattern = re.compile(r'pay(?:ment)?\s+(?:within|in)\s+(\d{1,3})\s+days', re.I)
    matches = []
    for p in pages:
        for m in pay_pattern.finditer(p.get('text') or ""):
            matches.append({"page": p['page'], "days": int(m.group(1)), "snippet": m.group(0)})
    conflicts = []
    if matches:
        days_set = sorted(set(m['days'] for m in matches))
        if len(days_set) > 1:
            conflicts.append({
                "type": "payment_term_mismatch",
                "severity": "High",
                "message": f"Different payment terms found: {days_set} days.",
                "evidence": matches
            })
    return conflicts

def check_retention_mismatch(pages):
    ret_pattern = re.compile(r'retention[^\d]{0,20}(\d{1,2})\s*%?', re.I)
    matches = []
    for p in pages:
        for m in ret_pattern.finditer(p.get('text') or ""):
            try:
                pct = int(m.group(1))
            except:
                continue
            matches.append({"page": p['page'], "pct": pct, "snippet": m.group(0)})
    conflicts = []
    if matches:
        pct_set = sorted(set(m['pct'] for m in matches))
        if len(pct_set) > 1:
            conflicts.append({
                "type": "retention_mismatch",
                "severity": "High",
                "message": f"Different retention percentages found: {pct_set}%.",
                "evidence": matches
            })
    return conflicts

def check_defect_liability_mismatch(pages):
    dl_pattern = re.compile(r'(defect liability period|defects liability|warranty period)[^\d]{0,60}(\d{1,3})\s*(months|month|years|year)?', re.I)
    matches = []
    for p in pages:
        for m in dl_pattern.finditer(p.get('text') or ""):
            val = int(m.group(2))
            unit = (m.group(3) or "months").lower()
            months = val * 12 if "year" in unit else val
            matches.append({"page": p['page'], "value": months, "raw": m.group(0)})
    conflicts = []
    if matches:
        vals = sorted(set(m['value'] for m in matches))
        if len(vals) > 1:
            conflicts.append({
                "type": "defect_liability_mismatch",
                "severity": "High",
                "message": f"Different defect liability lengths found (in months): {vals}.",
                "evidence": matches
            })
    return conflicts

def check_arbitration_vs_court_conflict(pages):
    arb_pattern = re.compile(r'\barbitrat', re.I)
    court_pattern = re.compile(r'\bcourt', re.I)
    arbs = []
    courts = []
    for p in pages:
        t = p.get('text') or ""
        if arb_pattern.search(t):
            arbs.append({"page": p['page']})
        if court_pattern.search(t):
            courts.append({"page": p['page']})
    conflicts = []
    if arbs and courts:
        conflicts.append({
            "type": "dispute_resolution_conflict",
            "severity": "Critical",
            "message": "Arbitration-related terms and court-related terms both found; possible conflict in dispute resolution or interim relief.",
            "evidence": {"arbitration_pages": arbs, "court_pages": courts}
        })
    return conflicts

CONFLICT_TO_CATEGORY = {
    "commencement_vs_site_possession": ("Contractual Disputes (Site Possession)", "Consider deeming commencement as actual possession or grant EOT + cost recovery. See 'SITE POSSESSION' clause."),
    "payment_term_mismatch": ("Contractual Disputes (Payment)", "Harmonize payment timeline; prefer main-body or latest-dated doc. Consider change-order for invoices/payments."),
    "retention_mismatch": ("Financial / Payment", "Clarify retention % in payment schedule; propose single retained % and update annex."),
    "defect_liability_mismatch": ("Quality / Warranty", "Adopt the stricter DLP or clarify which schedule governs; update annex."),
    "dispute_resolution_conflict": ("Dispute Resolution", "Resolve arbitration vs court language — prefer arbitration clause with interim relief carve-out."),
    # extend mapping for other checks you add
}

def map_conflict_to_practical_category(conflict):
    ctype = conflict.get("type")
    cat, hint = CONFLICT_TO_CATEGORY.get(ctype, ("Other / Administrative", "Review conflict and map manually."))
    conflict["category"] = cat
    conflict["resolution_hint"] = hint
    return conflict


# def run_conflict_detection(pages, full_text):
#     conflicts = []
#     conflicts.extend(check_commencement_vs_site_possession(pages, full_text))
#     conflicts.extend(check_payment_term_mismatch(pages))
#     conflicts.extend(check_retention_mismatch(pages))
#     conflicts.extend(check_defect_liability_mismatch(pages))
#     conflicts.extend(check_arbitration_vs_court_conflict(pages))
#     # extend with more checks if needed
#     return conflicts

def run_conflict_detection(pages, full_text):
    conflicts = []
    conflicts.extend(check_commencement_vs_site_possession(pages, full_text))
    conflicts.extend(check_payment_term_mismatch(pages))
    conflicts.extend(check_retention_mismatch(pages))
    conflicts.extend(check_defect_liability_mismatch(pages))
    conflicts.extend(check_arbitration_vs_court_conflict(pages))
    # extend more checks...
    # attach category + hint
    conflicts = [map_conflict_to_practical_category(c) for c in conflicts]
    return conflicts



def node_ocr_and_extract(pdf_bytes: bytes) -> Dict[str, Any]:
    st.info("Extracting text from PDF (pdfplumber preferred; fallback to OCR).")
    pages, full_text = extract_text_prefer_pdfplumber(pdf_bytes, dpi=200)
    return {"pages": pages, "full_text": full_text}

def node_chunk_text_from_full(raw_text: str) -> List[str]:
    st.info("Chunking text for UI and downstream (char-based chunks).")
    return simple_chunk_text(raw_text, size=3000, overlap=500)

def node_build_model_wrapper(_) -> Dict[str, Any]:
    st.info("Loading model and LangChain wrapper...")
    trans_pipe, tokenizer, llm_wrapper = build_generation_pipeline_and_wrapper(GEN_MODEL)
    return {"pipe": trans_pipe, "tokenizer": tokenizer, "llm": llm_wrapper}


# Streamlit UI

st.set_page_config(page_title="Contract CAD Chatbot - LangChain Direct LLM", layout="wide")
st.title("Automated Contract CAD Chatbot LangChain + Direct LLM Context")

# Sidebar controls
with st.sidebar:
    st.header("Controls")
    strategy = st.radio("Query strategy", options=["sliding", "concatenate"], index=0)
    window_tokens = st.number_input("Window tokens (approx)", min_value=512, max_value=64000, value=DEFAULT_WINDOW_TOKENS, step=256)
    overlap_tokens = st.number_input("Overlap tokens", min_value=0, max_value=window_tokens//2, value=DEFAULT_OVERLAP_TOKENS, step=64)
    reserved_qa_tokens = st.number_input("Reserved tokens for QA (generation)", min_value=64, max_value=4096, value=DEFAULT_RESERVED_QA_TOKENS, step=64)
    st.markdown("---")
    uploaded_file = st.file_uploader("Upload contract PDF (scanned or digital)", type=["pdf"])
    regen = st.button("Reprocess / Rebuild context")
    clear_all = st.button("Clear session")

if clear_all:
    for k in list(st.session_state.keys()):
        del st.session_state[k]
    st.experimental_rerun()

# Graph nodes & pipeline
graph = SimpleGraph()
graph.add_node(Node("OCR", lambda b: node_ocr_and_extract(b)))
graph.add_node(Node("Chunk", lambda info: node_chunk_text_from_full(info["full_text"])))
graph.add_node(Node("Model", node_build_model_wrapper))

# Build pipeline & process upload
if uploaded_file and (regen or not st.session_state.get("graph_built", False)):
    try:
        with st.spinner("Running pipeline (extract -> chunk -> model load)..."):
            pdf_bytes = uploaded_file.read()
            extract_info = node_ocr_and_extract(pdf_bytes)
            pages = extract_info["pages"]
            full_text = extract_info["full_text"]
            chunks = node_chunk_text_from_full(full_text)
            model_info = node_build_model_wrapper(None)
            st.session_state["pages"] = pages
            st.session_state["raw_text"] = full_text
            st.session_state["chunks"] = chunks
            st.session_state["trans_pipe"] = model_info["pipe"]
            st.session_state["tokenizer"] = model_info["tokenizer"]
            st.session_state["llm_wrapper"] = model_info["llm"]
            st.session_state["graph_built"] = True
        st.success("Processed PDF and loaded model.")
    except Exception as e:
        st.exception(e)
        st.error("Processing failed.")

# Setup memory-based chain for chat (LangChain Conversation)
if st.session_state.get("graph_built", False) and "chat_memory" not in st.session_state:
    st.session_state["chat_memory"] = ConversationBufferMemory(memory_key="chat_history", return_messages=True)

# Layout: Chat (left) and CAD/Compliance (right)
col1, col2 = st.columns([2, 1])

with col1:
    st.subheader("Chat with the contract (LangChain direct-LLM)")
    if not st.session_state.get("graph_built", False):
        st.info("Upload & Reprocess a PDF to enable chat.")
    else:
        if "conversation" not in st.session_state:
            st.session_state["conversation"] = []
        for turn in st.session_state["conversation"]:
            if turn["role"] == "user":
                st.chat_message("user").write(turn["text"])
            else:
                st.chat_message("assistant").write(turn["text"])

        user_input = st.chat_input("Ask about the contract (e.g., 'What is the payment clause?')")
        if user_input:
            st.session_state["conversation"].append({"role":"user","text":user_input})
            tokenizer = st.session_state["tokenizer"]
            llm_wrapper = st.session_state["llm_wrapper"]
            chunks = st.session_state["chunks"]
            chain = LLMChain(llm=llm_wrapper, prompt=CONTRACT_Q_PROMPT, memory=st.session_state["chat_memory"])
            with st.spinner("Asking LLM..."):
                res = ask_direct_langchain(llm_wrapper, tokenizer, chunks, user_input, strategy=strategy, window_tokens=window_tokens, overlap_tokens=overlap_tokens)
                answer = res.get("answer", "No answer.")
            st.session_state["conversation"].append({"role":"assistant","text":answer})
            st.chat_message("assistant").write(answer)
            if st.checkbox("Show debug windows & answers", key=f"dbg_chat_{len(st.session_state['conversation'])}"):
                st.write(res)

with col2:
    st.subheader("CAD Generator / Document")
    if not st.session_state.get("graph_built", False):
        st.info("Upload & Reprocess PDF to enable CAD generation.")
    else:
        if st.button("Generate CAD (JSON + DOCX + PDF)"):
            try:
                llm_wrapper = st.session_state["llm_wrapper"]
                tokenizer = st.session_state["tokenizer"]
                chunks = st.session_state["chunks"]
                pages = st.session_state["pages"]
                with st.spinner("Generating CAD (JSON/DOCX/PDF)..."):
                    json_path, docx_path, pdf_path = node_generate_cad_json_docx_pdf(llm_wrapper, tokenizer, chunks, pages, output_basename="generated_CAD", strategy=strategy, window_tokens=window_tokens, overlap_tokens=overlap_tokens)
                st.session_state["cad_json"] = json_path
                st.session_state["cad_docx"] = docx_path
                st.session_state["cad_pdf"] = pdf_path
                st.success("CAD generated as JSON, DOCX and PDF.")
            except Exception as e:
                st.exception(e)
                st.error("CAD generation failed.")

        if st.session_state.get("cad_json"):
            try:
                with open(st.session_state["cad_json"], "rb") as f:
                    st.download_button("Download CAD JSON", f.read(), file_name=Path(st.session_state['cad_json']).name)
            except Exception as e:
                st.error(f"Failed to open CAD JSON: {e}")
        if st.session_state.get("cad_docx"):
            try:
                with open(st.session_state["cad_docx"], "rb") as f:
                    st.download_button("Download CAD DOCX", f.read(), file_name=Path(st.session_state['cad_docx']).name)
            except Exception as e:
                st.error(f"Failed to open CAD DOCX: {e}")
        if st.session_state.get("cad_pdf"):
            try:
                with open(st.session_state["cad_pdf"], "rb") as f:
                    st.download_button("Download CAD PDF", f.read(), file_name=Path(st.session_state['cad_pdf']).name)
            except Exception as e:
                st.error(f"Failed to open CAD PDF: {e}")

    st.markdown("---")
    st.subheader("📑 Compliance Check")
    if not st.session_state.get("graph_built", False):
        st.info("Upload a contract and press Reprocess first.")
    else:
        rules_input = st.text_area("Enter compliance rules (one per line)", value="Does the contract specify payment terms?\nIs there a termination clause?\nWhat is the governing law?")
        if st.button("Run Compliance Check"):
            rules = [r.strip() for r in rules_input.splitlines() if r.strip()]
            llm_wrapper = st.session_state["llm_wrapper"]
            tokenizer = st.session_state["tokenizer"]
            chunks = st.session_state["chunks"]
            with st.spinner("Running compliance checks..."):
                results = compliance_check_json(chunks, rules, llm_wrapper, tokenizer, strategy=strategy, window_tokens=window_tokens, overlap_tokens=overlap_tokens)
            st.subheader("Compliance Report")
            df_rows = []
            for r in results:
                st.markdown(f"**Rule:** {r['rule']}  \n**Present:** {r['present']}  \n**Summary:** {r['summary']}  \n**Quote:** {r['quote']}  \n**Sources:** {r['sources']}  \n**Confidence:** {r['confidence']}")
                df_rows.append([r['rule'], r['present'], r['summary'], r['quote'], ", ".join(r['sources']), r['confidence']])
            if df_rows:
                st.dataframe(pd.DataFrame(df_rows, columns=["Rule", "Present", "Summary", "Quote", "Sources", "Confidence"]))

    st.markdown("---")
    st.subheader("⚠️ Conflict Detection")
    if not st.session_state.get("graph_built", False):
        st.info("Upload a contract and press Reprocess first.")
    else:
        if st.button("Run Conflict Detection"):
            try:
                pages = st.session_state.get("pages")
                full_text = st.session_state.get("raw_text")
                with st.spinner("Detecting conflicts..."):
                    conflicts = run_conflict_detection(pages, full_text)
                st.session_state["conflicts"] = conflicts
                if not conflicts:
                    st.success("No conflicts detected by automated checks.")
                else:
                    st.warning(f"{len(conflicts)} potential conflicts detected.")
                    for idx, c in enumerate(conflicts, start=1):
                        # st.markdown(f"### {idx}. Type: {c['type']}  \n**Severity:** {c['severity']}  \n**Message:** {c['message']}")
                        # st.json(c.get("evidence"))
                        st.markdown(f"### {idx}. Category: **{c.get('category')}**  \n**Type:** {c['type']}  \n**Severity:** {c['severity']}  \n**Message:** {c['message']}")
                        st.markdown(f"**Suggested quick fix:** {c.get('resolution_hint')}")
                        st.json(c.get("evidence"))
            except Exception as e:
                st.exception(e)
                st.error("Conflict detection failed.")

    st.markdown("---")
    st.subheader("Raw text & pages preview")
    if st.session_state.get("raw_text"):
        if st.checkbox("Show OCR/raw extracted text (first 15000 chars)"):
            st.text_area("Extracted Text", value=st.session_state["raw_text"][:15000], height=400)
    else:
        st.info("No extracted text available. Upload a PDF and press Reprocess.")

st.caption("LangChain + HuggingFace pipeline. Outputs: CAD JSON (audit), DOCX, PDF.")

# End of file
