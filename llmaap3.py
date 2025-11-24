# # app_streamlit_cad.py
# import os
# import io
# import json
# import textwrap
# import tempfile
# from pathlib import Path
# from typing import List, Dict, Any, Optional
# import logging
# import re
# from collections import defaultdict
# import streamlit as st
# import pdfplumber                    # pip install pdfplumber
# from reportlab.lib.pagesizes import A4
# from reportlab.pdfgen import canvas  # pip install reportlab
# from dateutil import parser as dateparser  # pip install python-dateutil

# from pdf2image import convert_from_bytes
# from PIL import Image
# import pytesseract
# import shutil
# import torch
# from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
# from docx import Document as DocxDocument
# from docx.shared import Pt
# import pandas as pd
# HF_TOKEN="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"
# # LangChain
# from langchain.chains import LLMChain
# from langchain_core.prompts import PromptTemplate
# from langchain.memory import ConversationBufferMemory
# from langchain_huggingface import HuggingFacePipeline

# # Optional: login to HF Hub via secrets/env (no hardcoded token!)
# try:
#     from huggingface_hub import login as hf_login
#     hf_token = st.secrets.get("HF_TOKEN", os.environ.get("HF_TOKEN"))
#     if hf_token:
#         try:
#             hf_login(token=hf_token)
#         except Exception as e:
#             st.warning(f"HuggingFace login failed: {e}")
# except Exception:
#     pass

# # -------------------------
# # Config & defaults
# # -------------------------
# try:
#     GEN_MODEL = None
#     try:
#         GEN_MODEL = st.secrets.get("GEN_MODEL")
#     except Exception:
#         GEN_MODEL = None
#     if not GEN_MODEL:
#         GEN_MODEL = os.environ.get("GEN_MODEL", "google/flan-t5-large")  # faster default for CPU
# except Exception:
#     GEN_MODEL = os.environ.get("GEN_MODEL", "google/flan-t5-large")

# DEFAULT_WINDOW_TOKENS = 512  # smaller windows for speed
# DEFAULT_OVERLAP_TOKENS = 128
# DEFAULT_RESERVED_QA_TOKENS = 512
# CHUNK_SIZE = 1000
# CHUNK_OVERLAP = 200

# # Ensure tesseract path
# tesseract_path = shutil.which("tesseract")
# if tesseract_path:
#     pytesseract.pytesseract.tesseract_cmd = tesseract_path

# # -------------------------
# # Small helper structures
# # -------------------------
# class Node:
#     def __init__(self, name: str, fn):
#         self.name = name
#         self.fn = fn
#         self.outputs = None
#     def run(self, *args, **kwargs):
#         self.outputs = self.fn(*args, **kwargs)
#         return self.outputs

# class SimpleGraph:
#     def __init__(self):
#         self.nodes = []
#     def add_node(self, node: Node):
#         self.nodes.append(node)
#     def run(self, start_node_index: int = 0, input=None, **kwargs):
#         data = None
#         for idx, node in enumerate(self.nodes[start_node_index:], start=start_node_index):
#             if idx == start_node_index:
#                 data = node.run(input, **kwargs)
#             else:
#                 data = node.run(data)
#         return data

# # -------------------------
# # PDF → Text extraction
# # -------------------------
# def pdf_bytes_to_images(pdf_bytes: bytes, dpi: int = 200) -> List[Image.Image]:
#     images = convert_from_bytes(pdf_bytes, dpi=dpi)
#     return images

# def ocr_image_to_text(image: Image.Image) -> str:
#     text = pytesseract.image_to_string(image)
#     return text

# def extract_text_from_pdf_bytes(pdf_bytes: bytes, dpi: int = 200) -> str:
#     pages = pdf_bytes_to_images(pdf_bytes, dpi=dpi)
#     texts = []
#     for i, img in enumerate(pages):
#         page_text = ocr_image_to_text(img)
#         header = f"\n\n--- PAGE {i+1} ---\n\n"
#         texts.append(header + page_text)
#     return "\n".join(texts)

# def extract_text_prefer_pdfplumber(pdf_bytes: bytes, dpi: int = 200):
#     """
#     Use pdfplumber to extract selectable text per page; fallback to OCR for pages with no text.
#     Returns: (pages_list, full_text)
#       pages_list: [{'page': n, 'text': str, 'is_ocr': bool}, ...]
#       full_text: a single string with --- PAGE n --- markers
#     """
#     pages = []
#     try:
#         with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
#             for i, page in enumerate(pdf.pages, start=1):
#                 try:
#                     txt = page.extract_text() or ""
#                 except Exception:
#                     txt = ""
#                 if txt and txt.strip():
#                     pages.append({"page": i, "text": txt, "is_ocr": False})
#                 else:
#                     try:
#                         pil_img = page.to_image(resolution=dpi).original
#                         ocr_txt = pytesseract.image_to_string(pil_img)
#                     except Exception:
#                         ocr_txt = ""
#                     pages.append({"page": i, "text": ocr_txt, "is_ocr": True})
#     except Exception as e:
#         logging.warning("pdfplumber failed: %s. Falling back to full OCR.", e)
#         ocr_full = extract_text_from_pdf_bytes(pdf_bytes, dpi=dpi)
#         page_texts = []
#         if "\f" in ocr_full:
#             page_texts = ocr_full.split("\f")
#         else:
#             page_texts = [ocr_full]
#         for idx, p_text in enumerate(page_texts, start=1):
#             pages.append({"page": idx, "text": p_text, "is_ocr": True})

#     full = []
#     for p in pages:
#         full.append(f"\n\n--- PAGE {p['page']} ---\n{p['text']}")
#     return pages, "\n".join(full)

# def simple_chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
#     if not text:
#         return []
#     joined = text.replace("\r\n", "\n")
#     chunks = []
#     start = 0
#     while start < len(joined):
#         end = min(len(joined), start + size)
#         chunks.append(joined[start:end])
#         if end == len(joined):
#             break
#         start = max(0, end - overlap)
#     return chunks

# # -------------------------
# # Token-based windowing
# # -------------------------
# def count_tokens(text: str, tokenizer) -> int:
#     return len(tokenizer.encode(text, add_special_tokens=False))

# def chunk_text_by_tokens(text: str, tokenizer, window_tokens: int, overlap_tokens: int) -> List[str]:
#     tokens = tokenizer.encode(text, add_special_tokens=False)
#     if len(tokens) <= window_tokens:
#         return [tokenizer.decode(tokens, skip_special_tokens=True)]
#     windows = []
#     start = 0
#     step = max(1, window_tokens - overlap_tokens)
#     while start < len(tokens):
#         end = min(start + window_tokens, len(tokens))
#         slice_ids = tokens[start:end]
#         windows.append(tokenizer.decode(slice_ids, skip_special_tokens=True))
#         if end == len(tokens):
#             break
#         start += step
#     return windows

# def make_full_contract_text(chunks: List[str], include_chunk_headers: bool = True) -> str:
#     parts = []
#     for i, c in enumerate(chunks):
#         if include_chunk_headers:
#             parts.append(f"\n\n--- CHUNK {i+1} ---\n{c}")
#         else:
#             parts.append(c)
#     return "\n\n".join(parts)

# # -------------------------
# # Prompts
# # -------------------------
# CONTRACT_Q_PROMPT = PromptTemplate(
#     input_variables=["contract", "question"],
#     template=textwrap.dedent("""\
#         You are a contract analyst. Use ONLY the CONTRACT text below to answer the question.
#         If the contract does not contain the answer, reply exactly: "Not found in contract."

#         CONTRACT:
#         {contract}

#         QUESTION:
#         {question}

#         ANSWER:""")
# )

# WINDOW_Q_PROMPT = PromptTemplate(
#     input_variables=["window", "question", "idx", "total"],
#     template=textwrap.dedent("""\
#         You are a contract analyst. Use ONLY the CONTRACT WINDOW below to answer the question.
#         If the window does not contain the answer, reply exactly: "Not found in contract."

#         CONTRACT WINDOW {idx}/{total}:
#         {window}

#         QUESTION:
#         {question}

#         ANSWER:""")
# )
# def _sanitize_json_like(txt: str) -> dict | None:
#     """Try to coerce model output to strict JSON."""
#     if not txt:
#         return None
#     s = txt.find('{'); e = txt.rfind('}')
#     if s == -1 or e == -1:
#         return None
#     cand = txt[s:e+1]

#     # Common fixes
#     cand = re.sub(r",\s*}", "}", cand)
#     cand = re.sub(r",\s*]", "]", cand)
#     # ensure true/false/null are lowercase (T5 usually okay)
#     cand = re.sub(r"\bTrue\b", "true", cand)
#     cand = re.sub(r"\bFalse\b", "false", cand)
#     cand = re.sub(r"\bNone\b", "null", cand)

#     try:
#         return json.loads(cand)
#     except Exception:
#         # attempt to add quotes to keys like: key: value -> "key": value (only for top-level-ish places)
#         def _quote_keys(m):
#             return f"\"{m.group(1)}\":"
#         cand2 = re.sub(r'(?m)^\s*([A-Za-z0-9_]+)\s*:', _quote_keys, cand)
#         try:
#             return json.loads(cand2)
#         except Exception:
#             return None

# REPAIR_JSON_PROMPT = PromptTemplate(
#     input_variables=["bad_text"],
#     template=textwrap.dedent("""\
#         Convert the following content into ONE strict JSON object that matches the CAD schema. 
#         Only output the JSON (start with '{' and end with '}').

#         CONTENT:
#         {bad_text}
#         """)
# )

# def _cad_candidate_window_score(w: str) -> int:
#     wl = w.lower()
#     score = 0
#     # favor windows with these signals:
#     for kw in ("employer", "contractor", "contract price", "payment", "retention", "mobilization", "bank guarantee",
#                "defects liability", "completion", "commencement", "arbitration", "governing law", "notice", "submittal", "method statement"):
#         if kw in wl:
#             score += 1
#     return score

# def _build_minimal_cad_from_text(text: str) -> dict:
#     """Last-resort minimal JSON using regex so we never hard-fail."""
#     t = text
#     out = {
#       "salient_features": {
#         "project_name": None, "project_location": None, "contract_ref": None,
#         "employer": {"name": None, "address": None, "contact": None, "sources": []},
#         "contractor": {"name": None, "address": None, "contact": None, "sources": []},
#         "contract_type": None,
#         "contract_value": {"amount": None, "currency": None, "source": []},
#         "price_escalation": {"applicable": False, "index": None, "formula": None, "source": []},
#         "award_date": None, "commencement_date": None, "scheduled_completion_date": None,
#         "defect_liability_period": None, "scope_overview": None, "sources": []
#       },
#       "important_submittals": [],
#       "notice_clauses": [],
#       "project_progress_clauses": [],
#       "payment": {
#         "contract_sum": {"amount": None, "currency": None, "source": []},
#         "mobilization_advance": {"applicable": False, "amount_pct": None, "security": None, "recovery": None},
#         "interim_payments": [],
#         "retention": {"percent": None, "release_condition": "", "source": []},
#         "final_payment": {"required_docs": [], "release_days_after_submission": None},
#         "escalation_clause": {"applicable": False, "details": "", "sources": []},
#         "sources": []
#       },
#       "risks_and_allocation": [],
#       "claims_disputes_arbitration": {"arbitration": {"applicable": False, "notes": "", "sources": []}, "dispute_forum": None, "claims_summary": [], "sources": []}
#     }

#     # Employer/Contractor names (very naive)
#     m_emp = re.search(r"Employer[:\s\-]+([^\n,]+)", t, re.I)
#     if m_emp: out["salient_features"]["employer"]["name"] = m_emp.group(1).strip()
#     m_con = re.search(r"Contractor[:\s\-]+([^\n,]+)", t, re.I)
#     if m_con: out["salient_features"]["contractor"]["name"] = m_con.group(1).strip()

#     # Contract price
#     m_price_num = re.search(r"Contract Price[:\s\-]+(?:INR|₹)?\s*([\d,]+)", t, re.I)
#     if m_price_num:
#         try:
#             amt = int(m_price_num.group(1).replace(",", ""))
#             out["payment"]["contract_sum"]["amount"] = amt
#             out["payment"]["contract_sum"]["currency"] = "INR"
#             out["salient_features"]["contract_value"]["amount"] = amt
#             out["salient_features"]["contract_value"]["currency"] = "INR"
#         except: pass

#     # Payment days (interim)
#     m_pay = re.search(r"(?:Interim Payments?|payable)\s.*?\bwithin\s+(\d{1,3})\s+days", t, re.I)
#     if m_pay:
#         try:
#             out["payment"]["interim_payments"] = [{"frequency":"monthly","certifier":"Engineer","payment_days_after_certification": int(m_pay.group(1))}]
#         except: pass

#     # Retention percent
#     m_ret = re.search(r"Retention(?: money)?\s*(?:of)?\s*(\d{1,2})\s*%?", t, re.I)
#     if m_ret:
#         try:
#             out["payment"]["retention"]["percent"] = int(m_ret.group(1))
#         except: pass

#     # DLP
#     m_dlp = re.search(r"(?:Defects? Liability.*?)(\d{1,3})\s*(months|month|years|year)", t, re.I)
#     if m_dlp:
#         val = int(m_dlp.group(1)); unit = m_dlp.group(2).lower()
#         out["salient_features"]["defect_liability_period"] = f"{val} {unit}"

#     # Arbitration/court
#     if "arbitration" in t.lower():
#         out["claims_disputes_arbitration"]["arbitration"]["applicable"] = True

#     return out

# CAD_JSON_PROMPT = PromptTemplate(
#     input_variables=["window_text"],
#     template=textwrap.dedent("""\
#         You are a contract analyst. Use ONLY the CONTRACT TEXT provided below. Produce ONE valid JSON object EXACTLY matching the schema. 
#         IMPORTANT: Output STRICT JSON only — no prose, no Markdown. Your output MUST start with '{{' and end with '}}'.

#         CONTRACT WINDOW:
#         {window_text}

#         SCHEMA:
#         {{
#           "salient_features": {{
#             "project_name": "string | null",
#             "project_location": "string | null",
#             "contract_ref": "string | null",
#             "employer": {{"name":"string | null","address":"string | null","contact":"string | null","sources":["page X"]}},
#             "contractor": {{"name":"string | null","address":"string | null","contact":"string | null","sources":["page X"]}},
#             "contract_type": "string | null",
#             "contract_value": {{"amount":number | null,"currency":"string | null","source":["page X"]}},
#             "price_escalation": {{"applicable": true|false, "index":"string|null", "formula":"string|null", "source":["page X"]}},
#             "award_date": "YYYY-MM-DD | null",
#             "commencement_date": "YYYY-MM-DD | null",
#             "scheduled_completion_date": "YYYY-MM-DD | null",
#             "defect_liability_period": "string | null",
#             "scope_overview": "string | null",
#             "sources": ["page X"]
#           }},
#           "important_submittals": [
#             {{"stage":"before|during|after|null","document":"string","due":"string","notes":"string","sources":["page X"]}}
#           ],
#           "notice_clauses": [
#             {{"event":"string","notifier":"string","recipient":"string","timeline":"string","method":"string","sources":["page X"]}}
#           ],
#           "project_progress_clauses": [
#             {{"topic":"string","summary":"string","sources":["page X","clause Y"]}}
#           ],
#           "payment": {{
#             "contract_sum": {{"amount": number | null, "currency":"string | null", "source":["page X"]}},
#             "mobilization_advance": {{"applicable": true|false,"amount_pct": number | null,"security":"string|null","recovery":"string|null"}},
#             "interim_payments": [{{"frequency":"string","certifier":"string","payment_days_after_certification":number | null}}],
#             "retention": {{"percent": number | null,"release_condition":"string","source":["page X"]}},
#             "final_payment": {{"required_docs":["string"],"release_days_after_submission": number | null}},
#             "escalation_clause": {{"applicable": true|false,"details":"string","sources":["page X"]}},
#             "sources": ["page X"]
#           }},
#           "risks_and_allocation": [
#             {{"risk":"string","severity":"Major|High|Medium|Low","probability":"Likely|Possible|Rare","responsibility":"Employer|Contractor|Shared","notes":"string","sources":["page X"]}}
#           ],
#           "claims_disputes_arbitration": {{
#             "arbitration": {{"applicable": true|false,"notes":"string","sources":["page X"]}},
#             "dispute_forum":"string|null",
#             "claims_summary":[{{"topic":"string","process":"string","sources":["page X"]}}],
#             "sources":["page X"]
#           }}
#         }}

#         EXAMPLE FORMAT (structure only):
#         {{"salient_features": {{"project_name": null, "sources":[]}}, "important_submittals": [], "notice_clauses": [], "project_progress_clauses": [], "payment": {{"sources":[]}}, "risks_and_allocation": [], "claims_disputes_arbitration": {{"sources":[]}}}}
#         """)
# )

# COMPLIANCE_JSON_PROMPT = PromptTemplate(
#     input_variables=["window_text", "rule"],
#     template=textwrap.dedent("""\
#         You are a contract analyst. Use ONLY the CONTRACT WINDOW below to answer the compliance rule.
#         Return STRICT JSON ONLY — no prose, no markdown, no surrounding text.

#         CONTRACT WINDOW:
#         {window_text}

#         RULE:
#         {rule}

#         VALID JSON EXAMPLES (copy the structure):
#         # Example present=true
#         {{"rule":"sample rule", "present": true, "summary":"short", "quote":"verbatim (<=200 chars)", "sources":["page X","clause Y"], "confidence":0.90}}

#         # Example present=false
#         {{"rule":"sample rule", "present": false, "summary": null, "quote": null, "sources": [], "confidence": 0.99}}

#         INSTRUCTIONS:
#         - If the window contains evidence answering the rule, return:
#           {{"rule":"{rule}", "present": true, "summary":"short summary", "quote":"verbatim quote (<=200 chars)", "sources":["page X","clause Y"], "confidence":0.90}}
#         - If not present, return:
#           {{"rule":"{rule}", "present": false, "summary": null, "quote": null, "sources": [], "confidence": 0.99}}
#         - STRICT JSON ONLY.
#         """)
# )



# # -------------------------
# # Generation pipelines (MAIN + COMPLIANCE)
# # -------------------------
# @st.cache_resource(show_spinner=False)
# def build_generation_pipelines(model_name: str = GEN_MODEL):
#     tokenizer = AutoTokenizer.from_pretrained(model_name, use_fast=True)
#     model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
#     device = 0 if torch.cuda.is_available() else -1

#     # MAIN: long outputs (CAD JSON, rich QA)
#     main_gen_kwargs = dict(
#         max_new_tokens=1024,
#         temperature=0.0,
#         do_sample=False,
#         top_p=1.0,
#         repetition_penalty=1.05,
#         no_repeat_ngram_size=3,
#         return_full_text=False,
#         truncation=True,
#     )
#     # COMPLIANCE: short/fast outputs
#     comp_gen_kwargs = dict(
#     max_new_tokens=256,  # was 192
#     temperature=0.0,
#     do_sample=False,
#     top_p=1.0,
#     repetition_penalty=1.05,
#     no_repeat_ngram_size=3,
#     return_full_text=False,
#     truncation=True,
#     )


#     trans_pipe = pipeline(
#         "text2text-generation",
#         model=model,
#         tokenizer=tokenizer,
#         device=device if device >= 0 else -1
#     )

#     llm_main = HuggingFacePipeline(pipeline=trans_pipe, model_kwargs=main_gen_kwargs)
#     llm_compliance = HuggingFacePipeline(pipeline=trans_pipe, model_kwargs=comp_gen_kwargs)

#     return tokenizer, llm_main, llm_compliance

# # -------------------------
# # JSON parsing helpers
# # -------------------------
# def _loose_json_grab(txt: str) -> Optional[dict]:
#     if not txt:
#         return None
#     s = txt.find('{'); e = txt.rfind('}')
#     if s == -1 or e == -1:
#         return None
#     cand = txt[s:e+1]
#     cand = re.sub(r",\s*}", "}", cand)
#     cand = re.sub(r",\s*]", "]", cand)
#     try:
#         return json.loads(cand)
#     except Exception:
#         return None

# # -------------------------
# # QA helpers
# # -------------------------
# def ask_with_concatenate_chain(llm_chain: LLMChain, contract_text: str, question: str) -> str:
#     return llm_chain.predict(contract=contract_text, question=question)

# def ask_with_sliding_chain(llm_chain: LLMChain, tokenizer, contract_text: str, question: str,
#                            window_tokens: int = DEFAULT_WINDOW_TOKENS, overlap_tokens: int = DEFAULT_OVERLAP_TOKENS) -> Dict[str, Any]:
#     windows = chunk_text_by_tokens(contract_text, tokenizer, window_tokens, overlap_tokens)
#     answers = []
#     for i, w in enumerate(windows, start=1):
#         out = llm_chain.predict(window=w, question=question, idx=i, total=len(windows))
#         answers.append({"window_index": i-1, "answer": (out or "").strip()})
#     non_empty = [a for a in answers if a["answer"] and "not found in contract" not in a["answer"].lower()]
#     if not non_empty:
#         longest = max(answers, key=lambda x: len(x["answer"] or ""))
#         return {"answer": (longest["answer"] or "Not found in contract.").strip(), "used_windows": [longest["window_index"]], "method": "sliding", "all_window_answers": answers}
#     freq = {}
#     for a in non_empty:
#         key = a["answer"].strip()
#         freq[key] = freq.get(key, 0) + 1
#     best_text = max(freq.items(), key=lambda kv: (kv[1], len(kv[0])))[0]
#     used_windows = [a["window_index"] for a in non_empty if a["answer"].strip() == best_text]
#     unique_answers = list({a["answer"].strip() for a in non_empty})
#     final_ans = best_text
#     if len(unique_answers) > 1:
#         final_ans = "\n\n--- AGGREGATED ANSWERS FROM MULTIPLE WINDOWS ---\n\n" + "\n\n---\n\n".join(unique_answers)
#     return {"answer": final_ans.strip(), "used_windows": used_windows, "method": "sliding", "all_window_answers": answers}

# def ask_direct_langchain(llm_wrapper: HuggingFacePipeline, tokenizer, full_text_for_windows: str, question: str, strategy: str = "sliding",
#                          window_tokens: int = DEFAULT_WINDOW_TOKENS, overlap_tokens: int = DEFAULT_OVERLAP_TOKENS, memory: Optional[ConversationBufferMemory]=None) -> Dict[str, Any]:
#     if strategy == "concatenate":
#         chain = LLMChain(llm=llm_wrapper, prompt=CONTRACT_Q_PROMPT, memory=memory)
#         ans = chain.predict(contract=full_text_for_windows, question=question)
#         return {"answer": (ans or "").strip(), "method": "concatenate", "used_windows": [0]}
#     else:
#         chain = LLMChain(llm=llm_wrapper, prompt=WINDOW_Q_PROMPT, memory=memory)
#         return ask_with_sliding_chain(chain, tokenizer, full_text_for_windows, question, window_tokens=window_tokens, overlap_tokens=overlap_tokens)

# # -------------------------
# # CAD Generation
# # -------------------------
# def find_quote_sources(quote: str, pages) -> list:
#     if not quote or not quote.strip():
#         return []
#     quote_norm = re.sub(r'\s+', ' ', quote.strip())[:400]
#     sources = []
#     for p in pages:
#         page_text_norm = re.sub(r'\s+', ' ', (p.get('text') or "").strip())
#         if quote_norm in page_text_norm:
#             sources.append(f"page {p['page']}")
#     if not sources and len(quote_norm) > 40:
#         sub = quote_norm[:80]
#         for p in pages:
#             if sub in (p.get('text') or ""):
#                 sources.append(f"page {p['page']}")
#     return sources

# def merge_json_objects(objs):
#     result = {}
#     for o in objs:
#         if not isinstance(o, dict):
#             continue
#         for k, v in o.items():
#             if v in (None, [], "", {}):
#                 continue
#             if k not in result:
#                 result[k] = v
#             else:
#                 if isinstance(v, list) and isinstance(result[k], list):
#                     for it in v:
#                         if it not in result[k]:
#                             result[k].append(it)
#                 elif isinstance(v, dict) and isinstance(result[k], dict):
#                     for kk, vv in v.items():
#                         if vv not in (None, "", [], {}):
#                             result[k].setdefault(kk, vv)
#                 else:
#                     pass
#     return result




# def node_generate_cad_json_docx_pdf(llm_wrapper, tokenizer, full_text_with_pages: str, pages, output_basename="generated_CAD",
#                                     window_tokens: int = DEFAULT_WINDOW_TOKENS, overlap_tokens: int = DEFAULT_OVERLAP_TOKENS):
#     """Robust CAD generator:
#        - picks top-k candidate windows,
#        - retries each window,
#        - sanitizes/repairs JSON,
#        - falls back to minimal JSON if needed.
#     """
#     # 0) Build windows (page markers preserved)
#     windows = chunk_text_by_tokens(full_text_with_pages, tokenizer, window_tokens, overlap_tokens)
#     if not windows:
#         raise RuntimeError("Empty text windows for CAD.")

#     # 1) Rank windows by simple keyword score, keep top-k
#     scored = sorted([(i, _cad_candidate_window_score(w), w) for i, w in enumerate(windows)], key=lambda x: x[1], reverse=True)
#     candidates = [w for _, _, w in scored[: min(8, max(3, len(scored)//3))]]  # top 3..8 windows

#     # 2) Ask model with strict prompt; retry + repair if needed
#     chain = LLMChain(llm=llm_wrapper, prompt=CAD_JSON_PROMPT)
#     json_responses = []
#     for w in candidates:
#         ok = False
#         for attempt in range(2):  # two attempts per window
#             raw = (chain.predict(window_text=w) or "").strip()
#             parsed = None
#             try:
#                 parsed = json.loads(raw)
#             except Exception:
#                 parsed = _sanitize_json_like(raw)

#             if not parsed:
#                 # repair step: ask model to convert its own text to JSON
#                 repair_chain = LLMChain(llm=llm_wrapper, prompt=REPAIR_JSON_PROMPT)
#                 repaired = (repair_chain.predict(bad_text=raw) or "").strip()
#                 try:
#                     parsed = json.loads(repaired)
#                 except Exception:
#                     parsed = _sanitize_json_like(repaired)

#             if parsed and isinstance(parsed, dict):
#                 json_responses.append(parsed)
#                 ok = True
#                 break  # next candidate
#         # continue to next candidate regardless

#     # 3) If still nothing, fall back to minimal JSON from regex over the best window
#     if not json_responses:
#         fallback_dict = _build_minimal_cad_from_text(candidates[0])
#         json_responses.append(fallback_dict)

#     # 4) Merge
#     merged = merge_json_objects(json_responses)

#     # 5) Attach sources for quotes/summaries where possible
#     def attach_sources_for_quotes(container):
#         if isinstance(container, dict):
#             for k, v in list(container.items()):
#                 if isinstance(v, str) and len(v) > 20:
#                     if any(keyword in k.lower() for keyword in ['quote','scope','summary','overview','text']):
#                         srcs = find_quote_sources(v, pages)
#                         if srcs:
#                             container.setdefault('sources', []).extend([s for s in srcs if s not in container.get('sources', [])])
#                 elif isinstance(v, (dict, list)):
#                     attach_sources_for_quotes(v)
#         elif isinstance(container, list):
#             for item in container:
#                 attach_sources_for_quotes(item)
#     attach_sources_for_quotes(merged)

#     # 6) Save JSON
#     json_path = Path(tempfile.gettempdir()) / f"{output_basename}.json"
#     with open(json_path, "w", encoding="utf8") as f:
#         json.dump(merged, f, indent=2, ensure_ascii=False)

#     # 7) DOCX (same as before)
#     docx = DocxDocument()
#     style = docx.styles['Normal']; font = style.font; font.name='Arial'; font.size = Pt(11)
#     docx.add_heading("Contract Appreciation Document (Generated)", level=1)
#     docx.add_paragraph("This document is generated automatically. Verify against contract before finalization.")

#     sf = merged.get("salient_features", {})
#     docx.add_heading("1. Salient Features", level=2)
#     if sf:
#         for k, v in sf.items():
#             if isinstance(v, dict):
#                 display_val = v.get('name') or v.get('amount') or json.dumps(v)
#                 src = ", ".join(v.get('sources', [])) if v.get('sources') else ""
#                 docx.add_paragraph(f"{k}: {display_val} {'(sources: '+src+')' if src else ''}")
#             else:
#                 docx.add_paragraph(f"{k}: {v}")
#     else:
#         docx.add_paragraph("Not found")

#     submittals = merged.get("important_submittals", [])
#     if submittals:
#         docx.add_heading("2. Important Submittals", level=2)
#         cols = ["stage", "document", "due", "notes", "sources"]
#         table = docx.add_table(rows=1, cols=len(cols))
#         for i, c in enumerate(cols):
#             table.rows[0].cells[i].text = c
#         for r in submittals:
#             rc = table.add_row().cells
#             rc[0].text = r.get("stage","")
#             rc[1].text = r.get("document","")
#             rc[2].text = r.get("due","")
#             rc[3].text = r.get("notes","")
#             rc[4].text = ", ".join(r.get("sources", []))
#     else:
#         docx.add_paragraph("No submittals found")

#     notices = merged.get("notice_clauses", [])
#     if notices:
#         docx.add_heading("3. Notice Clauses", level=2)
#         for n in notices:
#             docx.add_paragraph(f"{n.get('event')} | {n.get('notifier')} -> {n.get('recipient')} | timeline: {n.get('timeline')} | method: {n.get('method')} | sources: {', '.join(n.get('sources',[]))}")

#     docx.add_heading("4. Payment Summary", level=2)
#     pay = merged.get("payment", {})
#     if pay:
#         for k, v in pay.items():
#             docx.add_paragraph(f"{k}: {json.dumps(v) if isinstance(v,(dict,list)) else str(v)}")
#     else:
#         docx.add_paragraph("No payment information found")

#     docx.add_heading("5. Risks and Allocation", level=2)
#     risks = merged.get("risks_and_allocation", [])
#     if risks:
#         for r in risks:
#             docx.add_paragraph(json.dumps(r))
#     else:
#         docx.add_paragraph("No risks found")

#     docx.add_heading("6. Claims, Disputes & Arbitration", level=2)
#     docx.add_paragraph(json.dumps(merged.get("claims_disputes_arbitration", {}), indent=2))

#     docx_path = Path(tempfile.gettempdir()) / f"{output_basename}.docx"
#     docx.save(docx_path)

#     # 8) Simple PDF summary (unchanged, keep your existing draw code)
#     pdf_path = Path(tempfile.gettempdir()) / f"{output_basename}.pdf"
#     c = canvas.Canvas(str(pdf_path), pagesize=A4)
#     width, height = A4
#     margin = 50
#     y = height - margin

#     def write_wrapped_text(cobj, text, x, y, max_width, leading=12):
#         from reportlab.lib.utils import simpleSplit
#         lines = simpleSplit(text, "Helvetica", 10, max_width)
#         for ln in lines:
#             if y < margin + 40:
#                 cobj.showPage()
#                 y = height - margin
#             cobj.drawString(x, y, ln)
#             y -= leading
#         return y

#     c.setFont("Helvetica-Bold", 14)
#     c.drawString(margin, y, "Contract Appreciation Document (Generated)"); y -= 22
#     c.setFont("Helvetica", 10)
#     proj = sf.get("project_name") if isinstance(sf, dict) else None
#     c.drawString(margin, y, f"Project: {proj or 'N/A'}"); y -= 14
#     employer_name = sf.get("employer", {}).get("name") if isinstance(sf.get("employer"), dict) else None
#     contractor_name = sf.get("contractor", {}).get("name") if isinstance(sf.get("contractor"), dict) else None
#     c.drawString(margin, y, f"Employer: {employer_name or 'N/A'}"); y -= 14
#     c.drawString(margin, y, f"Contractor: {contractor_name or 'N/A'}"); y -= 18
#     scope = sf.get("scope_overview") if isinstance(sf, dict) else None
#     if isinstance(scope, str) and scope.strip():
#         y = write_wrapped_text(c, "Scope Overview: " + (scope[:1500]), margin, y, width - 2*margin, leading=12)
#     c.showPage(); c.save()

#     return str(json_path), str(docx_path), str(pdf_path)

# # -------------------------
# # Compliance check (faster with prefilter + early exit)
# # -------------------------
# def _window_keyword_meta(w: str) -> Dict[str, bool]:
#     wl = w.lower()
#     return {
#         "payment": ("payment" in wl) or ("ipc" in wl) or ("interim" in wl),
#         "retention": ("retention" in wl),
#         "termination": ("terminat" in wl),
#         "governing law": ("governing law" in wl) or ("law" in wl),
#         "suspension": ("suspens" in wl) or ("non-payment" in wl),
#         "safety": ("ppe" in wl) or ("safety" in wl),
#         "insurance": ("insurance" in wl) or ("car" in wl) or ("workmen" in wl) or ("tpl" in wl),
#         "notices": ("notice" in wl) or ("registered post" in wl) or ("email" in wl),
#         "records": ("record" in wl) or ("document" in wl) or ("retain" in wl),
#         "dispute": ("arbitrat" in wl) or ("court" in wl) or ("jurisdict" in wl),
#     }



# def _json_from_yesno(rule: str, window: str, guess_present: bool, quote: str | None, sources: list[str] | None, conf: float) -> dict:
#     return {
#         "rule": rule,
#         "present": bool(guess_present),
#         "summary": ("Found evidence in window" if guess_present else None),
#         "quote": (quote[:200] if quote else None),
#         "sources": sources or ([] if not guess_present else ["page ?"]),
#         "confidence": float(conf),
#     }

# _payment_re = re.compile(r"(payment|payable)\s+(?:within|in)\s+(\d{1,3})\s+days", re.I)
# _termination_re = re.compile(r"\bterminat(?:e|ion)\b", re.I)
# _govlaw_re = re.compile(r"\bgoverning\s+law\b.*?(india|indian|laws? of [A-Za-z ]+)", re.I | re.S)

# def _clean_quote(q: str) -> str:
#     if not q: return q
#     q = re.sub(r"---\s*PAGE\s+\d+\s*---", " ", q, flags=re.I)
#     q = q.replace("PAGE BREAK", " ")
#     q = re.sub(r"\s+", " ", q).strip()
#     return q[:200]

# def _infer_sources_from_quote(quote: str, pages) -> list[str]:
#     qn = _clean_quote(quote)
#     hits = []
#     for p in pages:
#         if qn and qn in (p.get("text") or ""):
#             hits.append(f"page {p['page']}")
#     return hits


# def _compliance_regex_fallback(rule: str, window: str) -> dict | None:
#     rl = rule.lower()
#     w = window

#     # A) payment terms
#     if "payment" in rl and ("term" in rl or "within" in rl or "days" in rl or "ipc" in rl):
#         m = _payment_re.search(w)
#         if m:
#             quote = m.group(0)
#             return _json_from_yesno(rule, w, True, quote,None, 0.75)

#     # B) termination
#     if "termination" in rl or "terminate" in rl:
#         m = _termination_re.search(w)
#         if m:
#             # try to grab a short line around it
#             line = ""
#             try:
#                 s = max(0, m.start()-60); e = min(len(w), m.end()+120)
#                 line = re.sub(r"\s+", " ", w[s:e]).strip()[:200]
#             except:
#                 pass
#             return _json_from_yesno(rule, w, True, line or m.group(0),None, 0.70)

#     # C) governing law
#     if "governing law" in rl or ("governing" in rl and "law" in rl):
#         m = _govlaw_re.search(w)
#         if m:
#             quote = re.sub(r"\s+", " ", m.group(0))[:200]
#             return _json_from_yesno(rule, w, True, quote,None, 0.80)

#     return None


# def compliance_check_json(full_text_with_pages: str, rules: List[str], llm_compliance: HuggingFacePipeline, tokenizer,
#                           window_tokens: int = DEFAULT_WINDOW_TOKENS, overlap_tokens: int = DEFAULT_OVERLAP_TOKENS):
#     chain = LLMChain(llm=llm_compliance, prompt=COMPLIANCE_JSON_PROMPT)
#     windows = chunk_text_by_tokens(full_text_with_pages, tokenizer, window_tokens, overlap_tokens)

#     metas = [{"w": w, "kw": _window_keyword_meta(w)} for w in windows]

#     results = []
#     for rule in rules:
#         rule_l = rule.lower()
#         candidates = []
#         for m in metas:
#             if any(k in rule_l and v for k, v in m["kw"].items()):
#                 candidates.append(m["w"])
#         if not candidates:
#             candidates = [m["w"] for m in metas]  # fallback

#         aggregated = []
#         for w in candidates:
#             raw = (chain.predict(window_text=w, rule=rule) or "").strip()
#             parsed = None
#             try:
#                 parsed = json.loads(raw)
#             except Exception:
#                 parsed = _loose_json_grab(raw)

#             if not parsed:
#                 # NEW: regex fallback builds a valid JSON if we clearly see evidence
#                 fallback = _compliance_regex_fallback(rule, w)
#                 if fallback:
#                     parsed = fallback

#             if parsed:
#                 aggregated.append(parsed)
#                 try:
#                     if parsed.get("present") and float(parsed.get("confidence", 0.0)) >= 0.80:
#                         break
#                 except Exception:
#                     pass

#         final = {"rule": rule, "present": False, "summary": None, "quote": None, "sources": [], "confidence": 0.0}
#         if final.get("present"):
#             if final.get("quote"):
#                 final["quote"] = _clean_quote(final["quote"])
#             if not final.get("sources"):
#                 final["sources"] = _infer_sources_from_quote(final.get("quote"), pages)
#         for a in aggregated:
#             if a.get("present"):
#                 final["present"] = True
#                 final["summary"] = a.get("summary") or final["summary"]
#                 final["quote"] = a.get("quote") or final["quote"]
#                 final["sources"].extend([s for s in a.get("sources", []) if s not in final["sources"]])
#                 try:
#                     final["confidence"] = max(final["confidence"], float(a.get("confidence", 0.0)))
#                 except Exception:
#                     pass
#         results.append(final)
#     return results


# # -------------------------
# # Conflict detection (regex)
# # -------------------------
# def check_commencement_vs_site_possession(pages, full_text):
#     date_pattern = r'((?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4})|(?:\d{1,2}\s+[A-Za-z]+\s+\d{4})|(?:[A-Za-z]+\s+\d{1,2},\s*\d{4}))'
#     site_pos_pattern = re.compile(r'(site possession|possession of the site|hand over (?:the )?site)[^\.\n]{0,250}(' + date_pattern + r')', re.I)
#     commence_pattern = re.compile(r'(commenc(?:e|ement) date|start date|date of commencement|date of commencement of works)[^\.\n]{0,250}(' + date_pattern + r')', re.I)

#     site_matches, comm_matches = [], []
#     for p in pages:
#         text = p.get('text') or ""
#         for m in site_pos_pattern.finditer(text):
#             raw = m.group(2)
#             try:
#                 dt = dateparser.parse(raw, dayfirst=True)
#             except Exception:
#                 dt = None
#             site_matches.append({"page": p['page'], "raw": raw, "date": dt, "snippet": m.group(0)})
#         for m in commence_pattern.finditer(text):
#             raw = m.group(2)
#             try:
#                 dt = dateparser.parse(raw, dayfirst=True)
#             except Exception:
#                 dt = None
#             comm_matches.append({"page": p['page'], "raw": raw, "date": dt, "snippet": m.group(0)})

#     conflicts = []
#     if site_matches and comm_matches:
#         comm_dates = [c['date'] for c in comm_matches if c['date']]
#         site_dates = [s['date'] for s in site_matches if s['date']]
#         if comm_dates and site_dates:
#             min_comm = min(comm_dates)
#             max_site = max(site_dates)
#             if max_site > min_comm:
#                 conflicts.append({
#                     "type": "commencement_vs_site_possession",
#                     "severity": "Critical",
#                     "message": f"Site possession ({max_site.date()}) occurs after commencement ({min_comm.date()}).",
#                     "evidence": {"commencement": comm_matches, "site_possession": site_matches}
#                 })
#     return conflicts

# def check_payment_term_mismatch(pages):
#     pay_pattern = re.compile(r'pay(?:ment)?\s+(?:within|in)\s+(\d{1,3})\s+days', re.I)
#     matches = []
#     for p in pages:
#         for m in pay_pattern.finditer(p.get('text') or ""):
#             try:
#                 matches.append({"page": p['page'], "days": int(m.group(1)), "snippet": m.group(0)})
#             except Exception:
#                 pass
#     conflicts = []
#     if matches:
#         days_set = sorted(set(m['days'] for m in matches))
#         if len(days_set) > 1:
#             conflicts.append({
#                 "type": "payment_term_mismatch",
#                 "severity": "High",
#                 "message": f"Different payment terms found: {days_set} days.",
#                 "evidence": matches
#             })
#     return conflicts

# def check_retention_mismatch(pages):
#     ret_pattern = re.compile(r'retention[^\d]{0,20}(\d{1,2})\s*%?', re.I)
#     matches = []
#     for p in pages:
#         for m in ret_pattern.finditer(p.get('text') or ""):
#             try:
#                 pct = int(m.group(1))
#                 matches.append({"page": p['page'], "pct": pct, "snippet": m.group(0)})
#             except Exception:
#                 continue
#     conflicts = []
#     if matches:
#         pct_set = sorted(set(m['pct'] for m in matches))
#         if len(pct_set) > 1:
#             conflicts.append({
#                 "type": "retention_mismatch",
#                 "severity": "High",
#                 "message": f"Different retention percentages found: {pct_set}%.",
#                 "evidence": matches
#             })
#     return conflicts

# def check_defect_liability_mismatch(pages):
#     dl_pattern = re.compile(r'(defect liability period|defects liability|warranty period)[^\d]{0,60}(\d{1,3})\s*(months|month|years|year)?', re.I)
#     matches = []
#     for p in pages:
#         for m in dl_pattern.finditer(p.get('text') or ""):
#             try:
#                 val = int(m.group(2))
#                 unit = (m.group(3) or "months").lower()
#                 months = val * 12 if "year" in unit else val
#                 matches.append({"page": p['page'], "value": months, "raw": m.group(0)})
#             except Exception:
#                 continue
#     conflicts = []
#     if matches:
#         vals = sorted(set(m['value'] for m in matches))
#         if len(vals) > 1:
#             conflicts.append({
#                 "type": "defect_liability_mismatch",
#                 "severity": "High",
#                 "message": f"Different defect liability lengths found (in months): {vals}.",
#                 "evidence": matches
#             })
#     return conflicts

# def check_arbitration_vs_court_conflict(pages):
#     arb_pattern = re.compile(r'\barbitrat', re.I)
#     court_pattern = re.compile(r'\bcourt', re.I)
#     arbs, courts = [], []
#     for p in pages:
#         t = p.get('text') or ""
#         if arb_pattern.search(t):
#             arbs.append({"page": p['page']})
#         if court_pattern.search(t):
#             courts.append({"page": p['page']})
#     conflicts = []
#     if arbs and courts:
#         conflicts.append({
#             "type": "dispute_resolution_conflict",
#             "severity": "Critical",
#             "message": "Arbitration-related terms and court-related terms both found; possible conflict in dispute resolution or interim relief.",
#             "evidence": {"arbitration_pages": arbs, "court_pages": courts}
#         })
#     return conflicts

# CONFLICT_TO_CATEGORY = {
#     "commencement_vs_site_possession": ("Contractual Disputes (Site Possession)", "Consider deeming commencement as actual possession or grant EOT + cost recovery. See 'SITE POSSESSION' clause."),
#     "payment_term_mismatch": ("Contractual Disputes (Payment)", "Harmonize payment timeline; prefer main-body or latest-dated doc. Consider change-order for invoices/payments."),
#     "retention_mismatch": ("Financial / Payment", "Clarify retention % in payment schedule; propose single retained % and update annex."),
#     "defect_liability_mismatch": ("Quality / Warranty", "Adopt the stricter DLP or clarify which schedule governs; update annex."),
#     "dispute_resolution_conflict": ("Dispute Resolution", "Resolve arbitration vs court language — prefer arbitration clause with interim relief carve-out."),
# }

# def map_conflict_to_practical_category(conflict):
#     ctype = conflict.get("type")
#     cat, hint = CONFLICT_TO_CATEGORY.get(ctype, ("Other / Administrative", "Review conflict and map manually."))
#     conflict["category"] = cat
#     conflict["resolution_hint"] = hint
#     return conflict

# def run_conflict_detection(pages, full_text):
#     conflicts = []
#     conflicts.extend(check_commencement_vs_site_possession(pages, full_text))
#     conflicts.extend(check_payment_term_mismatch(pages))
#     conflicts.extend(check_retention_mismatch(pages))
#     conflicts.extend(check_defect_liability_mismatch(pages))
#     conflicts.extend(check_arbitration_vs_court_conflict(pages))
#     conflicts = [map_conflict_to_practical_category(c) for c in conflicts]
#     return conflicts

# # -------------------------
# # Pipeline nodes
# # -------------------------
# def node_ocr_and_extract(pdf_bytes: bytes) -> Dict[str, Any]:
#     st.info("Extracting text from PDF (pdfplumber preferred; fallback to OCR).")
#     pages, full_text = extract_text_prefer_pdfplumber(pdf_bytes, dpi=200)
#     return {"pages": pages, "full_text": full_text}

# def node_chunk_text_from_full(raw_text: str) -> List[str]:
#     st.info("Chunking text for UI and downstream (char-based chunks).")
#     return simple_chunk_text(raw_text, size=3000, overlap=500)

# def node_build_model_wrappers(_) -> Dict[str, Any]:
#     st.info("Loading model and LangChain wrappers...")
#     tokenizer, llm_main, llm_comp = build_generation_pipelines(GEN_MODEL)
#     return {"tokenizer": tokenizer, "llm_main": llm_main, "llm_compliance": llm_comp}

# # -------------------------
# # Streamlit UI
# # -------------------------
# st.set_page_config(page_title="Contract CAD Chatbot - LangChain Direct LLM", layout="wide")
# st.title("Automated Contract CAD Generator")

# # Sidebar controls
# with st.sidebar:
#     st.header("Controls")
#     strategy = st.radio("Query strategy", options=["sliding", "concatenate"], index=0)
#     window_tokens = st.number_input("Window tokens (approx)", min_value=512, max_value=64000, value=DEFAULT_WINDOW_TOKENS, step=128)
#     overlap_tokens = st.number_input("Overlap tokens", min_value=0, max_value=max(1, window_tokens//2), value=DEFAULT_OVERLAP_TOKENS, step=32)
#     reserved_qa_tokens = st.number_input("Reserved tokens for QA (generation)", min_value=64, max_value=4096, value=DEFAULT_RESERVED_QA_TOKENS, step=64)
#     st.markdown("---")
#     uploaded_file = st.file_uploader("Upload contract PDF (scanned or digital)", type=["pdf"])
#     regen = st.button("Reprocess / Rebuild context")
#     clear_all = st.button("Clear session")

# if clear_all:
#     for k in list(st.session_state.keys()):
#         del st.session_state[k]
#     st.experimental_rerun()

# # Graph nodes & pipeline
# graph = SimpleGraph()
# graph.add_node(Node("OCR", lambda b: node_ocr_and_extract(b)))
# graph.add_node(Node("Chunk", lambda info: node_chunk_text_from_full(info["full_text"])))
# graph.add_node(Node("Model", node_build_model_wrappers))

# # Build pipeline & process upload
# if uploaded_file and (regen or not st.session_state.get("graph_built", False)):
#     try:
#         with st.spinner("Running pipeline (extract → chunk → model load)..."):
#             pdf_bytes = uploaded_file.read()
#             extract_info = node_ocr_and_extract(pdf_bytes)
#             pages = extract_info["pages"]
#             full_text = extract_info["full_text"]  # includes --- PAGE N --- markers
#             chunks = node_chunk_text_from_full(full_text)
#             model_info = node_build_model_wrappers(None)
#             st.session_state["pages"] = pages
#             st.session_state["raw_text"] = full_text
#             st.session_state["chunks"] = chunks
#             st.session_state["tokenizer"] = model_info["tokenizer"]
#             st.session_state["llm_main"] = model_info["llm_main"]
#             st.session_state["llm_compliance"] = model_info["llm_compliance"]
#             st.session_state["graph_built"] = True
#         st.success("Processed PDF and loaded model.")
#     except Exception as e:
#         st.exception(e)
#         st.error("Processing failed.")

# # Setup memory-based chain for chat (LangChain Conversation)
# if st.session_state.get("graph_built", False) and "chat_memory" not in st.session_state:
#     st.session_state["chat_memory"] = ConversationBufferMemory(memory_key="chat_history", return_messages=True)

# # Layout: Chat (left) and CAD/Compliance (right)
# col1, col2 = st.columns([2, 1])

# with col1:
#     st.subheader("Chat with the Contract")
#     if not st.session_state.get("graph_built", False):
#         st.info("Upload & Reprocess a PDF to enable chat.")
#     else:
#         if "conversation" not in st.session_state:
#             st.session_state["conversation"] = []
#         for turn in st.session_state["conversation"]:
#             if turn["role"] == "user":
#                 st.chat_message("user").write(turn["text"])
#             else:
#                 st.chat_message("assistant").write(turn["text"])

#         user_input = st.chat_input("Ask about the contract (e.g., 'What is the payment clause?')")
#         if user_input:
#             st.session_state["conversation"].append({"role":"user","text":user_input})
#             tokenizer = st.session_state["tokenizer"]
#             llm_main = st.session_state["llm_main"]
#             full_text_for_windows = st.session_state["raw_text"]  # KEEP PAGE HEADERS
#             chain_memory = st.session_state["chat_memory"]

#             with st.spinner("Asking LLM..."):
#                 res = ask_direct_langchain(
#                     llm_main, tokenizer, full_text_for_windows, user_input,
#                     strategy=strategy, window_tokens=window_tokens, overlap_tokens=overlap_tokens, memory=chain_memory
#                 )
#                 answer = res.get("answer", "No answer.")
#             st.session_state["conversation"].append({"role":"assistant","text":answer})
#             st.chat_message("assistant").write(answer)
#             if st.checkbox("Show debug windows & answers", key=f"dbg_chat_{len(st.session_state['conversation'])}"):
#                 st.write(res)

# with col2:
#     st.subheader("CAD Generator / Document")
#     if not st.session_state.get("graph_built", False):
#         st.info("Upload & Reprocess PDF to enable CAD generation.")
#     else:
#         if st.button("Generate CAD (JSON + DOCX + PDF)"):
#             try:
#                 llm_main = st.session_state["llm_main"]
#                 tokenizer = st.session_state["tokenizer"]
#                 full_text = st.session_state["raw_text"]   # keep page markers
#                 pages = st.session_state["pages"]
#                 with st.spinner("Generating CAD (JSON/DOCX/PDF)..."):
#                     json_path, docx_path, pdf_path = node_generate_cad_json_docx_pdf(
#                         llm_main, tokenizer, full_text, pages,
#                         output_basename="generated_CAD",
#                         window_tokens=window_tokens, overlap_tokens=overlap_tokens
#                     )
#                 st.session_state["cad_json"] = json_path
#                 st.session_state["cad_docx"] = docx_path
#                 st.session_state["cad_pdf"] = pdf_path
#                 st.success("CAD generated as JSON, DOCX and PDF.")
#             except Exception as e:
#                 st.exception(e)
#                 st.error("CAD generation failed.")

#         if st.session_state.get("cad_json"):
#             try:
#                 with open(st.session_state["cad_json"], "rb") as f:
#                     st.download_button("Download CAD JSON", f.read(), file_name=Path(st.session_state['cad_json']).name)
#             except Exception as e:
#                 st.error(f"Failed to open CAD JSON: {e}")
#         if st.session_state.get("cad_docx"):
#             try:
#                 with open(st.session_state["cad_docx"], "rb") as f:
#                     st.download_button("Download CAD DOCX", f.read(), file_name=Path(st.session_state['cad_docx']).name)
#             except Exception as e:
#                 st.error(f"Failed to open CAD DOCX: {e}")
#         if st.session_state.get("cad_pdf"):
#             try:
#                 with open(st.session_state["cad_pdf"], "rb") as f:
#                     st.download_button("Download CAD PDF", f.read(), file_name=Path(st.session_state['cad_pdf']).name)
#             except Exception as e:
#                 st.error(f"Failed to open CAD PDF: {e}")

#     st.markdown("---")
#     st.subheader("📑 Compliance Check")
#     if not st.session_state.get("graph_built", False):
#         st.info("Upload a contract and press Reprocess first.")
#     else:
#         default_rules = "Does the contract specify payment terms?\nIs there a termination clause?\nWhat is the governing law?"
#         rules_input = st.text_area("Enter compliance rules (one per line)", value=default_rules)
#         if st.button("Run Compliance Check"):
#             rules = [r.strip() for r in rules_input.splitlines() if r.strip()]
#             llm_compliance = st.session_state["llm_compliance"]
#             tokenizer = st.session_state["tokenizer"]
#             full_text = st.session_state["raw_text"]
#             with st.spinner("Running compliance checks..."):
#                 results = compliance_check_json(
#                     full_text, rules, llm_compliance, tokenizer,
#                     window_tokens=window_tokens, overlap_tokens=overlap_tokens
#                 )
#             st.subheader("Compliance Report")
#             df_rows = []
#             for r in results:
#                 st.markdown(f"**Rule:** {r['rule']}  \n**Present:** {r['present']}  \n**Summary:** {r['summary']}  \n**Quote:** {r['quote']}  \n**Sources:** {r['sources']}  \n**Confidence:** {r['confidence']}")
#                 df_rows.append([r['rule'], r['present'], r['summary'], r['quote'], ", ".join(r['sources']), r['confidence']])
#             if df_rows:
#                 st.dataframe(pd.DataFrame(df_rows, columns=["Rule", "Present", "Summary", "Quote", "Sources", "Confidence"]))

#     st.markdown("---")
#     st.subheader("⚠️ Conflict Detection")
#     if not st.session_state.get("graph_built", False):
#         st.info("Upload a contract and press Reprocess first.")
#     else:
#         if st.button("Run Conflict Detection"):
#             try:
#                 pages = st.session_state.get("pages")
#                 full_text = st.session_state.get("raw_text")
#                 with st.spinner("Detecting conflicts..."):
#                     conflicts = run_conflict_detection(pages, full_text)
#                 st.session_state["conflicts"] = conflicts
#                 if not conflicts:
#                     st.success("No conflicts detected by automated checks.")
#                 else:
#                     st.warning(f"{len(conflicts)} potential conflicts detected.")
#                     for idx, c in enumerate(conflicts, start=1):
#                         st.markdown(f"### {idx}. Category: **{c.get('category')}**  \n**Type:** {c['type']}  \n**Severity:** {c['severity']}  \n**Message:** {c['message']}")
#                         st.markdown(f"**Suggested quick fix:** {c.get('resolution_hint')}")
#                         st.json(c.get("evidence"))
#             except Exception as e:
#                 st.exception(e)
#                 st.error("Conflict detection failed.")

#     st.markdown("---")
#     st.subheader("Raw text & pages preview")
#     if st.session_state.get("raw_text"):
#         if st.checkbox("Show OCR/raw extracted text (first 15000 chars)"):
#             st.text_area("Extracted Text", value=st.session_state["raw_text"][:15000], height=400)
#     else:
#         st.info("No extracted text available. Upload a PDF and press Reprocess.")

# st.caption("LangChain + HuggingFace. Outputs: CAD JSON, DOCX, PDF. Faster compliance. Page-aware windows for better sources.")
# # End of file
# app_streamlit_cad.py
import os
import io
import json
import textwrap
import tempfile
from pathlib import Path
from typing import List, Dict, Any, Optional
import logging
import re
from collections import defaultdict
import streamlit as st
import pdfplumber                    # pip install pdfplumber
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas  # pip install reportlab
from dateutil import parser as dateparser  # pip install python-dateutil

from pdf2image import convert_from_bytes
from PIL import Image
import pytesseract
import shutil
import torch
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
from docx import Document as DocxDocument
from docx.shared import Pt
import pandas as pd

HF_TOKEN="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"

# LangChain
from langchain.chains import LLMChain
from langchain_core.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory
from langchain_huggingface import HuggingFacePipeline

# Optional: login to HF Hub via secrets/env (no hardcoded token!)
try:
    from huggingface_hub import login as hf_login
    hf_token = st.secrets.get("HF_TOKEN", os.environ.get("HF_TOKEN"))
    if hf_token:
        try:
            hf_login(token=hf_token)
        except Exception as e:
            st.warning(f"HuggingFace login failed: {e}")
except Exception:
    pass

try:
    GEN_MODEL = None
    try:
        GEN_MODEL = st.secrets.get("GEN_MODEL")
    except Exception:
        GEN_MODEL = None
    if not GEN_MODEL:
        GEN_MODEL = os.environ.get("GEN_MODEL", "google/flan-t5-base")  # faster default for CPU
except Exception:
    GEN_MODEL = os.environ.get("GEN_MODEL", "google/flan-t5-base")

DEFAULT_WINDOW_TOKENS = 512  # smaller windows for speed/safety
DEFAULT_OVERLAP_TOKENS = 128
DEFAULT_RESERVED_QA_TOKENS = 512
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200

# Ensure tesseract path
tesseract_path = shutil.which("tesseract")
if tesseract_path:
    pytesseract.pytesseract.tesseract_cmd = tesseract_path

# -------------------------
# Small helper structures
# -------------------------
class Node:
    def __init__(self, name: str, fn):
        self.name = name
        self.fn = fn
        self.outputs = None
    def run(self, *args, **kwargs):
        self.outputs = self.fn(*args, **kwargs)
        return self.outputs

class SimpleGraph:
    def __init__(self):
        self.nodes = []
    def add_node(self, node: Node):
        self.nodes.append(node)
    def run(self, start_node_index: int = 0, input=None, **kwargs):
        data = None
        for idx, node in enumerate(self.nodes[start_node_index:], start=start_node_index):
            if idx == start_node_index:
                data = node.run(input, **kwargs)
            else:
                data = node.run(data)
        return data

# -------------------------
# PDF → Text extraction
# -------------------------
def pdf_bytes_to_images(pdf_bytes: bytes, dpi: int = 200) -> List[Image.Image]:
    images = convert_from_bytes(pdf_bytes, dpi=dpi)
    return images

def ocr_image_to_text(image: Image.Image) -> str:
    text = pytesseract.image_to_string(image)
    return text

def extract_text_from_pdf_bytes(pdf_bytes: bytes, dpi: int = 200) -> str:
    pages = pdf_bytes_to_images(pdf_bytes, dpi=dpi)
    texts = []
    for i, img in enumerate(pages):
        page_text = ocr_image_to_text(img)
        header = f"\n\n--- PAGE {i+1} ---\n\n"
        texts.append(header + page_text)
    return "\n".join(texts)

def extract_text_prefer_pdfplumber(pdf_bytes: bytes, dpi: int = 200):
    """
    Use pdfplumber to extract selectable text per page; fallback to OCR for pages with no text.
    Returns: (pages_list, full_text)
      pages_list: [{'page': n, 'text': str, 'is_ocr': bool}, ...]
      full_text: a single string with --- PAGE n --- markers
    """
    pages = []
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                try:
                    txt = page.extract_text() or ""
                except Exception:
                    txt = ""
                if txt and txt.strip():
                    pages.append({"page": i, "text": txt, "is_ocr": False})
                else:
                    try:
                        pil_img = page.to_image(resolution=dpi).original
                        ocr_txt = pytesseract.image_to_string(pil_img)
                    except Exception:
                        ocr_txt = ""
                    pages.append({"page": i, "text": ocr_txt, "is_ocr": True})
    except Exception as e:
        logging.warning("pdfplumber failed: %s. Falling back to full OCR.", e)
        ocr_full = extract_text_from_pdf_bytes(pdf_bytes, dpi=dpi)
        page_texts = []
        if "\f" in ocr_full:
            page_texts = ocr_full.split("\f")
        else:
            page_texts = [ocr_full]
        for idx, p_text in enumerate(page_texts, start=1):
            pages.append({"page": idx, "text": p_text, "is_ocr": True})

    full = []
    for p in pages:
        full.append(f"\n\n--- PAGE {p['page']} ---\n{p['text']}")
    return pages, "\n".join(full)

def simple_chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    if not text:
        return []
    joined = text.replace("\r\n", "\n")
    chunks = []
    start = 0
    while start < len(joined):
        end = min(len(joined), start + size)
        chunks.append(joined[start:end])
        if end == len(joined):
            break
        start = max(0, end - overlap)
    return chunks

# -------------------------
# Token-based windowing
# -------------------------
def count_tokens(text: str, tokenizer) -> int:
    return len(tokenizer.encode(text, add_special_tokens=False))

def chunk_text_by_tokens(text: str, tokenizer, window_tokens: int, overlap_tokens: int) -> List[str]:
    tokens = tokenizer.encode(text, add_special_tokens=False)
    if len(tokens) <= window_tokens:
        return [tokenizer.decode(tokens, skip_special_tokens=True)]
    windows = []
    start = 0
    step = max(1, window_tokens - overlap_tokens)
    while start < len(tokens):
        end = min(start + window_tokens, len(tokens))
        slice_ids = tokens[start:end]
        windows.append(tokenizer.decode(slice_ids, skip_special_tokens=True))
        if end == len(tokens):
            break
        start += step
    return windows

def make_full_contract_text(chunks: List[str], include_chunk_headers: bool = True) -> str:
    parts = []
    for i, c in enumerate(chunks):
        if include_chunk_headers:
            parts.append(f"\n\n--- CHUNK {i+1} ---\n{c}")
        else:
            parts.append(c)
    return "\n\n".join(parts)

# -------------------------
# Prompts
# -------------------------
CONTRACT_Q_PROMPT = PromptTemplate(
    input_variables=["contract", "question"],
    template=textwrap.dedent("""\
        You are a contract analyst. Use ONLY the CONTRACT text below to answer the question.
        If the contract does not contain the answer, reply exactly: "Not found in contract."

        CONTRACT:
        {contract}

        QUESTION:
        {question}

        ANSWER:""")
)

WINDOW_Q_PROMPT = PromptTemplate(
    input_variables=["window", "question", "idx", "total"],
    template=textwrap.dedent("""\
        You are a contract analyst. Use ONLY the CONTRACT WINDOW below to answer the question.
        If the window does not contain the answer, reply exactly: "Not found in contract."

        CONTRACT WINDOW {idx}/{total}:
        {window}

        QUESTION:
        {question}

        ANSWER:""")
)

def _sanitize_json_like(txt: str) -> dict | None:
    """Try to coerce model output to strict JSON."""
    if not txt:
        return None
    s = txt.find('{'); e = txt.rfind('}')
    if s == -1 or e == -1:
        return None
    cand = txt[s:e+1]

    # Common fixes
    cand = re.sub(r",\s*}", "}", cand)
    cand = re.sub(r",\s*]", "]", cand)
    # ensure true/false/null are lowercase (T5 usually okay)
    cand = re.sub(r"\bTrue\b", "true", cand)
    cand = re.sub(r"\bFalse\b", "false", cand)
    cand = re.sub(r"\bNone\b", "null", cand)

    try:
        return json.loads(cand)
    except Exception:
        # attempt to add quotes to keys like: key: value -> "key": value (only for top-level-ish places)
        def _quote_keys(m):
            return f"\"{m.group(1)}\":"
        cand2 = re.sub(r'(?m)^\s*([A-Za-z0-9_]+)\s*:', _quote_keys, cand)
        try:
            return json.loads(cand2)
        except Exception:
            return None

REPAIR_JSON_PROMPT = PromptTemplate(
    input_variables=["bad_text"],
    template=textwrap.dedent("""\
        Convert the following content into ONE strict JSON object that matches the CAD schema.
        Only output the JSON (start with '{{' and end with '}}').

        CONTENT:
        {bad_text}
        """)
)

def _cad_candidate_window_score(w: str) -> int:
    wl = w.lower()
    score = 0
    # favor windows with these signals:
    for kw in ("employer", "contractor", "contract price", "payment", "retention", "mobilization", "bank guarantee",
               "defects liability", "completion", "commencement", "arbitration", "governing law", "notice", "submittal", "method statement"):
        if kw in wl:
            score += 1
    return score

def _build_minimal_cad_from_text(text: str) -> dict:
    """Last-resort minimal JSON using regex so we never hard-fail."""
    t = text
    out = {
      "salient_features": {
        "project_name": None, "project_location": None, "contract_ref": None,
        "employer": {"name": None, "address": None, "contact": None, "sources": []},
        "contractor": {"name": None, "address": None, "contact": None, "sources": []},
        "contract_type": None,
        "contract_value": {"amount": None, "currency": None, "source": []},
        "price_escalation": {"applicable": False, "index": None, "formula": None, "source": []},
        "award_date": None, "commencement_date": None, "scheduled_completion_date": None,
        "defect_liability_period": None, "scope_overview": None, "sources": []
      },
      "important_submittals": [],
      "notice_clauses": [],
      "project_progress_clauses": [],
      "payment": {
        "contract_sum": {"amount": None, "currency": None, "source": []},
        "mobilization_advance": {"applicable": False, "amount_pct": None, "security": None, "recovery": None},
        "interim_payments": [],
        "retention": {"percent": None, "release_condition": "", "source": []},
        "final_payment": {"required_docs": [], "release_days_after_submission": None},
        "escalation_clause": {"applicable": False, "details": "", "sources": []},
        "sources": []
      },
      "risks_and_allocation": [],
      "claims_disputes_arbitration": {"arbitration": {"applicable": False, "notes": "", "sources": []}, "dispute_forum": None, "claims_summary": [], "sources": []}
    }

    # Employer/Contractor names (very naive)
    m_emp = re.search(r"Employer[:\s\-]+([^\n,]+)", t, re.I)
    if m_emp: out["salient_features"]["employer"]["name"] = m_emp.group(1).strip()
    m_con = re.search(r"Contractor[:\s\-]+([^\n,]+)", t, re.I)
    if m_con: out["salient_features"]["contractor"]["name"] = m_con.group(1).strip()

    # Contract price
    m_price_num = re.search(r"Contract Price[:\s\-]+(?:INR|₹)?\s*([\d,]+)", t, re.I)
    if m_price_num:
        try:
            amt = int(m_price_num.group(1).replace(",", ""))
            out["payment"]["contract_sum"]["amount"] = amt
            out["payment"]["contract_sum"]["currency"] = "INR"
            out["salient_features"]["contract_value"]["amount"] = amt
            out["salient_features"]["contract_value"]["currency"] = "INR"
        except: pass

    # Payment days (interim)
    m_pay = re.search(r"(?:Interim Payments?|payable)\s.*?\bwithin\s+(\d{1,3})\s+days", t, re.I)
    if m_pay:
        try:
            out["payment"]["interim_payments"] = [{"frequency":"monthly","certifier":"Engineer","payment_days_after_certification": int(m_pay.group(1))}]
        except: pass

    # Retention percent
    m_ret = re.search(r"Retention(?: money)?\s*(?:of)?\s*(\d{1,2})\s*%?", t, re.I)
    if m_ret:
        try:
            out["payment"]["retention"]["percent"] = int(m_ret.group(1))
        except: pass

    # DLP
    m_dlp = re.search(r"(?:Defects? Liability.*?)(\d{1,3})\s*(months|month|years|year)", t, re.I)
    if m_dlp:
        val = int(m_dlp.group(1)); unit = m_dlp.group(2).lower()
        out["salient_features"]["defect_liability_period"] = f"{val} {unit}"

    # Arbitration/court
    if "arbitration" in t.lower():
        out["claims_disputes_arbitration"]["arbitration"]["applicable"] = True

    return out

CAD_JSON_PROMPT = PromptTemplate(
    input_variables=["window_text"],
    template=textwrap.dedent("""\
        You are a contract analyst. Use ONLY the CONTRACT TEXT provided below. Produce ONE valid JSON object EXACTLY matching the schema. 
        IMPORTANT: Output STRICT JSON only — no prose, no Markdown. Your output MUST start with '{{' and end with '}}'.

        CONTRACT WINDOW:
        {window_text}

        SCHEMA:
        {{
          "salient_features": {{
            "project_name": "string | null",
            "project_location": "string | null",
            "contract_ref": "string | null",
            "employer": {{"name":"string | null","address":"string | null","contact":"string | null","sources":["page X"]}},
            "contractor": {{"name":"string | null","address":"string | null","contact":"string | null","sources":["page X"]}},
            "contract_type": "string | null",
            "contract_value": {{"amount":number | null,"currency":"string | null","source":["page X"]}},
            "price_escalation": {{"applicable": true|false, "index":"string|null", "formula":"string|null", "source":["page X"]}},
            "award_date": "YYYY-MM-DD | null",
            "commencement_date": "YYYY-MM-DD | null",
            "scheduled_completion_date": "YYYY-MM-DD | null",
            "defect_liability_period": "string | null",
            "scope_overview": "string | null",
            "sources": ["page X"]
          }},
          "important_submittals": [
            {{"stage":"before|during|after|null","document":"string","due":"string","notes":"string","sources":["page X"]}}
          ],
          "notice_clauses": [
            {{"event":"string","notifier":"string","recipient":"string","timeline":"string","method":"string","sources":["page X"]}}
          ],
          "project_progress_clauses": [
            {{"topic":"string","summary":"string","sources":["page X","clause Y"]}}
          ],
          "payment": {{
            "contract_sum": {{"amount": number | null, "currency":"string | null", "source":["page X"]}},
            "mobilization_advance": {{"applicable": true|false,"amount_pct": number | null,"security":"string|null","recovery":"string|null"}},
            "interim_payments": [{{"frequency":"string","certifier":"string","payment_days_after_certification":number | null}}],
            "retention": {{"percent": number | null,"release_condition":"string","source":["page X"]}},
            "final_payment": {{"required_docs":["string"],"release_days_after_submission": number | null}},
            "escalation_clause": {{"applicable": true|false,"details":"string","sources":["page X"]}},
            "sources": ["page X"]
          }},
          "risks_and_allocation": [
            {{"risk":"string","severity":"Major|High|Medium|Low","probability":"Likely|Possible|Rare","responsibility":"Employer|Contractor|Shared","notes":"string","sources":["page X"]}}
          ],
          "claims_disputes_arbitration": {{
            "arbitration": {{"applicable": true|false,"notes":"string","sources":["page X"]}},
            "dispute_forum":"string|null",
            "claims_summary":[{{"topic":"string","process":"string","sources":["page X"]}}],
            "sources":["page X"]
          }}
        }}

        EXAMPLE FORMAT (structure only):
        {{"salient_features": {{"project_name": null, "sources":[]}}, "important_submittals": [], "notice_clauses": [], "project_progress_clauses": [], "payment": {{"sources":[]}}, "risks_and_allocation": [], "claims_disputes_arbitration": {{"sources":[]}}}}
        """)
)

COMPLIANCE_JSON_PROMPT = PromptTemplate(
    input_variables=["window_text", "rule"],
    template=textwrap.dedent("""\
        You are a contract analyst. Use ONLY the CONTRACT WINDOW below to answer the compliance rule.
        Return STRICT JSON ONLY — no prose, no markdown, no surrounding text.

        CONTRACT WINDOW:
        {window_text}

        RULE:
        {rule}

        VALID JSON EXAMPLES (copy the structure):
        # Example present=true
        {{"rule":"sample rule", "present": true, "summary":"short", "quote":"verbatim (<=200 chars)", "sources":["page X","clause Y"], "confidence":0.90}}

        # Example present=false
        {{"rule":"sample rule", "present": false, "summary": null, "quote": null, "sources": [], "confidence": 0.99}}

        INSTRUCTIONS:
        - If the window contains evidence answering the rule, return:
          {{"rule":"{rule}", "present": true, "summary":"short summary", "quote":"verbatim quote (<=200 chars)", "sources":["page X","clause Y"], "confidence":0.90}}
        - If not present, return:
          {{"rule":"{rule}", "present": false, "summary": null, "quote": null, "sources": [], "confidence": 0.99}}
        - STRICT JSON ONLY.
        """)
)

# -------------------------
# Generation pipelines (MAIN + COMPLIANCE)
# -------------------------
@st.cache_resource(show_spinner=False)
def build_generation_pipelines(model_name: str = GEN_MODEL):
    tokenizer = AutoTokenizer.from_pretrained(model_name, use_fast=True)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
    device = 0 if torch.cuda.is_available() else -1

    # MAIN: long outputs (CAD JSON, rich QA)
    main_gen_kwargs = dict(
        max_new_tokens=1024,
        temperature=0.0,
        do_sample=False,
        top_p=1.0,
        repetition_penalty=1.05,
        no_repeat_ngram_size=3,
        return_full_text=False,
        truncation=True,
    )
    # COMPLIANCE: short/fast outputs
    comp_gen_kwargs = dict(
        max_new_tokens=256,
        temperature=0.0,
        do_sample=False,
        top_p=1.0,
        repetition_penalty=1.05,
        no_repeat_ngram_size=3,
        return_full_text=False,
        truncation=True,
    )

    trans_pipe = pipeline(
        "text2text-generation",
        model=model,
        tokenizer=tokenizer,
        device=device if device >= 0 else -1
    )

    llm_main = HuggingFacePipeline(pipeline=trans_pipe, model_kwargs=main_gen_kwargs)
    llm_compliance = HuggingFacePipeline(pipeline=trans_pipe, model_kwargs=comp_gen_kwargs)

    return tokenizer, llm_main, llm_compliance

# -------------------------
# JSON parsing helpers
# -------------------------
def _loose_json_grab(txt: str) -> Optional[dict]:
    if not txt:
        return None
    s = txt.find('{'); e = txt.rfind('}')
    if s == -1 or e == -1:
        return None
    cand = txt[s:e+1]
    cand = re.sub(r",\s*}", "}", cand)
    cand = re.sub(r",\s*]", "]", cand)
    try:
        return json.loads(cand)
    except Exception:
        return None

# -------------------------
# QA helpers
# -------------------------
def ask_with_concatenate_chain(llm_chain: LLMChain, contract_text: str, question: str) -> str:
    return llm_chain.predict(contract=contract_text, question=question)

def ask_with_sliding_chain(llm_chain: LLMChain, tokenizer, contract_text: str, question: str,
                           window_tokens: int = DEFAULT_WINDOW_TOKENS, overlap_tokens: int = DEFAULT_OVERLAP_TOKENS) -> Dict[str, Any]:
    windows = chunk_text_by_tokens(contract_text, tokenizer, window_tokens, overlap_tokens)
    answers = []
    for i, w in enumerate(windows, start=1):
        out = llm_chain.predict(window=w, question=question, idx=i, total=len(windows))
        answers.append({"window_index": i-1, "answer": (out or "").strip()})
    non_empty = [a for a in answers if a["answer"] and "not found in contract" not in a["answer"].lower()]
    if not non_empty:
        longest = max(answers, key=lambda x: len(x["answer"] or ""))
        return {"answer": (longest["answer"] or "Not found in contract.").strip(), "used_windows": [longest["window_index"]], "method": "sliding", "all_window_answers": answers}
    freq = {}
    for a in non_empty:
        key = a["answer"].strip()
        freq[key] = freq.get(key, 0) + 1
    best_text = max(freq.items(), key=lambda kv: (kv[1], len(kv[0])))[0]
    used_windows = [a["window_index"] for a in non_empty if a["answer"].strip() == best_text]
    unique_answers = list({a["answer"].strip() for a in non_empty})
    final_ans = best_text
    if len(unique_answers) > 1:
        final_ans = "\n\n--- AGGREGATED ANSWERS FROM MULTIPLE WINDOWS ---\n\n" + "\n\n---\n\n".join(unique_answers)
    return {"answer": final_ans.strip(), "used_windows": used_windows, "method": "sliding", "all_window_answers": answers}

def ask_direct_langchain(llm_wrapper: HuggingFacePipeline, tokenizer, full_text_for_windows: str, question: str, strategy: str = "sliding",
                         window_tokens: int = DEFAULT_WINDOW_TOKENS, overlap_tokens: int = DEFAULT_OVERLAP_TOKENS, memory: Optional[ConversationBufferMemory]=None) -> Dict[str, Any]:
    if strategy == "concatenate":
        chain = LLMChain(llm=llm_wrapper, prompt=CONTRACT_Q_PROMPT, memory=memory)
        ans = chain.predict(contract=full_text_for_windows, question=question)
        return {"answer": (ans or "").strip(), "method": "concatenate", "used_windows": [0]}
    else:
        chain = LLMChain(llm=llm_wrapper, prompt=WINDOW_Q_PROMPT, memory=memory)
        return ask_with_sliding_chain(chain, tokenizer, full_text_for_windows, question, window_tokens=window_tokens, overlap_tokens=overlap_tokens)

# -------------------------
# CAD Generation
# -------------------------
def find_quote_sources(quote: str, pages) -> list:
    if not quote or not quote.strip():
        return []
    quote_norm = re.sub(r'\s+', ' ', quote.strip())[:400]
    sources = []
    for p in pages:
        page_text_norm = re.sub(r'\s+', ' ', (p.get('text') or "").strip())
        if quote_norm in page_text_norm:
            sources.append(f"page {p['page']}")
    if not sources and len(quote_norm) > 40:
        sub = quote_norm[:80]
        for p in pages:
            if sub in (p.get('text') or ""):
                sources.append(f"page {p['page']}")
    return sources

def merge_json_objects(objs):
    result = {}
    for o in objs:
        if not isinstance(o, dict):
            continue
        for k, v in o.items():
            if v in (None, [], "", {}):
                continue
            if k not in result:
                result[k] = v
            else:
                if isinstance(v, list) and isinstance(result[k], list):
                    for it in v:
                        if it not in result[k]:
                            result[k].append(it)
                elif isinstance(v, dict) and isinstance(result[k], dict):
                    for kk, vv in v.items():
                        if vv not in (None, "", [], {}):
                            result[k].setdefault(kk, vv)
                else:
                    pass
    return result

def node_generate_cad_json_docx_pdf(llm_wrapper, tokenizer, full_text_with_pages: str, pages, output_basename="generated_CAD",
                                    window_tokens: int = DEFAULT_WINDOW_TOKENS, overlap_tokens: int = DEFAULT_OVERLAP_TOKENS):
    """Robust CAD generator:
       - picks top-k candidate windows,
       - retries each window,
       - sanitizes/repairs JSON,
       - falls back to minimal JSON if needed.
    """
    # 0) Build windows (page markers preserved)
    windows = chunk_text_by_tokens(full_text_with_pages, tokenizer, window_tokens, overlap_tokens)
    if not windows:
        raise RuntimeError("Empty text windows for CAD.")

    # 1) Rank windows by simple keyword score, keep top-k
    scored = sorted([(i, _cad_candidate_window_score(w), w) for i, w in enumerate(windows)], key=lambda x: x[1], reverse=True)
    candidates = [w for _, _, w in scored[: min(8, max(3, len(scored)//3))]]  # top 3..8 windows

    # 2) Ask model with strict prompt; retry + repair if needed
    chain = LLMChain(llm=llm_wrapper, prompt=CAD_JSON_PROMPT)
    json_responses = []
    for w in candidates:
        ok = False
        for attempt in range(2):  # two attempts per window
            raw = (chain.predict(window_text=w) or "").strip()
            parsed = None
            try:
                parsed = json.loads(raw)
            except Exception:
                parsed = _sanitize_json_like(raw)

            if not parsed:
                # repair step: ask model to convert its own text to JSON
                repair_chain = LLMChain(llm=llm_wrapper, prompt=REPAIR_JSON_PROMPT)
                repaired = (repair_chain.predict(bad_text=raw) or "").strip()
                try:
                    parsed = json.loads(repaired)
                except Exception:
                    parsed = _sanitize_json_like(repaired)

            if parsed and isinstance(parsed, dict):
                json_responses.append(parsed)
                ok = True
                break  # next candidate
        # continue to next candidate regardless

    # 3) If still nothing, fall back to minimal JSON from regex over the best window
    if not json_responses:
        fallback_dict = _build_minimal_cad_from_text(candidates[0])
        json_responses.append(fallback_dict)

    # 4) Merge
    merged = merge_json_objects(json_responses)

    # 5) Attach sources for quotes/summaries where possible
    def attach_sources_for_quotes(container):
        if isinstance(container, dict):
            for k, v in list(container.items()):
                if isinstance(v, str) and len(v) > 20:
                    if any(keyword in k.lower() for keyword in ['quote','scope','summary','overview','text']):
                        srcs = find_quote_sources(v, pages)
                        if srcs:
                            container.setdefault('sources', []).extend([s for s in srcs if s not in container.get('sources', [])])
                elif isinstance(v, (dict, list)):
                    attach_sources_for_quotes(v)
        elif isinstance(container, list):
            for item in container:
                attach_sources_for_quotes(item)
    attach_sources_for_quotes(merged)

    # 6) Save JSON
    json_path = Path(tempfile.gettempdir()) / f"{output_basename}.json"
    with open(json_path, "w", encoding="utf8") as f:
        json.dump(merged, f, indent=2, ensure_ascii=False)

    # 7) DOCX (same as before)
    docx = DocxDocument()
    style = docx.styles['Normal']; font = style.font; font.name='Arial'; font.size = Pt(11)
    docx.add_heading("Contract Appreciation Document (Generated)", level=1)
    docx.add_paragraph("This document is generated automatically. Verify against contract before finalization.")

    sf = merged.get("salient_features", {})
    docx.add_heading("1. Salient Features", level=2)
    if sf:
        for k, v in sf.items():
            if isinstance(v, dict):
                display_val = v.get('name') or v.get('amount') or json.dumps(v)
                src = ", ".join(v.get('sources', [])) if v.get('sources') else ""
                docx.add_paragraph(f"{k}: {display_val} {'(sources: '+src+')' if src else ''}")
            else:
                docx.add_paragraph(f"{k}: {v}")
    else:
        docx.add_paragraph("Not found")

    submittals = merged.get("important_submittals", [])
    if submittals:
        docx.add_heading("2. Important Submittals", level=2)
        cols = ["stage", "document", "due", "notes", "sources"]
        table = docx.add_table(rows=1, cols=len(cols))
        for i, c in enumerate(cols):
            table.rows[0].cells[i].text = c
        for r in submittals:
            rc = table.add_row().cells
            rc[0].text = r.get("stage","")
            rc[1].text = r.get("document","")
            rc[2].text = r.get("due","")
            rc[3].text = r.get("notes","")
            rc[4].text = ", ".join(r.get("sources", []))
    else:
        docx.add_paragraph("No submittals found")

    notices = merged.get("notice_clauses", [])
    if notices:
        docx.add_heading("3. Notice Clauses", level=2)
        for n in notices:
            docx.add_paragraph(f"{n.get('event')} | {n.get('notifier')} -> {n.get('recipient')} | timeline: {n.get('timeline')} | method: {n.get('method')} | sources: {', '.join(n.get('sources',[]))}")

    docx.add_heading("4. Payment Summary", level=2)
    pay = merged.get("payment", {})
    if pay:
        for k, v in pay.items():
            docx.add_paragraph(f"{k}: {json.dumps(v) if isinstance(v,(dict,list)) else str(v)}")
    else:
        docx.add_paragraph("No payment information found")

    docx.add_heading("5. Risks and Allocation", level=2)
    risks = merged.get("risks_and_allocation", [])
    if risks:
        for r in risks:
            docx.add_paragraph(json.dumps(r))
    else:
        docx.add_paragraph("No risks found")

    docx.add_heading("6. Claims, Disputes & Arbitration", level=2)
    docx.add_paragraph(json.dumps(merged.get("claims_disputes_arbitration", {}), indent=2))

    docx_path = Path(tempfile.gettempdir()) / f"{output_basename}.docx"
    docx.save(docx_path)

    # 8) Simple PDF summary (unchanged, keep your existing draw code)
    pdf_path = Path(tempfile.gettempdir()) / f"{output_basename}.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=A4)
    width, height = A4
    margin = 50
    y = height - margin

    def write_wrapped_text(cobj, text, x, y, max_width, leading=12):
        from reportlab.lib.utils import simpleSplit
        lines = simpleSplit(text, "Helvetica", 10, max_width)
        for ln in lines:
            if y < margin + 40:
                cobj.showPage()
                y = height - margin
            cobj.drawString(x, y, ln)
            y -= leading
        return y

    c.setFont("Helvetica-Bold", 14)
    c.drawString(margin, y, "Contract Appreciation Document (Generated)"); y -= 22
    c.setFont("Helvetica", 10)
    proj = sf.get("project_name") if isinstance(sf, dict) else None
    c.drawString(margin, y, f"Project: {proj or 'N/A'}"); y -= 14
    employer_name = sf.get("employer", {}).get("name") if isinstance(sf.get("employer"), dict) else None
    contractor_name = sf.get("contractor", {}).get("name") if isinstance(sf.get("contractor"), dict) else None
    c.drawString(margin, y, f"Employer: {employer_name or 'N/A'}"); y -= 14
    c.drawString(margin, y, f"Contractor: {contractor_name or 'N/A'}"); y -= 18
    scope = sf.get("scope_overview") if isinstance(sf, dict) else None
    if isinstance(scope, str) and scope.strip():
        y = write_wrapped_text(c, "Scope Overview: " + (scope[:1500]), margin, y, width - 2*margin, leading=12)
    c.showPage(); c.save()

    return str(json_path), str(docx_path), str(pdf_path)

# -------------------------
# Compliance check (faster with prefilter + early exit)
# -------------------------
def _window_keyword_meta(w: str) -> Dict[str, bool]:
    wl = w.lower()
    return {
        "payment": ("payment" in wl) or ("ipc" in wl) or ("interim" in wl),
        "retention": ("retention" in wl),
        "termination": ("terminat" in wl),
        "governing law": ("governing law" in wl) or ("law" in wl),
        "suspension": ("suspens" in wl) or ("non-payment" in wl),
        "safety": ("ppe" in wl) or ("safety" in wl),
        "insurance": ("insurance" in wl) or ("car" in wl) or ("workmen" in wl) or ("tpl" in wl),
        "notices": ("notice" in wl) or ("registered post" in wl) or ("email" in wl),
        "records": ("record" in wl) or ("document" in wl) or ("retain" in wl),
        "dispute": ("arbitrat" in wl) or ("court" in wl) or ("jurisdict" in wl),
    }

def _json_from_yesno(rule: str, window: str, guess_present: bool, quote: str | None, sources: list[str] | None, conf: float) -> dict:
    return {
        "rule": rule,
        "present": bool(guess_present),
        "summary": ("Found evidence in window" if guess_present else None),
        "quote": (quote[:200] if quote else None),
        "sources": sources or ([] if not guess_present else ["page ?"]),
        "confidence": float(conf),
    }

_payment_re = re.compile(r"(payment|payable)\s+(?:within|in)\s+(\d{1,3})\s+days", re.I)
_termination_re = re.compile(r"\bterminat(?:e|ion)\b", re.I)
_govlaw_re = re.compile(r"\bgoverning\s+law\b.*?(india|indian|laws? of [A-Za-z ]+)", re.I | re.S)

def _clean_quote(q: str) -> str:
    if not q: return q
    q = re.sub(r"---\s*PAGE\s+\d+\s*---", " ", q, flags=re.I)
    q = q.replace("PAGE BREAK", " ")
    q = re.sub(r"\s+", " ", q).strip()
    return q[:200]

def _infer_sources_from_quote(quote: str, pages) -> list[str]:
    qn = _clean_quote(quote)
    hits = []
    for p in pages:
        if qn and qn in (p.get("text") or ""):
            hits.append(f"page {p['page']}")
    return hits

def _compliance_regex_fallback(rule: str, window: str) -> dict | None:
    rl = rule.lower()
    w = window

    # A) payment terms
    if "payment" in rl and ("term" in rl or "within" in rl or "days" in rl or "ipc" in rl):
        m = _payment_re.search(w)
        if m:
            quote = m.group(0)
            return _json_from_yesno(rule, w, True, quote, None, 0.75)

    # B) termination
    if "termination" in rl or "terminate" in rl:
        m = _termination_re.search(w)
        if m:
            # try to grab a short line around it
            line = ""
            try:
                s = max(0, m.start()-60); e = min(len(w), m.end()+120)
                line = re.sub(r"\s+", " ", w[s:e]).strip()[:200]
            except:
                pass
            return _json_from_yesno(rule, w, True, line or m.group(0), None, 0.70)

    # C) governing law
    if "governing law" in rl or ("governing" in rl and "law" in rl):
        m = _govlaw_re.search(w)
        if m:
            quote = re.sub(r"\s+", " ", m.group(0))[:200]
            return _json_from_yesno(rule, w, True, quote, None, 0.80)

    return None

def compliance_check_json(full_text_with_pages: str, rules: List[str], llm_compliance: HuggingFacePipeline, tokenizer,
                          window_tokens: int = DEFAULT_WINDOW_TOKENS, overlap_tokens: int = DEFAULT_OVERLAP_TOKENS):
    chain = LLMChain(llm=llm_compliance, prompt=COMPLIANCE_JSON_PROMPT)
    windows = chunk_text_by_tokens(full_text_with_pages, tokenizer, window_tokens, overlap_tokens)

    metas = [{"w": w, "kw": _window_keyword_meta(w)} for w in windows]

    results = []
    for rule in rules:
        rule_l = rule.lower()
        candidates = []
        for m in metas:
            if any(k in rule_l and v for k, v in m["kw"].items()):
                candidates.append(m["w"])
        if not candidates:
            candidates = [m["w"] for m in metas]  # fallback

        aggregated = []
        for w in candidates:
            raw = (chain.predict(window_text=w, rule=rule) or "").strip()
            parsed = None
            try:
                parsed = json.loads(raw)
            except Exception:
                parsed = _loose_json_grab(raw)

            if not parsed:
                # regex fallback builds a valid JSON if we clearly see evidence
                fallback = _compliance_regex_fallback(rule, w)
                if fallback:
                    parsed = fallback

            if parsed:
                aggregated.append(parsed)
                try:
                    if parsed.get("present") and float(parsed.get("confidence", 0.0)) >= 0.80:
                        break
                except Exception:
                    pass

        final = {"rule": rule, "present": False, "summary": None, "quote": None, "sources": [], "confidence": 0.0}
        for a in aggregated:
            if a.get("present"):
                final["present"] = True
                final["summary"] = a.get("summary") or final["summary"]
                final["quote"] = a.get("quote") or final["quote"]
                final["sources"].extend([s for s in a.get("sources", []) if s not in final["sources"]])
                try:
                    final["confidence"] = max(final["confidence"], float(a.get("confidence", 0.0)))
                except Exception:
                    pass

        results.append(final)
    return results

# -------------------------
# Conflict detection (regex)
# -------------------------
def check_commencement_vs_site_possession(pages, full_text):
    date_pattern = r'((?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4})|(?:\d{1,2}\s+[A-Za-z]+\s+\d{4})|(?:[A-Za-z]+\s+\d{1,2},\s*\d{4}))'
    site_pos_pattern = re.compile(r'(site possession|possession of the site|hand over (?:the )?site)[^\.\n]{0,250}(' + date_pattern + r')', re.I)
    commence_pattern = re.compile(r'(commenc(?:e|ement) date|start date|date of commencement|date of commencement of works)[^\.\n]{0,250}(' + date_pattern + r')', re.I)

    site_matches, comm_matches = [], []
    for p in pages:
        text = p.get('text') or ""
        for m in site_pos_pattern.finditer(text):
            raw = m.group(2)
            try:
                dt = dateparser.parse(raw, dayfirst=True)
            except Exception:
                dt = None
            site_matches.append({"page": p['page'], "raw": raw, "date": dt, "snippet": m.group(0)})
        for m in commence_pattern.finditer(text):
            raw = m.group(2)
            try:
                dt = dateparser.parse(raw, dayfirst=True)
            except Exception:
                dt = None
            comm_matches.append({"page": p['page'], "raw": raw, "date": dt, "snippet": m.group(0)})

    conflicts = []
    if site_matches and comm_matches:
        comm_dates = [c['date'] for c in comm_matches if c['date']]
        site_dates = [s['date'] for s in site_matches if s['date']]
        if comm_dates and site_dates:
            min_comm = min(comm_dates)
            max_site = max(site_dates)
            if max_site > min_comm:
                conflicts.append({
                    "type": "commencement_vs_site_possession",
                    "severity": "Critical",
                    "message": f"Site possession ({max_site.date()}) occurs after commencement ({min_comm.date()}).",
                    "evidence": {"commencement": comm_matches, "site_possession": site_matches}
                })
    return conflicts

def check_payment_term_mismatch(pages):
    pay_pattern = re.compile(r'pay(?:ment)?\s+(?:within|in)\s+(\d{1,3})\s+days', re.I)
    matches = []
    for p in pages:
        for m in pay_pattern.finditer(p.get('text') or ""):
            try:
                matches.append({"page": p['page'], "days": int(m.group(1)), "snippet": m.group(0)})
            except Exception:
                pass
    conflicts = []
    if matches:
        days_set = sorted(set(m['days'] for m in matches))
        if len(days_set) > 1:
            conflicts.append({
                "type": "payment_term_mismatch",
                "severity": "High",
                "message": f"Different payment terms found: {days_set} days.",
                "evidence": matches
            })
    return conflicts

def check_retention_mismatch(pages):
    ret_pattern = re.compile(r'retention[^\d]{0,20}(\d{1,2})\s*%?', re.I)
    matches = []
    for p in pages:
        for m in ret_pattern.finditer(p.get('text') or ""):
            try:
                pct = int(m.group(1))
                matches.append({"page": p['page'], "pct": pct, "snippet": m.group(0)})
            except Exception:
                continue
    conflicts = []
    if matches:
        pct_set = sorted(set(m['pct'] for m in matches))
        if len(pct_set) > 1:
            conflicts.append({
                "type": "retention_mismatch",
                "severity": "High",
                "message": f"Different retention percentages found: {pct_set}%.",
                "evidence": matches
            })
    return conflicts

def check_defect_liability_mismatch(pages):
    dl_pattern = re.compile(r'(defect liability period|defects liability|warranty period)[^\d]{0,60}(\d{1,3})\s*(months|month|years|year)?', re.I)
    matches = []
    for p in pages:
        for m in dl_pattern.finditer(p.get('text') or ""):
            try:
                val = int(m.group(2))
                unit = (m.group(3) or "months").lower()
                months = val * 12 if "year" in unit else val
                matches.append({"page": p['page'], "value": months, "raw": m.group(0)})
            except Exception:
                continue
    conflicts = []
    if matches:
        vals = sorted(set(m['value'] for m in matches))
        if len(vals) > 1:
            conflicts.append({
                "type": "defect_liability_mismatch",
                "severity": "High",
                "message": f"Different defect liability lengths found (in months): {vals}.",
                "evidence": matches
            })
    return conflicts

def check_arbitration_vs_court_conflict(pages):
    arb_pattern = re.compile(r'\barbitrat', re.I)
    court_pattern = re.compile(r'\bcourt', re.I)
    arbs, courts = [], []
    for p in pages:
        t = p.get('text') or ""
        if arb_pattern.search(t):
            arbs.append({"page": p['page']})
        if court_pattern.search(t):
            courts.append({"page": p['page']})
    conflicts = []
    if arbs and courts:
        conflicts.append({
            "type": "dispute_resolution_conflict",
            "severity": "Critical",
            "message": "Arbitration-related terms and court-related terms both found; possible conflict in dispute resolution or interim relief.",
            "evidence": {"arbitration_pages": arbs, "court_pages": courts}
        })
    return conflicts

CONFLICT_TO_CATEGORY = {
    "commencement_vs_site_possession": ("Contractual Disputes (Site Possession)", "Consider deeming commencement as actual possession or grant EOT + cost recovery. See 'SITE POSSESSION' clause."),
    "payment_term_mismatch": ("Contractual Disputes (Payment)", "Harmonize payment timeline; prefer main-body or latest-dated doc. Consider change-order for invoices/payments."),
    "retention_mismatch": ("Financial / Payment", "Clarify retention % in payment schedule; propose single retained % and update annex."),
    "defect_liability_mismatch": ("Quality / Warranty", "Adopt the stricter DLP or clarify which schedule governs; update annex."),
    "dispute_resolution_conflict": ("Dispute Resolution", "Resolve arbitration vs court language — prefer arbitration clause with interim relief carve-out."),
}

def map_conflict_to_practical_category(conflict):
    ctype = conflict.get("type")
    cat, hint = CONFLICT_TO_CATEGORY.get(ctype, ("Other / Administrative", "Review conflict and map manually."))
    conflict["category"] = cat
    conflict["resolution_hint"] = hint
    return conflict

def run_conflict_detection(pages, full_text):
    conflicts = []
    conflicts.extend(check_commencement_vs_site_possession(pages, full_text))
    conflicts.extend(check_payment_term_mismatch(pages))
    conflicts.extend(check_retention_mismatch(pages))
    conflicts.extend(check_defect_liability_mismatch(pages))
    conflicts.extend(check_arbitration_vs_court_conflict(pages))
    conflicts = [map_conflict_to_practical_category(c) for c in conflicts]
    return conflicts

# -------------------------
# Pipeline nodes
# -------------------------
def node_ocr_and_extract(pdf_bytes: bytes) -> Dict[str, Any]:
    st.info("Extracting text from PDF (pdfplumber preferred; fallback to OCR).")
    pages, full_text = extract_text_prefer_pdfplumber(pdf_bytes, dpi=200)
    return {"pages": pages, "full_text": full_text}

def node_chunk_text_from_full(raw_text: str) -> List[str]:
    st.info("Chunking text for UI and downstream (char-based chunks).")
    return simple_chunk_text(raw_text, size=3000, overlap=500)

def node_build_model_wrappers(_) -> Dict[str, Any]:
    st.info("Loading model and LangChain wrappers...")
    tokenizer, llm_main, llm_comp = build_generation_pipelines(GEN_MODEL)
    return {"tokenizer": tokenizer, "llm_main": llm_main, "llm_compliance": llm_comp}

# -------------------------
# Streamlit UI
# -------------------------
st.set_page_config(page_title="Contract CAD Chatbot - LangChain Direct LLM", layout="wide")
st.title("Automated Contract CAD Generator")

# Sidebar controls
with st.sidebar:
    st.header("Controls")
    strategy = st.radio("Query strategy", options=["sliding", "concatenate"], index=0)
    window_tokens = st.number_input("Window tokens (approx)", min_value=512, max_value=64000, value=DEFAULT_WINDOW_TOKENS, step=128)
    overlap_tokens = st.number_input("Overlap tokens", min_value=0, max_value=max(1, window_tokens//2), value=DEFAULT_OVERLAP_TOKENS, step=32)
    reserved_qa_tokens = st.number_input("Reserved tokens for QA (generation)", min_value=64, max_value=4096, value=DEFAULT_RESERVED_QA_TOKENS, step=64)
    st.markdown("---")
    uploaded_file = st.file_uploader("Upload contract PDF (scanned or digital)", type=["pdf"])
    regen = st.button("Reprocess / Rebuild context")
    clear_all = st.button("Clear session")

if clear_all:
    for k in list(st.session_state.keys()):
        del st.session_state[k]
    st.experimental_rerun()

# Graph nodes & pipeline
graph = SimpleGraph()
graph.add_node(Node("OCR", lambda b: node_ocr_and_extract(b)))
graph.add_node(Node("Chunk", lambda info: node_chunk_text_from_full(info["full_text"])))
graph.add_node(Node("Model", node_build_model_wrappers))

# Build pipeline & process upload
if uploaded_file and (regen or not st.session_state.get("graph_built", False)):
    try:
        with st.spinner("Running pipeline (extract → chunk → model load)..."):
            pdf_bytes = uploaded_file.read()
            extract_info = node_ocr_and_extract(pdf_bytes)
            pages = extract_info["pages"]
            full_text = extract_info["full_text"]  # includes --- PAGE N --- markers
            chunks = node_chunk_text_from_full(full_text)
            model_info = node_build_model_wrappers(None)
            st.session_state["pages"] = pages
            st.session_state["raw_text"] = full_text
            st.session_state["chunks"] = chunks
            st.session_state["tokenizer"] = model_info["tokenizer"]
            st.session_state["llm_main"] = model_info["llm_main"]
            st.session_state["llm_compliance"] = model_info["llm_compliance"]
            st.session_state["graph_built"] = True
        st.success("Processed PDF and loaded model.")
    except Exception as e:
        st.exception(e)
        st.error("Processing failed.")

# Setup memory-based chain for chat (LangChain Conversation)
if st.session_state.get("graph_built", False) and "chat_memory" not in st.session_state:
    st.session_state["chat_memory"] = ConversationBufferMemory(memory_key="chat_history", return_messages=True)

# Layout: Chat (left) and CAD/Compliance (right)
col1, col2 = st.columns([2, 1])

with col1:
    st.subheader("Chat with the Contract")
    if not st.session_state.get("graph_built", False):
        st.info("Upload & Reprocess a PDF to enable chat.")
    else:
        if "conversation" not in st.session_state:
            st.session_state["conversation"] = []
        for turn in st.session_state["conversation"]:
            if turn["role"] == "user":
                st.chat_message("user").write(turn["text"])
            else:
                st.chat_message("assistant").write(turn["text"])

        user_input = st.chat_input("Ask about the contract (e.g., 'What is the payment clause?')")
        if user_input:
            st.session_state["conversation"].append({"role":"user","text":user_input})
            tokenizer = st.session_state["tokenizer"]
            llm_main = st.session_state["llm_main"]
            full_text_for_windows = st.session_state["raw_text"]  # KEEP PAGE HEADERS
            chain_memory = st.session_state["chat_memory"]

            with st.spinner("Asking LLM..."):
                res = ask_direct_langchain(
                    llm_main, tokenizer, full_text_for_windows, user_input,
                    strategy=strategy, window_tokens=window_tokens, overlap_tokens=overlap_tokens, memory=chain_memory
                )
                answer = res.get("answer", "No answer.")
            st.session_state["conversation"].append({"role":"assistant","text":answer})
            st.chat_message("assistant").write(answer)
            if st.checkbox("Show debug windows & answers", key=f"dbg_chat_{len(st.session_state['conversation'])}"):
                st.write(res)

with col2:
    st.subheader("CAD Generator / Document")
    if not st.session_state.get("graph_built", False):
        st.info("Upload & Reprocess PDF to enable CAD generation.")
    else:
        if st.button("Generate CAD (JSON + DOCX + PDF)"):
            try:
                llm_main = st.session_state["llm_main"]
                tokenizer = st.session_state["tokenizer"]
                full_text = st.session_state["raw_text"]   # keep page markers
                pages = st.session_state["pages"]
                with st.spinner("Generating CAD (JSON/DOCX/PDF)..."):
                    json_path, docx_path, pdf_path = node_generate_cad_json_docx_pdf(
                        llm_main, tokenizer, full_text, pages,
                        output_basename="generated_CAD",
                        window_tokens=window_tokens, overlap_tokens=overlap_tokens
                    )
                st.session_state["cad_json"] = json_path
                st.session_state["cad_docx"] = docx_path
                st.session_state["cad_pdf"] = pdf_path
                st.success("CAD generated as JSON, DOCX and PDF.")
            except Exception as e:
                st.exception(e)
                st.error("CAD generation failed.")

        if st.session_state.get("cad_json"):
            try:
                with open(st.session_state["cad_json"], "rb") as f:
                    st.download_button("Download CAD JSON", f.read(), file_name=Path(st.session_state['cad_json']).name)
            except Exception as e:
                st.error(f"Failed to open CAD JSON: {e}")
        if st.session_state.get("cad_docx"):
            try:
                with open(st.session_state["cad_docx"], "rb") as f:
                    st.download_button("Download CAD DOCX", f.read(), file_name=Path(st.session_state['cad_docx']).name)
            except Exception as e:
                st.error(f"Failed to open CAD DOCX: {e}")
        if st.session_state.get("cad_pdf"):
            try:
                with open(st.session_state["cad_pdf"], "rb") as f:
                    st.download_button("Download CAD PDF", f.read(), file_name=Path(st.session_state['cad_pdf']).name)
            except Exception as e:
                st.error(f"Failed to open CAD PDF: {e}")

    st.markdown("---")
    st.subheader("📑 Compliance Check")
    if not st.session_state.get("graph_built", False):
        st.info("Upload a contract and press Reprocess first.")
    else:
        default_rules = "Does the contract specify payment terms?\nIs there a termination clause?\nWhat is the governing law?"
        rules_input = st.text_area("Enter compliance rules (one per line)", value=default_rules)
        if st.button("Run Compliance Check"):
            rules = [r.strip() for r in rules_input.splitlines() if r.strip()]
            llm_compliance = st.session_state["llm_compliance"]
            tokenizer = st.session_state["tokenizer"]
            full_text = st.session_state["raw_text"]
            with st.spinner("Running compliance checks..."):
                results = compliance_check_json(
                    full_text, rules, llm_compliance, tokenizer,
                    window_tokens=window_tokens, overlap_tokens=overlap_tokens
                )
            st.subheader("Compliance Report")
            df_rows = []
            for r in results:
                st.markdown(f"**Rule:** {r['rule']}  \n**Present:** {r['present']}  \n**Summary:** {r['summary']}  \n**Quote:** {r['quote']}  \n**Sources:** {r['sources']}  \n**Confidence:** {r['confidence']}")
                df_rows.append([r['rule'], r['present'], r['summary'], r['quote'], ", ".join(r['sources']), r['confidence']])
            if df_rows:
                st.dataframe(pd.DataFrame(df_rows, columns=["Rule", "Present", "Summary", "Quote", "Sources", "Confidence"]))

    st.markdown("---")
    st.subheader("⚠️ Conflict Detection")
    if not st.session_state.get("graph_built", False):
        st.info("Upload a contract and press Reprocess first.")
    else:
        if st.button("Run Conflict Detection"):
            try:
                pages = st.session_state.get("pages")
                full_text = st.session_state.get("raw_text")
                with st.spinner("Detecting conflicts..."):
                    conflicts = run_conflict_detection(pages, full_text)
                st.session_state["conflicts"] = conflicts
                if not conflicts:
                    st.success("No conflicts detected by automated checks.")
                else:
                    st.warning(f"{len(conflicts)} potential conflicts detected.")
                    for idx, c in enumerate(conflicts, start=1):
                        st.markdown(f"### {idx}. Category: **{c.get('category')}**  \n**Type:** {c['type']}  \n**Severity:** {c['severity']}  \n**Message:** {c['message']}")
                        st.markdown(f"**Suggested quick fix:** {c.get('resolution_hint')}")
                        st.json(c.get("evidence"))
            except Exception as e:
                st.exception(e)
                st.error("Conflict detection failed.")

    st.markdown("---")
    st.subheader("Raw text & pages preview")
    if st.session_state.get("raw_text"):
        if st.checkbox("Show OCR/raw extracted text (first 15000 chars)"):
            st.text_area("Extracted Text", value=st.session_state["raw_text"][:15000], height=400)
    else:
        st.info("No extracted text available. Upload a PDF and press Reprocess.")

st.caption("LangChain + HuggingFace. Outputs: CAD JSON, DOCX, PDF. Faster compliance. Page-aware windows for better sources.")
# End of file
