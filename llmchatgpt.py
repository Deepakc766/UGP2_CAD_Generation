# # app_streamlit_cad_chatgpt.py
# import os
# import io
# import json
# import textwrap
# import tempfile
# from pathlib import Path
# from typing import List, Dict, Any, Optional
# import logging
# import re
# import streamlit as st
# import pdfplumber
# from reportlab.lib.pagesizes import A4
# from reportlab.pdfgen import canvas
# from dateutil import parser as dateparser

# from pdf2image import convert_from_bytes
# from PIL import Image
# import pytesseract
# import shutil
# from docx import Document as DocxDocument
# from docx.shared import Pt
# import pandas as pd

# # ============ NEW: OpenAI ChatGPT client ============
# # Requires: pip install openai==1.*
# from openai import OpenAI

# # -------------------------
# # Config
# # -------------------------
# DEFAULT_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")  # fast & capable; you can switch to gpt-4.1
# MAX_OUTPUT_TOKENS = 4096  # for JSON responses
# TEMPERATURE = 0.0

# # Ensure tesseract path if present
# tesseract_path = shutil.which("tesseract")
# if tesseract_path:
#     pytesseract.pytesseract.tesseract_cmd = tesseract_path

# # -------------------------
# # OpenAI helpers
# # -------------------------
# @st.cache_resource(show_spinner=False)
# def build_openai_client():
#     api_key = st.secrets.get("OPENAI_API_KEY", os.environ.get("OPENAI_API_KEY"))
#     if not api_key:
#         st.error("Missing OpenAI API key. Set in Streamlit secrets as OPENAI_API_KEY or env var.")
#         st.stop()
#     client = OpenAI(api_key=api_key)
#     return client

# def chat_json(
#     client: OpenAI,
#     model: str,
#     system_prompt: str,
#     user_prompt: str,
#     max_tokens: int = MAX_OUTPUT_TOKENS,
#     temperature: float = TEMPERATURE,
# ) -> str:
#     """
#     Calls Chat Completions with response_format=json_object to force strict JSON.
#     Returns raw JSON string from the assistant.
#     """
#     resp = client.chat.completions.create(
#         model=model,
#         temperature=temperature,
#         max_tokens=max_tokens,
#         response_format={"type": "json_object"},
#         messages=[
#             {"role": "system", "content": system_prompt.strip()},
#             {"role": "user", "content": user_prompt.strip()},
#         ],
#     )
#     return (resp.choices[0].message.content or "").strip()

# def safe_json_loads(txt: str) -> Optional[dict]:
#     if not txt:
#         return None
#     try:
#         return json.loads(txt)
#     except Exception:
#         # last-resort: try to extract {...}
#         s = txt.find("{"); e = txt.rfind("}")
#         if s != -1 and e != -1:
#             cand = txt[s:e+1]
#             cand = re.sub(r",\s*}", "}", cand)
#             cand = re.sub(r",\s*]", "]", cand)
#             try:
#                 return json.loads(cand)
#             except Exception:
#                 return None
#         return None

# # -------------------------
# # PDF → Text extraction
# # -------------------------
# def pdf_bytes_to_images(pdf_bytes: bytes, dpi: int = 200) -> List[Image.Image]:
#     return convert_from_bytes(pdf_bytes, dpi=dpi)

# def ocr_image_to_text(image: Image.Image) -> str:
#     return pytesseract.image_to_string(image)

# def extract_text_from_pdf_bytes(pdf_bytes: bytes, dpi: int = 200) -> str:
#     pages = pdf_bytes_to_images(pdf_bytes, dpi=dpi)
#     texts = []
#     for i, img in enumerate(pages):
#         page_text = ocr_image_to_text(img)
#         header = f"\n\n--- PAGE {i+1} ---\n\n"
#         texts.append(header + page_text)
#     return "\n".join(texts)

# def extract_text_prefer_pdfplumber(pdf_bytes: bytes, dpi: int = 200):
#     """
#     Use pdfplumber to extract selectable text per page; fallback to OCR for pages with no text.
#     Returns: (pages_list, full_text)
#       pages_list: [{'page': n, 'text': str, 'is_ocr': bool}, ...]
#       full_text: string with --- PAGE n --- markers
#     """
#     pages = []
#     try:
#         with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
#             for i, page in enumerate(pdf.pages, start=1):
#                 try:
#                     txt = page.extract_text() or ""
#                 except Exception:
#                     txt = ""
#                 if txt.strip():
#                     pages.append({"page": i, "text": txt, "is_ocr": False})
#                 else:
#                     try:
#                         pil_img = page.to_image(resolution=dpi).original
#                         ocr_txt = pytesseract.image_to_string(pil_img)
#                     except Exception:
#                         ocr_txt = ""
#                     pages.append({"page": i, "text": ocr_txt, "is_ocr": True})
#     except Exception as e:
#         logging.warning("pdfplumber failed: %s. Falling back to OCR for whole PDF.", e)
#         ocr_full = extract_text_from_pdf_bytes(pdf_bytes, dpi=dpi)
#         page_texts = [p for p in ocr_full.split("\f")] if "\f" in ocr_full else [ocr_full]
#         for idx, p_text in enumerate(page_texts, start=1):
#             pages.append({"page": idx, "text": p_text, "is_ocr": True})

#     full = []
#     for p in pages:
#         full.append(f"\n\n--- PAGE {p['page']} ---\n{p['text']}")
#     return pages, "\n".join(full)

# # -------------------------
# # Prompts (ONE-SHOT, full contract)
# # -------------------------
# # -------------------------
# # Prompts (REPLACE only CAD_SYSTEM and CAD_USER_TMPL)
# # -------------------------

# CAD_SYSTEM = """
# You are a senior contracts analyst.
# Output STRICT JSON ONLY (no markdown, no comments, no prose outside the JSON).
# Never hallucinate; if not explicitly supported by contract text, use null.
# Prefer concise phrasing. Normalize dates to YYYY-MM-DD when explicit.
# For every factual extraction, attach a sources array with page/ clause markers (e.g., "page 7", "clause 37.4").
# """

# # Style guide reflecting the sample CAD sections:
# CAD_STYLE_GUIDE = """
# Sections and order (must match exactly):
# 1. Salient Features of the Contract
# 2. List of Important Submittals
# 3. Notice / Information Clauses
# 4. Important clauses pertaining to project progress - EOT, Escalation, Variation, Suspension, etc.
# 5. Payment Clause
# 6. Identify Risk, Risk Matrix, and Risk Allocation
# 7. Claim, Dispute, and Arbitration clause
# """

# CAD_USER_TMPL = """Use ONLY the CONTRACT TEXT to produce ONE JSON object exactly matching the schema.
# Respond with strict JSON (start with '{{' and end with '}}'). No extra text.

# CONTRACT TEXT (page markers included):
# {contract_text}

# STYLE GUIDE:
# {style_guide}

# SCHEMA:
# {{
#   "salient_features": [
#     {{"description":"string","clause_or_ref":"string ","details":"string ","sources":["page X"]}}
#   ],
#   "important_submittals": [
#     {{"stage":"before|during|after ","document":"string","due":"string ","notes":"string ","clause_or_ref":"string ","sources":["page X"]}}
#   ],
#   "notice_information_clauses": [
#     {{"description":"string","clause_or_ref":"string ","details":"string","sources":["page X"]}}
#   ],
#   "project_progress_clauses": [
#     {{"topic":"string","clause_or_ref":"string ","summary":"string","sources":["page X"]}}
#   ],
#   "payment": {{
#     "type_of_contract":"string",
#     "mode_of_payment":"string",
#     "measurement_and_evaluation":"string ",
#     "billing_timing":"string ",
#     "other_terms":"string ",
#     "contract_sum": {{"amount": number , "currency":"string ","sources":["page X"]}},
#     "mobilization_advance": {{"applicable": true|false, "amount_pct": number , "interest":"string ", "security":"string ", "release_conditions":"string ","sources":["page X"]}},
#     "retention": {{"percent": number ,"release_condition":"string ","sources":["page X"]}},
#     "interim_payments": [{{"frequency":"string ","certifier":"string ","payment_days_after_certification": number ,"sources":["page X"]}}],
#     "price_escalation": {{"applicable": true|false, "basis":"string ", "formula_or_reference":"string ","sources":["page X"]}},
#     "final_payment": {{"required_docs":["string"],"release_days_after_submission": number ,"sources":["page X"]}},
#     "sources":["page X"]
#   }},
#   "risk_matrix": {{
#     "severity_buckets":[{{"label":"string","criteria":"string"}}],
#     "probability_buckets":[{{"label":"string","criteria":"string"}}],
#     "matrix_legend": "string "
#   }},
#   "risks_and_allocation": [
#     {{"id":"string","category":"string","risk_element":"string","probability":"Most Likely|Likely|Occasional|Unlikely|Remote ","severity":"Very Severe|Major|Significant|Minor|Insignificant ","rating":"Low|Medium|High|Critical ","owner":"Employer|Contractor|Shared ","mitigation":"string ","sources":["page X"]}}
#   ],
#   "claims_disputes_arbitration": {{
#     "arbitration": {{"applicable": true|false,"forum_or_rules":"string ","notes":"string ","sources":["page X"]}},
#     "court_jurisdiction":"string ",
#     "claims_summary":[{{"topic":"string","payable_or_not":"Payable|Not Payable|Unclear","notes":"string","sources":["page X"]}}],
#     "dispute_areas":[{{"description":"string","sources":["page X"]}}],
#     "excusable_delays":[{{"item":"string","sources":["page X"]}}],
#     "delay_compensation_clause":"string ",
#     "sources":["page X"]
#   }}
# }}

# CRITICAL RULES:
# - Fill fields ONLY if supported by the contract text. If unknown, use null.
# - Every extracted fact MUST include "sources" with page and, if present, clause references.
# - Keep text concise and specific (no generic boilerplate). Use lists instead of long paragraphs where helpful.
# - Do NOT invent clause numbers; if absent, keep clause_or_ref null and rely on page numbers in "sources".
# """



# COMPLIANCE_SYSTEM = """You are a contracts compliance checker. Always return STRICT JSON (no markdown)."""

# COMPLIANCE_USER_TMPL = """Use ONLY the CONTRACT TEXT to evaluate the rule. Return ONE strict JSON object.

# CONTRACT TEXT:
# {contract_text}

# RULE:
# {rule}

# Return JSON in one of these shapes:
# {{"rule":"{rule}","present": true,"summary":"short","quote":"<=200 chars verbatim","sources":["page X","clause Y"],"confidence":0.9}}
# or
# {{"rule":"{rule}","present": false,"summary": null,"quote": null,"sources": [],"confidence": 0.99}}
# """

# QA_SYSTEM = """You are a contracts Q&A assistant. Answer ONLY from the contract text."""

# QA_USER_TMPL = """CONTRACT TEXT:
# {contract_text}

# QUESTION:
# {question}

# If unknown, reply exactly: "Not found in contract."
# """

# # -------------------------
# # One-shot CAD generation (ChatGPT)
# # -------------------------
# def node_generate_cad_json_docx_pdf_openai(
#     client: OpenAI,
#     model: str,
#     full_text_with_pages: str,
#     pages,
#     output_basename: str = "generated_CAD"
# ):
#     # 1) Call LLM once for the whole contract
#     user_prompt = CAD_USER_TMPL.format(contract_text=full_text_with_pages)
#     raw = chat_json(client, model, CAD_SYSTEM, user_prompt)
#     data = safe_json_loads(raw)

#     if not data or not isinstance(data, dict):
#         # one repair attempt by asking model to emit the same content as strict JSON
#         repair_prompt = f"Convert this content into a SINGLE strict JSON object (no commentary):\n\n{raw}"
#         repaired = chat_json(client, model, CAD_SYSTEM, repair_prompt)
#         data = safe_json_loads(repaired)
#         if not data:
#             raise RuntimeError("Failed to obtain valid CAD JSON from LLM.")

#     # 2) Save JSON
#     json_path = Path(tempfile.gettempdir()) / f"{output_basename}.json"
#     with open(json_path, "w", encoding="utf8") as f:
#         json.dump(data, f, ensure_ascii=False, indent=2)

#     # 3) DOCX rendering
#     docx = DocxDocument()
#     style = docx.styles["Normal"]; style.font.name="Arial"; style.font.size = Pt(11)
#     docx.add_heading("Contract Appreciation Document (Generated)", level=1)
#     docx.add_paragraph("Generated via ChatGPT one-shot CAD. Verify against the contract before finalization.")

#     sf = data.get("salient_features", {})
#     docx.add_heading("1. Salient Features", level=2)
#     if sf:
#         for k, v in sf.items():
#             if isinstance(v, dict):
#                 # heuristics to print a readable value
#                 printable = v.get("name") or v.get("amount") or v.get("summary") or json.dumps(v, ensure_ascii=False)
#                 docx.add_paragraph(f"{k}: {printable}")
#             else:
#                 docx.add_paragraph(f"{k}: {v}")
#     else:
#         docx.add_paragraph("Not found")

#     submittals = data.get("important_submittals", [])
#     docx.add_heading("2. Important Submittals", level=2)
#     if submittals:
#         cols = ["stage","document","due","notes","sources"]
#         table = docx.add_table(rows=1, cols=len(cols))
#         for i, c in enumerate(cols):
#             table.rows[0].cells[i].text = c
#         for r in submittals:
#             rc = table.add_row().cells
#             rc[0].text = (r.get("stage") or "")
#             rc[1].text = (r.get("document") or "")
#             rc[2].text = (r.get("due") or "")
#             rc[3].text = (r.get("notes") or "")
#             rc[4].text = ", ".join(r.get("sources") or [])
#     else:
#         docx.add_paragraph("No submittals listed.")

#     notices = data.get("notice_clauses", [])
#     docx.add_heading("3. Notice Clauses", level=2)
#     if notices:
#         for n in notices:
#             docx.add_paragraph(
#                 f"{n.get('event')} | {n.get('notifier')} → {n.get('recipient')} | "
#                 f"timeline: {n.get('timeline')} | method: {n.get('method')} | "
#                 f"sources: {', '.join(n.get('sources', []))}"
#             )
#     else:
#         docx.add_paragraph("No notice clauses listed.")

#     docx.add_heading("4. Payment Summary", level=2)
#     pay = data.get("payment", {})
#     if pay:
#         for k, v in pay.items():
#             docx.add_paragraph(f"{k}: {json.dumps(v, ensure_ascii=False) if isinstance(v,(dict,list)) else str(v)}")
#     else:
#         docx.add_paragraph("No payment information found.")

#     docx.add_heading("5. Risks and Allocation", level=2)
#     risks = data.get("risks_and_allocation", [])
#     if risks:
#         for r in risks:
#             docx.add_paragraph(json.dumps(r, ensure_ascii=False))
#     else:
#         docx.add_paragraph("No risks listed.")

#     docx.add_heading("6. Claims, Disputes & Arbitration", level=2)
#     cda = data.get("claims_disputes_arbitration", {})
#     docx.add_paragraph(json.dumps(cda, ensure_ascii=False, indent=2))

#     docx_path = Path(tempfile.gettempdir()) / f"{output_basename}.docx"
#     docx.save(docx_path)

#     # 4) Summary PDF (same style as before)
#     pdf_path = Path(tempfile.gettempdir()) / f"{output_basename}.pdf"
#     c = canvas.Canvas(str(pdf_path), pagesize=A4)
#     width, height = A4
#     margin = 50
#     y = height - margin

#     def write_wrapped_text(cobj, text, x, y, max_width, leading=12):
#         from reportlab.lib.utils import simpleSplit
#         lines = simpleSplit(text, "Helvetica", 10, max_width)
#         for ln in lines:
#             if y < margin + 40:
#                 cobj.showPage()
#                 y = height - margin
#             cobj.drawString(x, y, ln)
#             y -= leading
#         return y

#     c.setFont("Helvetica-Bold", 14)
#     c.drawString(margin, y, "Contract Appreciation Document (Generated)"); y -= 22
#     c.setFont("Helvetica", 10)
#     proj = sf.get("project_name") if isinstance(sf, dict) else None
#     c.drawString(margin, y, f"Project: {proj or 'N/A'}"); y -= 14
#     employer_name = sf.get("employer", {}).get("name") if isinstance(sf.get("employer"), dict) else None
#     contractor_name = sf.get("contractor", {}).get("name") if isinstance(sf.get("contractor"), dict) else None
#     c.drawString(margin, y, f"Employer: {employer_name or 'N/A'}"); y -= 14
#     c.drawString(margin, y, f"Contractor: {contractor_name or 'N/A'}"); y -= 18
#     scope = sf.get("scope_overview") if isinstance(sf, dict) else None
#     if isinstance(scope, str) and scope.strip():
#         y = write_wrapped_text(c, "Scope Overview: " + (scope[:1500]), margin, y, width - 2*margin, leading=12)
#     c.showPage(); c.save()

#     return str(data), str(json_path), str(docx_path), str(pdf_path)

# # -------------------------
# # Compliance (one-shot per rule)
# # -------------------------
# def compliance_check_json_openai(
#     client: OpenAI,
#     model: str,
#     full_text_with_pages: str,
#     rules: List[str]
# ) -> List[Dict[str, Any]]:
#     results = []
#     for rule in rules:
#         uprompt = COMPLIANCE_USER_TMPL.format(contract_text=full_text_with_pages, rule=rule)
#         raw = chat_json(client, model, COMPLIANCE_SYSTEM, uprompt, max_tokens=800)
#         obj = safe_json_loads(raw)
#         if not obj:
#             # fallback: mark not present with low confidence
#             obj = {"rule": rule, "present": False, "summary": None, "quote": None, "sources": [], "confidence": 0.50}
#         results.append(obj)
#     return results

# # -------------------------
# # Conflict detection (regex; unchanged)
# # -------------------------
# def check_commencement_vs_site_possession(pages, full_text):
#     date_pattern = r'((?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4})|(?:\d{1,2}\s+[A-Za-z]+\s+\d{4})|(?:[A-Za-z]+\s+\d{1,2},\s*\d{4}))'
#     site_pos_pattern = re.compile(r'(site possession|possession of the site|hand over (?:the )?site)[^\.\n]{0,250}(' + date_pattern + r')', re.I)
#     commence_pattern = re.compile(r'(commenc(?:e|ement) date|start date|date of commencement|date of commencement of works)[^\.\n]{0,250}(' + date_pattern + r')', re.I)

#     site_matches, comm_matches = [], []
#     for p in pages:
#         text = p.get('text') or ""
#         for m in site_pos_pattern.finditer(text):
#             raw = m.group(2)
#             try:
#                 dt = dateparser.parse(raw, dayfirst=True)
#             except Exception:
#                 dt = None
#             site_matches.append({"page": p['page'], "raw": raw, "date": dt, "snippet": m.group(0)})
#         for m in commence_pattern.finditer(text):
#             raw = m.group(2)
#             try:
#                 dt = dateparser.parse(raw, dayfirst=True)
#             except Exception:
#                 dt = None
#             comm_matches.append({"page": p['page'], "raw": raw, "date": dt, "snippet": m.group(0)})

#     conflicts = []
#     if site_matches and comm_matches:
#         comm_dates = [c['date'] for c in comm_matches if c['date']]
#         site_dates = [s['date'] for s in site_matches if s['date']]
#         if comm_dates and site_dates:
#             min_comm = min(comm_dates)
#             max_site = max(site_dates)
#             if max_site > min_comm:
#                 conflicts.append({
#                     "type": "commencement_vs_site_possession",
#                     "severity": "Critical",
#                     "message": f"Site possession ({max_site.date()}) occurs after commencement ({min_comm.date()}).",
#                     "evidence": {"commencement": comm_matches, "site_possession": site_matches}
#                 })
#     return conflicts

# def check_payment_term_mismatch(pages):
#     pay_pattern = re.compile(r'pay(?:ment)?\s+(?:within|in)\s+(\d{1,3})\s+days', re.I)
#     matches = []
#     for p in pages:
#         for m in pay_pattern.finditer(p.get('text') or ""):
#             try:
#                 matches.append({"page": p['page'], "days": int(m.group(1)), "snippet": m.group(0)})
#             except Exception:
#                 pass
#     conflicts = []
#     if matches:
#         days_set = sorted(set(m['days'] for m in matches))
#         if len(days_set) > 1:
#             conflicts.append({
#                 "type": "payment_term_mismatch",
#                 "severity": "High",
#                 "message": f"Different payment terms found: {days_set} days.",
#                 "evidence": matches
#             })
#     return conflicts

# def check_retention_mismatch(pages):
#     ret_pattern = re.compile(r'retention[^\d]{0,20}(\d{1,2})\s*%?', re.I)
#     matches = []
#     for p in pages:
#         for m in ret_pattern.finditer(p.get('text') or ""):
#             try:
#                 pct = int(m.group(1))
#                 matches.append({"page": p['page'], "pct": pct, "snippet": m.group(0)})
#             except Exception:
#                 continue
#     conflicts = []
#     if matches:
#         pct_set = sorted(set(m['pct'] for m in matches))
#         if len(pct_set) > 1:
#             conflicts.append({
#                 "type": "retention_mismatch",
#                 "severity": "High",
#                 "message": f"Different retention percentages found: {pct_set}%.",
#                 "evidence": matches
#             })
#     return conflicts

# def check_defect_liability_mismatch(pages):
#     dl_pattern = re.compile(r'(defect liability period|defects liability|warranty period)[^\d]{0,60}(\d{1,3})\s*(months|month|years|year)?', re.I)
#     matches = []
#     for p in pages:
#         for m in dl_pattern.finditer(p.get('text') or ""):
#             try:
#                 val = int(m.group(2))
#                 unit = (m.group(3) or "months").lower()
#                 months = val * 12 if "year" in unit else val
#                 matches.append({"page": p['page'], "value": months, "raw": m.group(0)})
#             except Exception:
#                 continue
#     conflicts = []
#     if matches:
#         vals = sorted(set(m['value'] for m in matches))
#         if len(vals) > 1:
#             conflicts.append({
#                 "type": "defect_liability_mismatch",
#                 "severity": "High",
#                 "message": f"Different defect liability lengths found (in months): {vals}.",
#                 "evidence": matches
#             })
#     return conflicts

# def check_arbitration_vs_court_conflict(pages):
#     arb_pattern = re.compile(r'\barbitrat', re.I)
#     court_pattern = re.compile(r'\bcourt', re.I)
#     arbs, courts = [], []
#     for p in pages:
#         t = p.get('text') or ""
#         if arb_pattern.search(t):
#             arbs.append({"page": p['page']})
#         if court_pattern.search(t):
#             courts.append({"page": p['page']})
#     conflicts = []
#     if arbs and courts:
#         conflicts.append({
#             "type": "dispute_resolution_conflict",
#             "severity": "Critical",
#             "message": "Arbitration and court references coexist; possible conflict in dispute resolution.",
#             "evidence": {"arbitration_pages": arbs, "court_pages": courts}
#         })
#     return conflicts

# CONFLICT_TO_CATEGORY = {
#     "commencement_vs_site_possession": ("Contractual Disputes (Site Possession)", "Deem commencement as possession or grant EOT + cost; see Site Possession."),
#     "payment_term_mismatch": ("Contractual Disputes (Payment)", "Harmonize payment timeline; prefer main body or latest-dated doc."),
#     "retention_mismatch": ("Financial / Payment", "Clarify a single retention % and update annex."),
#     "defect_liability_mismatch": ("Quality / Warranty", "Adopt stricter DLP or clarify governing schedule."),
#     "dispute_resolution_conflict": ("Dispute Resolution", "Prefer arbitration with interim relief carve-out; remove ambiguity."),
# }

# def map_conflict_to_practical_category(conflict):
#     ctype = conflict.get("type")
#     cat, hint = CONFLICT_TO_CATEGORY.get(ctype, ("Other / Administrative", "Review conflict and map manually."))
#     conflict["category"] = cat
#     conflict["resolution_hint"] = hint
#     return conflict

# def run_conflict_detection(pages, full_text):
#     conflicts = []
#     conflicts.extend(check_commencement_vs_site_possession(pages, full_text))
#     conflicts.extend(check_payment_term_mismatch(pages))
#     conflicts.extend(check_retention_mismatch(pages))
#     conflicts.extend(check_defect_liability_mismatch(pages))
#     conflicts.extend(check_arbitration_vs_court_conflict(pages))
#     conflicts = [map_conflict_to_practical_category(c) for c in conflicts]
#     return conflicts

# # -------------------------
# # Streamlit UI
# # -------------------------
# st.set_page_config(page_title="Contract CAD (ChatGPT One-Shot)", layout="wide")
# st.title("Automated Contract CAD Generator — ChatGPT (One-Shot, No Chunking)")

# with st.sidebar:
#     st.header("Model & Controls")
#     model_name = st.text_input("OpenAI model", value=DEFAULT_MODEL, help="e.g., gpt-4.1, gpt-4o, gpt-4o-mini")
#     uploaded_file = st.file_uploader("Upload contract PDF (scanned or digital)", type=["pdf"])
#     regen = st.button("Extract / Rebuild context")
#     clear_all = st.button("Clear session")

# if clear_all:
#     for k in list(st.session_state.keys()):
#         del st.session_state[k]
#     st.experimental_rerun()

# # Extract & store
# if uploaded_file and (regen or not st.session_state.get("graph_built", False)):
#     try:
#         with st.spinner("Extracting text (pdfplumber preferred; OCR fallback)..."):
#             client = build_openai_client()
#             pdf_bytes = uploaded_file.read()
#             pages, full_text = extract_text_prefer_pdfplumber(pdf_bytes, dpi=200)
#             st.session_state["pages"] = pages
#             st.session_state["raw_text"] = full_text  # includes --- PAGE N --- markers
#             st.session_state["graph_built"] = True
#         st.success("PDF processed.")
#     except Exception as e:
#         st.exception(e)
#         st.error("Processing failed.")

# # Layout
# col1, col2 = st.columns([2, 1])

# # Chat Q&A (one-shot on full text)
# with col1:
#     st.subheader("Chat with the Contract (One-Shot)")
#     if not st.session_state.get("graph_built", False):
#         st.info("Upload & Extract a PDF to enable chat.")
#     else:
#         if "conversation" not in st.session_state:
#             st.session_state["conversation"] = []
#         for turn in st.session_state["conversation"]:
#             st.chat_message(turn["role"]).write(turn["text"])

#         user_q = st.chat_input("Ask about the contract (e.g., 'What is the payment clause?')")
#         if user_q:
#             st.session_state["conversation"].append({"role": "user", "text": user_q})
#             client = build_openai_client()
#             uprompt = QA_USER_TMPL.format(contract_text=st.session_state["raw_text"], question=user_q)
#             try:
#                 resp = client.chat.completions.create(
#                     model=model_name,
#                     temperature=0,
#                     max_tokens=900,
#                     messages=[
#                         {"role": "system", "content": QA_SYSTEM},
#                         {"role": "user", "content": uprompt},
#                     ],
#                 )
#                 ans = (resp.choices[0].message.content or "").strip()
#             except Exception as e:
#                 ans = f"Error: {e}"
#             st.session_state["conversation"].append({"role": "assistant", "text": ans})
#             st.chat_message("assistant").write(ans)

# # CAD / Compliance / Conflicts
# with col2:
#     st.subheader("CAD Generator")
#     if not st.session_state.get("graph_built", False):
#         st.info("Upload & Extract a PDF to enable CAD generation.")
#     else:
#         if st.button("Generate CAD (JSON + DOCX + PDF)"):
#             try:
#                 client = build_openai_client()
#                 full_text = st.session_state["raw_text"]
#                 pages = st.session_state["pages"]
#                 with st.spinner("Generating one-shot CAD via ChatGPT..."):
#                     _, json_path, docx_path, pdf_path = node_generate_cad_json_docx_pdf_openai(
#                         client, model_name, full_text, pages, output_basename="generated_CAD"
#                     )
#                 st.session_state["cad_json"] = json_path
#                 st.session_state["cad_docx"] = docx_path
#                 st.session_state["cad_pdf"] = pdf_path
#                 st.success("CAD generated.")
#             except Exception as e:
#                 st.exception(e)
#                 st.error("CAD generation failed.")

#         if st.session_state.get("cad_json"):
#             try:
#                 with open(st.session_state["cad_json"], "rb") as f:
#                     st.download_button("Download CAD JSON", f.read(), file_name=Path(st.session_state['cad_json']).name)
#             except Exception as e:
#                 st.error(f"Failed to open CAD JSON: {e}")
#         if st.session_state.get("cad_docx"):
#             try:
#                 with open(st.session_state["cad_docx"], "rb") as f:
#                     st.download_button("Download CAD DOCX", f.read(), file_name=Path(st.session_state['cad_docx']).name)
#             except Exception as e:
#                 st.error(f"Failed to open CAD DOCX: {e}")
#         if st.session_state.get("cad_pdf"):
#             try:
#                 with open(st.session_state["cad_pdf"], "rb") as f:
#                     st.download_button("Download CAD PDF", f.read(), file_name=Path(st.session_state['cad_pdf']).name)
#             except Exception as e:
#                 st.error(f"Failed to open CAD PDF: {e}")

#     st.markdown("---")
#     st.subheader("📑 Compliance Check (One-Shot per rule)")
#     if not st.session_state.get("graph_built", False):
#         st.info("Upload a contract first.")
#     else:
#         default_rules = "Does the contract specify payment terms?\nIs there a termination clause?\nWhat is the governing law?"
#         rules_input = st.text_area("Enter compliance rules (one per line)", value=default_rules)
#         if st.button("Run Compliance Check"):
#             try:
#                 rules = [r.strip() for r in rules_input.splitlines() if r.strip()]
#                 client = build_openai_client()
#                 full_text = st.session_state["raw_text"]
#                 with st.spinner("Checking compliance..."):
#                     results = compliance_check_json_openai(client, model_name, full_text, rules)
#                 st.subheader("Compliance Report")
#                 df_rows = []
#                 for r in results:
#                     st.markdown(
#                         f"**Rule:** {r.get('rule')}  \n"
#                         f"**Present:** {r.get('present')}  \n"
#                         f"**Summary:** {r.get('summary')}  \n"
#                         f"**Quote:** {r.get('quote')}  \n"
#                         f"**Sources:** {r.get('sources')}  \n"
#                         f"**Confidence:** {r.get('confidence')}"
#                     )
#                     df_rows.append([
#                         r.get("rule"), r.get("present"), r.get("summary"),
#                         r.get("quote"), ", ".join(r.get("sources", [])), r.get("confidence")
#                     ])
#                 if df_rows:
#                     st.dataframe(pd.DataFrame(df_rows, columns=["Rule","Present","Summary","Quote","Sources","Confidence"]))
#             except Exception as e:
#                 st.exception(e)
#                 st.error("Compliance check failed.")

#     st.markdown("---")
#     st.subheader("⚠️ Conflict Detection")
#     if not st.session_state.get("graph_built", False):
#         st.info("Upload a contract first.")
#     else:
#         if st.button("Run Conflict Detection"):
#             try:
#                 pages = st.session_state.get("pages")
#                 full_text = st.session_state.get("raw_text")
#                 with st.spinner("Detecting conflicts..."):
#                     conflicts = run_conflict_detection(pages, full_text)
#                 st.session_state["conflicts"] = conflicts
#                 if not conflicts:
#                     st.success("No conflicts detected by automated checks.")
#                 else:
#                     st.warning(f"{len(conflicts)} potential conflicts detected.")
#                     for idx, c in enumerate(conflicts, start=1):
#                         st.markdown(
#                             f"### {idx}. Category: **{c.get('category')}**  \n"
#                             f"**Type:** {c['type']}  \n**Severity:** {c['severity']}  \n"
#                             f"**Message:** {c['message']}"
#                         )
#                         st.markdown(f"**Suggested quick fix:** {c.get('resolution_hint')}")
#                         st.json(c.get("evidence"))
#             except Exception as e:
#                 st.exception(e)
#                 st.error("Conflict detection failed.")

#     st.markdown("---")
#     st.subheader("Raw text & pages preview")
#     if st.session_state.get("raw_text"):
#         if st.checkbox("Show extracted text (first 15000 chars)"):
#             st.text_area("Extracted Text", value=st.session_state["raw_text"][:15000], height=400)
#     else:
#         st.info("No extracted text yet. Upload a PDF and click Extract.")

# st.caption("ChatGPT One-Shot CAD • No chunking, no windowing • JSON-enforced responses")
# # End of file



# app_streamlit_cad_chatgpt.py
import os
import io
import json
import textwrap
import tempfile
from pathlib import Path
from typing import List, Dict, Any, Optional
import logging
import re
import streamlit as st
import pdfplumber
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from dateutil import parser as dateparser

from pdf2image import convert_from_bytes
from PIL import Image
import pytesseract
import shutil
from docx import Document as DocxDocument
from docx.shared import Pt
import pandas as pd
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# OLD (optional): from reportlab.pdfgen import canvas
# Either delete it or alias so it can't be used by accident:
from reportlab.pdfgen import canvas as pdfcanvas  # <- avoid name "canvas" = c.setFont trap
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# ============ OpenAI ChatGPT client ============
# Requires: pip install openai==1.*
from openai import OpenAI

# -------------------------
# Config (UPDATED)
# -------------------------
# DEFAULT_MODEL = os.environ.get("OPENAI_MODEL", "gpt-5.1")  # large-context model (~150k tokens)
DEFAULT_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4.1-mini")
MAX_OUTPUT_TOKENS = int(os.environ.get("OPENAI_MAX_OUTPUT_TOKENS", "8000"))
TEMPERATURE = 0.0

# Ensure tesseract path if present
tesseract_path = shutil.which("tesseract")
if tesseract_path:
    pytesseract.pytesseract.tesseract_cmd = tesseract_path

# -------------------------
# OpenAI helpers
# -------------------------

@st.cache_resource(show_spinner=False)
def build_openai_client():
    api_key = st.secrets.get("OPENAI_API_KEY", os.environ.get("OPENAI_API_KEY"))
    if not api_key:
        st.error("Missing OpenAI API key. Set in Streamlit secrets as OPENAI_API_KEY or env var.")
        st.stop()
    client = OpenAI(api_key=api_key)
    return client

def chat_json(
    client: OpenAI,
    model: str,
    system_prompt: str,
    user_prompt: str,
    max_tokens: int = MAX_OUTPUT_TOKENS,
    temperature: float = TEMPERATURE,
) -> str:
    """
    Calls Chat Completions with response_format=json_object to force strict JSON.
    Returns raw JSON string from the assistant.
    """
    resp = client.chat.completions.create(
        model=model,
        temperature=temperature,
        max_tokens=max_tokens,
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": system_prompt.strip()},
            {"role": "user", "content": user_prompt.strip()},
        ],
    )
    return (resp.choices[0].message.content or "").strip()

def safe_json_loads(txt: str) -> Optional[dict]:
    if not txt:
        return None
    try:
        return json.loads(txt)
    except Exception:
        # last-resort: try to extract {...}
        s = txt.find("{"); e = txt.rfind("}")
        if s != -1 and e != -1:
            cand = txt[s:e+1]
            cand = re.sub(r",\s*}", "}", cand)
            cand = re.sub(r",\s*]", "]", cand)
            try:
                return json.loads(cand)
            except Exception:
                return None
        return None

# -------------------------
# PDF → Text extraction
# -------------------------
def pdf_bytes_to_images(pdf_bytes: bytes, dpi: int = 200) -> List[Image.Image]:
    return convert_from_bytes(pdf_bytes, dpi=dpi)

def ocr_image_to_text(image: Image.Image) -> str:
    return pytesseract.image_to_string(image)

def extract_text_from_pdf_bytes(pdf_bytes: bytes, dpi: int = 200) -> str:
    pages = pdf_bytes_to_images(pdf_bytes, dpi=dpi)
    texts = []
    for i, img in enumerate(pages):
        page_text = ocr_image_to_text(img)
        header = f"\n\n--- PAGE {i+1} ---\n\n"
        texts.append(header + page_text)
    return "\n".join(texts)

def extract_text_prefer_pdfplumber(pdf_bytes: bytes, dpi: int = 200):
    """
    Use pdfplumber to extract selectable text per page; fallback to OCR for pages with no text.
    Returns: (pages_list, full_text)
      pages_list: [{'page': n, 'text': str, 'is_ocr': bool}, ...]
      full_text: string with --- PAGE n --- markers
    """
    pages = []
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                try:
                    txt = page.extract_text() or ""
                except Exception:
                    txt = ""
                if txt.strip():
                    pages.append({"page": i, "text": txt, "is_ocr": False})
                else:
                    try:
                        pil_img = page.to_image(resolution=dpi).original
                        ocr_txt = pytesseract.image_to_string(pil_img)
                    except Exception:
                        ocr_txt = ""
                    pages.append({"page": i, "text": ocr_txt, "is_ocr": True})
    except Exception as e:
        logging.warning("pdfplumber failed: %s. Falling back to OCR for whole PDF.", e)
        ocr_full = extract_text_from_pdf_bytes(pdf_bytes, dpi=dpi)
        page_texts = [p for p in ocr_full.split("\f")] if "\f" in ocr_full else [ocr_full]
        for idx, p_text in enumerate(page_texts, start=1):
            pages.append({"page": idx, "text": p_text, "is_ocr": True})

    full = []
    for p in pages:
        full.append(f"\n\n--- PAGE {p['page']} ---\n{p['text']}")
    return pages, "\n".join(full)

# -------------------------
# Prompts (ONE-SHOT, full contract) — UPDATED CAD prompts
# -------------------------
CAD_SYSTEM = """
You are a senior contracts analyst.
Output STRICT JSON ONLY (no markdown, no comments, no prose outside the JSON).
Never hallucinate; if not explicitly supported by contract text, use null.
Prefer concise phrasing. Normalize dates to YYYY-MM-DD when explicit.
For every factual extraction, attach a sources array with page/ clause markers (e.g., "page 7", "clause 37.4").
"""

CAD_STYLE_GUIDE = """
Sections and order (must match exactly):
1. Salient Features of the Contract
2. List of Important Submittals
3. Notice / Information Clauses
4. Important clauses pertaining to project progress - EOT, Escalation, Variation, Suspension, etc.
5. Payment Clause
6. Identify Risk, Risk Matrix, and Risk Allocation
7. Claim, Dispute, and Arbitration clause
"""

CAD_USER_TMPL = """Use ONLY the CONTRACT TEXT to produce ONE JSON object exactly matching the schema.
Respond with strict JSON (start with '{{' and end with '}}'). No extra text.

CONTRACT TEXT (page markers included):
{contract_text}

STYLE GUIDE:
{style_guide}

SCHEMA:
{{
  "salient_features": [
    {{"description":"string","clause_or_ref":"string|null","details":"string|null","sources":["page X"]}}
  ],
  "important_submittals": [
    {{"stage":"before|during|after|null","document":"string","due":"string|null","notes":"string|null","clause_or_ref":"string|null","sources":["page X"]}}
  ],
  "notice_information_clauses": [
    {{"description":"string","clause_or_ref":"string|null","details":"string","sources":["page X"]}}
  ],
  "project_progress_clauses": [
    {{"topic":"string","clause_or_ref":"string|null","summary":"string","sources":["page X"]}}
  ],
  "payment": {{
    "type_of_contract":"string|null",
    "mode_of_payment":"string|null",
    "measurement_and_evaluation":"string|null",
    "billing_timing":"string|null",
    "other_terms":"string|null",
    "contract_sum": {{"amount": number|null, "currency":"string|null","sources":["page X"]}},
    "mobilization_advance": {{"applicable": true|false, "amount_pct": number|null, "interest":"string|null", "security":"string|null", "release_conditions":"string|null","sources":["page X"]}},
    "retention": {{"percent": number|null,"release_condition":"string|null","sources":["page X"]}},
    "interim_payments": [{{"frequency":"string|null","certifier":"string|null","payment_days_after_certification": number|null,"sources":["page X"]}}],
    "price_escalation": {{"applicable": true|false, "basis":"string|null", "formula_or_reference":"string|null","sources":["page X"]}},
    "final_payment": {{"required_docs":["string"],"release_days_after_submission": number|null,"sources":["page X"]}},
    "sources":["page X"]
  }},
  "risk_matrix": {{
    "severity_buckets":[{{"label":"string","criteria":"string"}}],
    "probability_buckets":[{{"label":"string","criteria":"string"}}],
    "matrix_legend": "string|null"
  }},
  "risks_and_allocation": [
    {{"id":"string","category":"string","risk_element":"string","probability":"Most Likely|Likely|Occasional|Unlikely|Remote|null","severity":"Very Severe|Major|Significant|Minor|Insignificant|null","rating":"Low|Medium|High|Critical|null","owner":"Employer|Contractor|Shared|null","mitigation":"string|null","sources":["page X"]}}
  ],
  "claims_disputes_arbitration": {{
    "arbitration": {{"applicable": true|false,"forum_or_rules":"string|null","notes":"string|null","sources":["page X"]}},
    "court_jurisdiction":"string|null",
    "claims_summary":[{{"topic":"string","payable_or_not":"Payable|Not Payable|Unclear","notes":"string","sources":["page X"]}}],
    "dispute_areas":[{{"description":"string","sources":["page X"]}}],
    "excusable_delays":[{{"item":"string","sources":["page X"]}}],
    "delay_compensation_clause":"string|null",
    "sources":["page X"]
  }}
}}

CRITICAL RULES:
- Fill fields ONLY if supported by the contract text. If unknown, use null.
- Every extracted fact MUST include "sources" with page and, if present, clause references.
- Keep text concise and specific (no generic boilerplate). Use lists instead of long paragraphs where helpful.
- Do NOT invent clause numbers; if absent, keep clause_or_ref null and rely on page numbers in "sources".
"""

COMPLIANCE_SYSTEM = """You are a contracts compliance checker. Always return STRICT JSON (no markdown)."""

COMPLIANCE_USER_TMPL = """Use ONLY the CONTRACT TEXT to evaluate the rule. Return ONE strict JSON object.

CONTRACT TEXT:
{contract_text}

RULE:
{rule}

Return JSON in one of these shapes:
{{"rule":"{rule}","present": true,"summary":"short","quote":"<=200 chars verbatim","sources":["page X","clause Y"],"confidence":0.9}}
or
{{"rule":"{rule}","present": false,"summary": null,"quote": null,"sources": [],"confidence": 0.99}}
"""

QA_SYSTEM = """You are a contracts Q&A assistant. Answer ONLY from the contract text."""

QA_USER_TMPL = """CONTRACT TEXT:
{contract_text}

QUESTION:
{question}

If unknown, reply exactly: "Not found in contract."
"""
# --- ReportLab Platypus (table PDF) ---
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# --- Robust rendering helpers ---
def _get(d, key, default=""):
    """Safe dict getter."""
    return d.get(key, default) if isinstance(d, dict) else default

def _as_list_of_dicts(seq, primary_key):
    """
    Normalize a sequence to a list of dicts.
    - dict -> keep
    - str  -> wrap as {primary_key: str}
    - else -> skip
    """
    out = []
    if not isinstance(seq, (list, tuple)):
        return out
    for it in seq:
        if isinstance(it, dict):
            out.append(it)
        elif isinstance(it, str):
            out.append({primary_key: it})
    return out

# -------------------------
# One-shot CAD generation (UPDATED)
# -------------------------
# def node_generate_cad_json_docx_pdf_openai(
#     client: OpenAI,
#     model: str,
#     full_text_with_pages: str,
#     pages,
#     output_basename: str = "generated_CAD"
# ):
#     # 0) Trim obvious whitespace noise to save tokens (but keep page markers)
#     compact_text = re.sub(r"[ \t]+", " ", full_text_with_pages)
#     compact_text = re.sub(r"\n{3,}", "\n\n", compact_text)

#     # 1) Call LLM once for the whole contract
#     user_prompt = CAD_USER_TMPL.format(
#         contract_text=compact_text,
#         style_guide=CAD_STYLE_GUIDE.strip()
#     )
#     raw = chat_json(client, model, CAD_SYSTEM, user_prompt, max_tokens=MAX_OUTPUT_TOKENS)

#     # 2) Parse/repair JSON once; then a stricter repair if needed
#     data = safe_json_loads(raw)
#     if not data or not isinstance(data, dict):
#         repair_prompt = (
#             "Convert this content into ONE strict JSON object that matches the previously provided schema. "
#             "Do not add comments or prose—JSON only.\n\n" + (raw or "")
#         )
#         repaired = chat_json(client, model, CAD_SYSTEM, repair_prompt, max_tokens=MAX_OUTPUT_TOKENS)
#         data = safe_json_loads(repaired)
#         if not data:
#             # Last resort: ask to re-emit from source with the schema only
#             reemit_prompt = (
#                 "Re-extract from the CONTRACT TEXT using the exact schema previously given. "
#                 "Output STRICT JSON only. If a field is unknown, set it to null.\n\n"
#                 f"CONTRACT TEXT:\n{compact_text}\n\nSTYLE GUIDE:\n{CAD_STYLE_GUIDE.strip()}"
#             )
#             reemitted = chat_json(client, model, CAD_SYSTEM, reemit_prompt, max_tokens=MAX_OUTPUT_TOKENS)
#             data = safe_json_loads(reemitted)
#             if not data:
#                 raise RuntimeError("Failed to obtain valid CAD JSON from LLM.")

#     # 3) Save JSON
#     json_path = Path(tempfile.gettempdir()) / f"{output_basename}.json"
#     with open(json_path, "w", encoding="utf8") as f:
#         json.dump(data, f, ensure_ascii=False, indent=2)

#     # 4) DOCX rendering — mirrors the sample CAD layout
#     docx = DocxDocument()
#     style = docx.styles["Normal"]; style.font.name = "Arial"; style.font.size = Pt(11)

#     # Title
#     docx.add_paragraph().add_run().add_break()
#     docx.add_heading("Contract Appreciation Document (Generated)", level=1)

#     # ---- 1. Salient Features of the Contract ----
#     docx.add_heading("1. Salient Features of the Contract", level=2)
#     sf_rows = data.get("salient_features", []) or []
#     if sf_rows:
#         table = docx.add_table(rows=1, cols=4)
#         hdrs = ["Sl. No.", "Description", "Clause No.", "Details"]
#         for i, h in enumerate(hdrs): table.rows[0].cells[i].text = h
#         for idx, r in enumerate(sf_rows, start=1):
#             row = table.add_row().cells
#             row[0].text = str(idx)
#             row[1].text = r.get("description") or ""
#             row[2].text = r.get("clause_or_ref") or ""
#             row[3].text = r.get("details") or ""
#     else:
#         docx.add_paragraph("No salient features found.")

#     # ---- 2. List of Important Submittals ----
#     docx.add_heading("2. List of Important Submittals", level=2)
#     subs = data.get("important_submittals", []) or []
#     if subs:
#         table = docx.add_table(rows=1, cols=6)
#         for i, h in enumerate(["Sl. No.", "Stage", "Document", "Due", "Notes", "Clause No."]):
#             table.rows[0].cells[i].text = h
#         for idx, s in enumerate(subs, start=1):
#             c = table.add_row().cells
#             c[0].text = str(idx)
#             c[1].text = (s.get("stage") or "")
#             c[2].text = (s.get("document") or "")
#             c[3].text = (s.get("due") or "")
#             c[4].text = (s.get("notes") or "")
#             c[5].text = (s.get("clause_or_ref") or "")
#     else:
#         docx.add_paragraph("No submittals listed.")

#     # ---- 3. Notice / Information Clauses ----
#     docx.add_heading("3. Notice / Information Clauses", level=2)
#     notices = data.get("notice_information_clauses", []) or []
#     if notices:
#         table = docx.add_table(rows=1, cols=4)
#         for i, h in enumerate(["Sl. No.", "Description", "Clause No.", "Details"]):
#             table.rows[0].cells[i].text = h
#         for idx, n in enumerate(notices, start=1):
#             c = table.add_row().cells
#             c[0].text = str(idx)
#             c[1].text = n.get("description") or ""
#             c[2].text = n.get("clause_or_ref") or ""
#             c[3].text = n.get("details") or ""
#     else:
#         docx.add_paragraph("No notice/information clauses found.")

#     # ---- 4. Project Progress (EOT, Escalation, Variation, Suspension, etc.) ----
#     docx.add_heading("4. Important clauses pertaining to project progress - EOT, Escalation, Variation, Suspension, etc.", level=2)
#     ppcs = data.get("project_progress_clauses", []) or []
#     if ppcs:
#         table = docx.add_table(rows=1, cols=4)
#         for i, h in enumerate(["Sl. No.", "Topic", "Clause No.", "Summary"]):
#             table.rows[0].cells[i].text = h
#         for idx, ppc in enumerate(ppcs, start=1):
#             c = table.add_row().cells
#             c[0].text = str(idx)
#             c[1].text = ppc.get("topic") or ""
#             c[2].text = ppc.get("clause_or_ref") or ""
#             c[3].text = ppc.get("summary") or ""
#     else:
#         docx.add_paragraph("No project progress clauses found.")

#     # ---- 5. Payment Clause ----
#     docx.add_heading("5. Payment Clause", level=2)
#     pay = data.get("payment", {}) or {}
#     def _p(label, val): docx.add_paragraph(f"{label}: {val or 'N/A'}")
#     _p("Type of Contract", pay.get("type_of_contract"))
#     _p("Mode of Payment", pay.get("mode_of_payment"))
#     _p("Measurement and Evaluations", pay.get("measurement_and_evaluation"))
#     _p("Billing Timing", pay.get("billing_timing"))
#     _p("Other Terms", pay.get("other_terms"))
#     # Key money terms compactly:
#     money_lines = []
#     if isinstance(pay.get("contract_sum"), dict):
#         cs = pay["contract_sum"]
#         if cs.get("amount") is not None:
#             money_lines.append(f"Contract Sum: {cs.get('amount')} {cs.get('currency') or ''}".strip())
#     if isinstance(pay.get("mobilization_advance"), dict):
#         ma = pay["mobilization_advance"]
#         money_lines.append(
#             f"Mobilization Advance: {'Yes' if ma.get('applicable') else 'No'}"
#             + (f", {ma.get('amount_pct')}% @ {ma.get('interest')}" if ma.get('applicable') and ma.get('amount_pct') is not None else "")
#         )
#     if isinstance(pay.get("retention"), dict):
#         rt = pay["retention"]
#         if rt.get("percent") is not None:
#             money_lines.append(f"Retention: {rt.get('percent')}% ({rt.get('release_condition') or 'release terms N/A'})")
#     if isinstance(pay.get("price_escalation"), dict):
#         es = pay["price_escalation"]
#         money_lines.append(f"Price Escalation: {'Applicable' if es.get('applicable') else 'Not applicable'}"
#                            + (f" ({es.get('basis') or es.get('formula_or_reference')})" if es.get('applicable') else ""))
#     if money_lines:
#         for ln in money_lines: docx.add_paragraph("• " + ln)
#     ips = pay.get("interim_payments") or []
#     if ips:
#         docx.add_paragraph("Interim Payments:")
#         for ip in ips:
#             docx.add_paragraph(f"  - {ip.get('frequency') or 'N/A'}; Certifier: {ip.get('certifier') or 'N/A'}; "
#                                f"DAYS after certification: {ip.get('payment_days_after_certification') or 'N/A'}")

#     # ---- 6. Risk Matrix & Allocation ----
#     docx.add_heading("6. Identify Risk, Risk Matrix, and Risk Allocation", level=2)
#     rm = data.get("risk_matrix", {}) or {}
#     sev = rm.get("severity_buckets") or []
#     prob = rm.get("probability_buckets") or []
#     if sev:
#         docx.add_paragraph("Severity Buckets:")
#         for s in sev: docx.add_paragraph(f"  - {s.get('label')}: {s.get('criteria')}")
#     if prob:
#         docx.add_paragraph("Probability Buckets:")
#         for p in prob: docx.add_paragraph(f"  - {p.get('label')}: {p.get('criteria')}")
#     risks = data.get("risks_and_allocation", []) or []
#     if risks:
#         table = docx.add_table(rows=1, cols=9)
#         hdrs = ["ID","Category","Risk Element","Probability","Severity","Rating","Owner","Mitigation","Sources"]
#         for i, h in enumerate(hdrs): table.rows[0].cells[i].text = h
#         for r in risks:
#             c = table.add_row().cells
#             c[0].text = r.get("id") or ""
#             c[1].text = r.get("category") or ""
#             c[2].text = r.get("risk_element") or ""
#             c[3].text = r.get("probability") or ""
#             c[4].text = r.get("severity") or ""
#             c[5].text = r.get("rating") or ""
#             c[6].text = r.get("owner") or ""
#             c[7].text = r.get("mitigation") or ""
#             c[8].text = ", ".join(r.get("sources") or [])
#     else:
#         docx.add_paragraph("No risks listed.")

#     # ---- 7. Claims, Dispute, and Arbitration ----
#     docx.add_heading("7. Claim, Dispute, and Arbitration clause", level=2)
#     cda = data.get("claims_disputes_arbitration", {}) or {}
#     def _pp(label, val):
#         if val: docx.add_paragraph(f"{label}: {val}")
#     arb = cda.get("arbitration") or {}
#     _pp("Arbitration", f"Applicable: {'Yes' if arb.get('applicable') else 'No'}; Forum/Rules: {arb.get('forum_or_rules') or 'N/A'}")
#     _pp("Court Jurisdiction", cda.get("court_jurisdiction") or "N/A")
#     claims = cda.get("claims_summary") or []
#     if claims:
#         docx.add_paragraph("Claims Summary:")
#         for cl in claims:
#             docx.add_paragraph(f"  - {cl.get('topic')}: {cl.get('payable_or_not') or 'Unclear'} — {cl.get('notes') or ''}")
#     disputes = cda.get("dispute_areas") or []
#     if disputes:
#         docx.add_paragraph("Dispute Areas:")
#         for d in disputes:
#             docx.add_paragraph(f"  - {d.get('description')}")
#     exds = cda.get("excusable_delays") or []
#     if exds:
#         docx.add_paragraph("Excusable Delays:")
#         for e in exds:
#             docx.add_paragraph(f"  - {e.get('item')}")
#     _pp("Delay Compensation Clause", cda.get("delay_compensation_clause"))

#     # Save DOCX
#     docx_path = Path(tempfile.gettempdir()) / f"{output_basename}.docx"
#     docx.save(docx_path)

#     # # 5) 1-page Summary PDF (compact)
#     # # pdf_path = Path(tempfile.gettempdir()) / f"{output_basename}.pdf"
#     # # c = canvas.Canvas(str(pdf_path), pagesize=A4)
#     # # width, height = A4
#     # # margin = 50
#     # # y = height - margin
#     from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
#     from reportlab.lib import colors
#     from reportlab.lib.styles import getSampleStyleSheet
#     # 5) 1-page Summary PDF (tabular, Platypus)\
#     # 5) 1-page Summary PDF (tabular, Platypus only — no canvas usage)

#     pdf_path = Path(tempfile.gettempdir()) / f"{output_basename}.pdf"

#     doc = SimpleDocTemplate(
#         str(pdf_path),
#         pagesize=A4,
#         rightMargin=40, leftMargin=40, topMargin=60, bottomMargin=40
#     )

#     styles = getSampleStyleSheet()
#     styles["Normal"].fontName = "Helvetica"
#     styles["Normal"].fontSize = 10
#     styles["Normal"].leading = 12

#     h2 = ParagraphStyle(
#         name="H2", parent=styles["Heading2"], fontName="Helvetica-Bold",
#         spaceAfter=8
#     )
#     h3 = ParagraphStyle(
#         name="H3", parent=styles["Heading3"], fontName="Helvetica-Bold",
#         spaceBefore=6, spaceAfter=6
#     )

#     elements = []
#     elements.append(Paragraph("Contract Appreciation Document (Summary)", h2))
#     elements.append(Spacer(1, 6))
#     elements.append(Paragraph("Key Salient Features", h3))
#     elements.append(Spacer(1, 4))

#     # Pull salient features (already built above)
#     sf_rows = data.get("salient_features", []) or []
#     table_data = [["Sl. No.", "Description", "Details"]]
#     for i, row in enumerate(sf_rows[:6], start=1):
#         desc = (row.get("description") or "").strip()
#         det  = (row.get("details") or "").strip()
#         table_data.append([
#             str(i),
#             Paragraph(desc, styles["Normal"]),
#             Paragraph(det,  styles["Normal"]),
#         ])

#     table = Table(table_data, colWidths=[45, 210, 285], repeatRows=1)
#     table.setStyle(TableStyle([
#         ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#F0F0F0")),
#         ("GRID", (0,0), (-1,-1), 0.5, colors.black),
#         ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
#         ("ALIGN", (0,0), (0,-1), "CENTER"),  # Sl. No.
#         ("VALIGN", (0,0), (-1,-1), "TOP"),
#         ("LEFTPADDING", (0,0), (-1,-1), 6),
#         ("RIGHTPADDING", (0,0), (-1,-1), 6),
#         ("TOPPADDING", (0,0), (-1,-1), 4),
#         ("BOTTOMPADDING", (0,0), (-1,-1), 4),
#     ]))
#     elements.append(table)

#     doc.build(elements)

def node_generate_cad_json_docx_pdf_openai(
    client: OpenAI,
    model: str,
    full_text_with_pages: str,
    pages,
    output_basename: str = "generated_CAD"
):
    # 0) Trim obvious whitespace noise to save tokens (but keep page markers)
    compact_text = re.sub(r"[ \t]+", " ", full_text_with_pages)
    compact_text = re.sub(r"\n{3,}", "\n\n", compact_text)

    # 1) Call LLM once for the whole contract
    user_prompt = CAD_USER_TMPL.format(
        contract_text=compact_text,
        style_guide=CAD_STYLE_GUIDE.strip()
    )
    raw = chat_json(client, model, CAD_SYSTEM, user_prompt, max_tokens=MAX_OUTPUT_TOKENS)

    # 2) Parse/repair JSON once; then a stricter repair if needed
    data = safe_json_loads(raw)
    if not data or not isinstance(data, dict):
        repair_prompt = (
            "Convert this content into ONE strict JSON object that matches the previously provided schema. "
            "Do not add comments or prose—JSON only.\n\n" + (raw or "")
        )
        repaired = chat_json(client, model, CAD_SYSTEM, repair_prompt, max_tokens=MAX_OUTPUT_TOKENS)
        data = safe_json_loads(repaired)
        if not data:
            # Last resort: re-extract from source with schema
            reemit_prompt = (
                "Re-extract from the CONTRACT TEXT using the exact schema previously given. "
                "Output STRICT JSON only. If a field is unknown, set it to null.\n\n"
                f"CONTRACT TEXT:\n{compact_text}\n\nSTYLE GUIDE:\n{CAD_STYLE_GUIDE.strip()}"
            )
            reemitted = chat_json(client, model, CAD_SYSTEM, reemit_prompt, max_tokens=MAX_OUTPUT_TOKENS)
            data = safe_json_loads(reemitted)
            if not data:
                raise RuntimeError("Failed to obtain valid CAD JSON from LLM.")

    # 3) Save JSON
    json_path = Path(tempfile.gettempdir()) / f"{output_basename}.json"
    with open(json_path, "w", encoding="utf8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    # -------- Normalize all sections defensively --------
    salient_features = _as_list_of_dicts(data.get("salient_features", []), "description")
    important_submittals = _as_list_of_dicts(data.get("important_submittals", []), "document")
    notice_information = _as_list_of_dicts(data.get("notice_information_clauses", []), "description")
    project_progress = _as_list_of_dicts(data.get("project_progress_clauses", []), "topic")

    pay = data.get("payment", {})
    pay = pay if isinstance(pay, dict) else {}

    risk_matrix = data.get("risk_matrix", {})
    risk_matrix = risk_matrix if isinstance(risk_matrix, dict) else {}
    sev_buckets = _as_list_of_dicts(risk_matrix.get("severity_buckets", []), "label")
    prob_buckets = _as_list_of_dicts(risk_matrix.get("probability_buckets", []), "label")

    risks = _as_list_of_dicts(data.get("risks_and_allocation", []), "risk_element")

    cda = data.get("claims_disputes_arbitration", {})
    cda = cda if isinstance(cda, dict) else {}
    arbitration = cda.get("arbitration", {})
    arbitration = arbitration if isinstance(arbitration, dict) else {}
    claims = _as_list_of_dicts(cda.get("claims_summary", []), "topic")
    disputes = _as_list_of_dicts(cda.get("dispute_areas", []), "description")
    excusable_delays = _as_list_of_dicts(cda.get("excusable_delays", []), "item")

    # 4) DOCX rendering — mirrors the sample CAD layout
    docx = DocxDocument()
    style = docx.styles["Normal"]; style.font.name = "Arial"; style.font.size = Pt(11)

    # Title
    docx.add_paragraph().add_run().add_break()
    docx.add_heading("Contract Appreciation Document (Generated)", level=1)

    # ---- 1. Salient Features of the Contract ----
    docx.add_heading("1. Salient Features of the Contract", level=2)
    if salient_features:
        table = docx.add_table(rows=1, cols=4)
        hdrs = ["Sl. No.", "Description", "Clause No.", "Details"]
        for i, h in enumerate(hdrs): table.rows[0].cells[i].text = h
        for idx, r in enumerate(salient_features, start=1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = _get(r, "description")
            row[2].text = _get(r, "clause_or_ref")
            row[3].text = _get(r, "details")
    else:
        docx.add_paragraph("No salient features found.")

    # ---- 2. List of Important Submittals ----
    docx.add_heading("2. List of Important Submittals", level=2)
    if important_submittals:
        table = docx.add_table(rows=1, cols=6)
        for i, h in enumerate(["Sl. No.", "Stage", "Document", "Due", "Notes", "Clause No."]):
            table.rows[0].cells[i].text = h
        for idx, s in enumerate(important_submittals, start=1):
            c = table.add_row().cells
            c[0].text = str(idx)
            c[1].text = _get(s, "stage")
            c[2].text = _get(s, "document")
            c[3].text = _get(s, "due")
            c[4].text = _get(s, "notes")
            c[5].text = _get(s, "clause_or_ref")
    else:
        docx.add_paragraph("No submittals listed.")

    # ---- 3. Notice / Information Clauses ----
    docx.add_heading("3. Notice / Information Clauses", level=2)
    if notice_information:
        table = docx.add_table(rows=1, cols=4)
        for i, h in enumerate(["Sl. No.", "Description", "Clause No.", "Details"]):
            table.rows[0].cells[i].text = h
        for idx, n in enumerate(notice_information, start=1):
            c = table.add_row().cells
            c[0].text = str(idx)
            c[1].text = _get(n, "description")
            c[2].text = _get(n, "clause_or_ref")
            c[3].text = _get(n, "details")
    else:
        docx.add_paragraph("No notice/information clauses found.")

    # ---- 4. Project Progress (EOT, Escalation, Variation, Suspension, etc.) ----
    docx.add_heading("4. Important clauses pertaining to project progress - EOT, Escalation, Variation, Suspension, etc.", level=2)
    if project_progress:
        table = docx.add_table(rows=1, cols=4)
        for i, h in enumerate(["Sl. No.", "Topic", "Clause No.", "Summary"]):
            table.rows[0].cells[i].text = h
        for idx, ppc in enumerate(project_progress, start=1):
            c = table.add_row().cells
            c[0].text = str(idx)
            c[1].text = _get(ppc, "topic")
            c[2].text = _get(ppc, "clause_or_ref")
            c[3].text = _get(ppc, "summary")
    else:
        docx.add_paragraph("No project progress clauses found.")

    # ---- 5. Payment Clause ----
    docx.add_heading("5. Payment Clause", level=2)
    def _p(label, val): docx.add_paragraph(f"{label}: {val if (val or val==0) else 'N/A'}")
    _p("Type of Contract", _get(pay, "type_of_contract"))
    _p("Mode of Payment", _get(pay, "mode_of_payment"))
    _p("Measurement and Evaluations", _get(pay, "measurement_and_evaluation"))
    _p("Billing Timing", _get(pay, "billing_timing"))
    _p("Other Terms", _get(pay, "other_terms"))

    money_lines = []
    cs = pay.get("contract_sum") if isinstance(pay.get("contract_sum"), dict) else {}
    if isinstance(cs, dict) and (cs.get("amount") is not None):
        money_lines.append(f"Contract Sum: {cs.get('amount')} {cs.get('currency') or ''}".strip())

    ma = pay.get("mobilization_advance") if isinstance(pay.get("mobilization_advance"), dict) else {}
    if isinstance(ma, dict):
        if ma.get("applicable") is True:
            if ma.get("amount_pct") is not None:
                money_lines.append(f"Mobilization Advance: Yes, {ma.get('amount_pct')}% @ {ma.get('interest') or 'N/A'}")
            else:
                money_lines.append("Mobilization Advance: Yes")
        elif ma.get("applicable") is False:
            money_lines.append("Mobilization Advance: No")

    rt = pay.get("retention") if isinstance(pay.get("retention"), dict) else {}
    if isinstance(rt, dict) and (rt.get("percent") is not None):
        money_lines.append(f"Retention: {rt.get('percent')}% ({rt.get('release_condition') or 'release terms N/A'})")

    es = pay.get("price_escalation") if isinstance(pay.get("price_escalation"), dict) else {}
    if isinstance(es, dict):
        if es.get("applicable") is True:
            basis = es.get("basis") or es.get("formula_or_reference")
            money_lines.append(f"Price Escalation: Applicable ({basis or 'basis/formula N/A'})")
        elif es.get("applicable") is False:
            money_lines.append("Price Escalation: Not applicable")

    for ln in money_lines:
        docx.add_paragraph("• " + ln)

    ips = pay.get("interim_payments") if isinstance(pay.get("interim_payments"), list) else []
    if ips:
        docx.add_paragraph("Interim Payments:")
        for ip in ips:
            ip = ip if isinstance(ip, dict) else {}
            docx.add_paragraph(
                f"  - {_get(ip, 'frequency', 'N/A')}; "
                f"Certifier: {_get(ip, 'certifier', 'N/A')}; "
                f"DAYS after certification: {_get(ip, 'payment_days_after_certification', 'N/A')}"
            )

    # ---- 6. Risk Matrix & Allocation ----
    docx.add_heading("6. Identify Risk, Risk Matrix, and Risk Allocation", level=2)
    if sev_buckets:
        docx.add_paragraph("Severity Buckets:")
        for s in sev_buckets: docx.add_paragraph(f"  - {_get(s,'label')}: {_get(s,'criteria')}")
    if prob_buckets:
        docx.add_paragraph("Probability Buckets:")
        for p in prob_buckets: docx.add_paragraph(f"  - {_get(p,'label')}: {_get(p,'criteria')}")

    if risks:
        table = docx.add_table(rows=1, cols=9)
        hdrs = ["ID","Category","Risk Element","Probability","Severity","Rating","Owner","Mitigation","Sources"]
        for i, h in enumerate(hdrs): table.rows[0].cells[i].text = h
        for r in risks:
            c = table.add_row().cells
            c[0].text = _get(r, "id")
            c[1].text = _get(r, "category")
            c[2].text = _get(r, "risk_element")
            c[3].text = _get(r, "probability")
            c[4].text = _get(r, "severity")
            c[5].text = _get(r, "rating")
            c[6].text = _get(r, "owner")
            c[7].text = _get(r, "mitigation")
            c[8].text = ", ".join(r.get("sources", [])) if isinstance(r, dict) else ""
    else:
        docx.add_paragraph("No risks listed.")

    # ---- 7. Claims, Dispute, and Arbitration ----
    docx.add_heading("7. Claim, Dispute, and Arbitration clause", level=2)
    docx.add_paragraph(
        f"Arbitration: Applicable: {'Yes' if _get(arbitration,'applicable') else 'No'}; "
        f"Forum/Rules: {_get(arbitration,'forum_or_rules','N/A')}"
    )
    docx.add_paragraph(f"Court Jurisdiction: {_get(cda,'court_jurisdiction','N/A')}")

    if claims:
        docx.add_paragraph("Claims Summary:")
        for cl in claims:
            topic = _get(cl, "topic")
            payable = _get(cl, "payable_or_not", "Unclear")
            notes = _get(cl, "notes")
            docx.add_paragraph(f"  - {topic}: {payable} — {notes}")

    if disputes:
        docx.add_paragraph("Dispute Areas:")
        for d in disputes:
            docx.add_paragraph(f"  - {_get(d, 'description')}")

    if excusable_delays:
        docx.add_paragraph("Excusable Delays:")
        for e in excusable_delays:
            docx.add_paragraph(f"  - {_get(e, 'item')}")

    delay_comp = _get(cda, "delay_compensation_clause")
    if delay_comp:
        docx.add_paragraph(f"Delay Compensation Clause: {delay_comp}")

    # Save DOCX
    docx_path = Path(tempfile.gettempdir()) / f"{output_basename}.docx"
    docx.save(docx_path)

    # 5) 1-page Summary PDF (tabular, Platypus only — no canvas usage)
    pdf_path = Path(tempfile.gettempdir()) / f"{output_basename}.pdf"
    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        rightMargin=40, leftMargin=40, topMargin=60, bottomMargin=40
    )
    styles = getSampleStyleSheet()
    styles["Normal"].fontName = "Helvetica"
    styles["Normal"].fontSize = 10
    styles["Normal"].leading = 12
    h2 = ParagraphStyle(name="H2", parent=styles["Heading2"], fontName="Helvetica-Bold", spaceAfter=8)
    h3 = ParagraphStyle(name="H3", parent=styles["Heading3"], fontName="Helvetica-Bold", spaceBefore=6, spaceAfter=6)

    elements = []
    elements.append(Paragraph("Contract Appreciation Document (Summary)", h2))
    elements.append(Spacer(1, 6))
    elements.append(Paragraph("Key Salient Features", h3))
    elements.append(Spacer(1, 4))

    # Use normalized salient_features
    table_data = [["Sl. No.", "Description", "Details"]]
    for i, row in enumerate(salient_features[:6], start=1):
        desc = _get(row, "description").strip()
        det  = _get(row, "details").strip()
        table_data.append([str(i), Paragraph(desc, styles["Normal"]), Paragraph(det, styles["Normal"])])

    table = Table(table_data, colWidths=[45, 210, 285], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#F0F0F0")),
        ("GRID", (0,0), (-1,-1), 0.5, colors.black),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("ALIGN", (0,0), (0,-1), "CENTER"),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("LEFTPADDING", (0,0), (-1,-1), 6),
        ("RIGHTPADDING", (0,0), (-1,-1), 6),
        ("TOPPADDING", (0,0), (-1,-1), 4),
        ("BOTTOMPADDING", (0,0), (-1,-1), 4),
    ]))
    elements.append(table)
    doc.build(elements)

    return str(data), str(json_path), str(docx_path), str(pdf_path)





    # def write_wrapped_text(cobj, text, x, y, max_width, leading=12):
    #     from reportlab.lib.utils import simpleSplit
    #     from reportlab.lib.pagesizes import A4  # Import A4 to get dimensions
    #     width, height = A4  # Define width and height using A4
    #     lines = simpleSplit(text, "Helvetica", 10, max_width)
    #     for ln in lines:
    #         if y < margin + 40:
    #             cobj.showPage()
    #             y = height - margin
    #         cobj.drawString(x, y, ln); y -= leading
    #     return y

    # margin = 50  # Define margin here
    # c.setFont("Helvetica-Bold", 14)
    # c.drawString(margin, y, "Contract Appreciation Document (Summary)"); y -= 22
    # c.setFont("Helvetica", 10)
    # if sf_rows:
    #     y = write_wrapped_text(c, "Key Salient Features:", margin, y, width - 2*margin)
    #     for i, row in enumerate(sf_rows[:6], start=1):
    #         desc = row.get("description") or ""
    #         details = row.get("details") or ""
    #         y = write_wrapped_text(c, f"  {i}. {desc}: {details}", margin, y, width - 2*margin)
    # c.showPage(); c.save()

    # return str(data), str(json_path), str(docx_path), str(pdf_path)

# -------------------------
# Compliance (one-shot per rule) — UNCHANGED
# -------------------------
def compliance_check_json_openai(
    client: OpenAI,
    model: str,
    full_text_with_pages: str,
    rules: List[str]
) -> List[Dict[str, Any]]:
    results = []
    for rule in rules:
        uprompt = COMPLIANCE_USER_TMPL.format(contract_text=full_text_with_pages, rule=rule)
        raw = chat_json(client, model, COMPLIANCE_SYSTEM, uprompt, max_tokens=800)
        obj = safe_json_loads(raw)
        if not obj:
            # fallback: mark not present with low confidence
            obj = {"rule": rule, "present": False, "summary": None, "quote": None, "sources": [], "confidence": 0.50}
        results.append(obj)
    return results

# -------------------------
# Conflict detection (regex) — UNCHANGED
# -------------------------
def check_commencement_vs_site_possession(pages, full_text):
    date_pattern = r'((?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4})|(?:\d{1,2}\s+[A-Za-z]+\s+\d{4})|(?:[A-Za-z]+\s+\d{1,2},\s*\d{4}))'
    site_pos_pattern = re.compile(r'(site possession|possession of the site|hand over (?:the )?site)[^\.\n]{0,250}(' + date_pattern + r')', re.I)
    commence_pattern = re.compile(r'(commenc(?:e|ement) date|start date|date of commencement|date of commencement of works)[^\.\n]{0,250}(' + date_pattern + r')', re.I)

    site_matches, comm_matches = [], []
    for p in pages:
        text = p.get('text') or ""
        for m in site_pos_pattern.finditer(text):
            raw = m.group(2)
            try:
                dt = dateparser.parse(raw, dayfirst=True)
            except Exception:
                dt = None
            site_matches.append({"page": p['page'], "raw": raw, "date": dt, "snippet": m.group(0)})
        for m in commence_pattern.finditer(text):
            raw = m.group(2)
            try:
                dt = dateparser.parse(raw, dayfirst=True)
            except Exception:
                dt = None
            comm_matches.append({"page": p['page'], "raw": raw, "date": dt, "snippet": m.group(0)})

    conflicts = []
    if site_matches and comm_matches:
        comm_dates = [c['date'] for c in comm_matches if c['date']]
        site_dates = [s['date'] for s in site_matches if s['date']]
        if comm_dates and site_dates:
            min_comm = min(comm_dates)
            max_site = max(site_dates)
            if max_site > min_comm:
                conflicts.append({
                    "type": "commencement_vs_site_possession",
                    "severity": "Critical",
                    "message": f"Site possession ({max_site.date()}) occurs after commencement ({min_comm.date()}).",
                    "evidence": {"commencement": comm_matches, "site_possession": site_matches}
                })
    return conflicts

def check_payment_term_mismatch(pages):
    pay_pattern = re.compile(r'pay(?:ment)?\s+(?:within|in)\s+(\d{1,3})\s+days', re.I)
    matches = []
    for p in pages:
        for m in pay_pattern.finditer(p.get('text') or ""):
            try:
                matches.append({"page": p['page'], "days": int(m.group(1)), "snippet": m.group(0)})
            except Exception:
                pass
    conflicts = []
    if matches:
        days_set = sorted(set(m['days'] for m in matches))
        if len(days_set) > 1:
            conflicts.append({
                "type": "payment_term_mismatch",
                "severity": "High",
                "message": f"Different payment terms found: {days_set} days.",
                "evidence": matches
            })
    return conflicts

def check_retention_mismatch(pages):
    ret_pattern = re.compile(r'retention[^\d]{0,20}(\d{1,2})\s*%?', re.I)
    matches = []
    for p in pages:
        for m in ret_pattern.finditer(p.get('text') or ""):
            try:
                pct = int(m.group(1))
                matches.append({"page": p['page'], "pct": pct, "snippet": m.group(0)})
            except Exception:
                continue
    conflicts = []
    if matches:
        pct_set = sorted(set(m['pct'] for m in matches))
        if len(pct_set) > 1:
            conflicts.append({
                "type": "retention_mismatch",
                "severity": "High",
                "message": f"Different retention percentages found: {pct_set}%.",
                "evidence": matches
            })
    return conflicts

def check_defect_liability_mismatch(pages):
    dl_pattern = re.compile(r'(defect liability period|defects liability|warranty period)[^\d]{0,60}(\d{1,3})\s*(months|month|years|year)?', re.I)
    matches = []
    for p in pages:
        for m in dl_pattern.finditer(p.get('text') or ""):
            try:
                val = int(m.group(2))
                unit = (m.group(3) or "months").lower()
                months = val * 12 if "year" in unit else val
                matches.append({"page": p['page'], "value": months, "raw": m.group(0)})
            except Exception:
                continue
    conflicts = []
    if matches:
        vals = sorted(set(m['value'] for m in matches))
        if len(vals) > 1:
            conflicts.append({
                "type": "defect_liability_mismatch",
                "severity": "High",
                "message": f"Different defect liability lengths found (in months): {vals}.",
                "evidence": matches
            })
    return conflicts

def check_arbitration_vs_court_conflict(pages):
    arb_pattern = re.compile(r'\barbitrat', re.I)
    court_pattern = re.compile(r'\bcourt', re.I)
    arbs, courts = [], []
    for p in pages:
        t = p.get('text') or ""
        if arb_pattern.search(t):
            arbs.append({"page": p['page']})
        if court_pattern.search(t):
            courts.append({"page": p['page']})
    conflicts = []
    if arbs and courts:
        conflicts.append({
            "type": "dispute_resolution_conflict",
            "severity": "Critical",
            "message": "Arbitration and court references coexist; possible conflict in dispute resolution.",
            "evidence": {"arbitration_pages": arbs, "court_pages": courts}
        })
    return conflicts

CONFLICT_TO_CATEGORY = {
    "commencement_vs_site_possession": ("Contractual Disputes (Site Possession)", "Deem commencement as possession or grant EOT + cost; see Site Possession."),
    "payment_term_mismatch": ("Contractual Disputes (Payment)", "Harmonize payment timeline; prefer main body or latest-dated doc."),
    "retention_mismatch": ("Financial / Payment", "Clarify a single retention % and update annex."),
    "defect_liability_mismatch": ("Quality / Warranty", "Adopt stricter DLP or clarify governing schedule."),
    "dispute_resolution_conflict": ("Dispute Resolution", "Prefer arbitration with interim relief carve-out; remove ambiguity."),
}

def map_conflict_to_practical_category(conflict):
    ctype = conflict.get("type")
    cat, hint = CONFLICT_TO_CATEGORY.get(ctype, ("Other / Administrative", "Review conflict and map manually."))
    conflict["category"] = cat
    conflict["resolution_hint"] = hint
    return conflict

def run_conflict_detection(pages, full_text):
    conflicts = []
    conflicts.extend(check_commencement_vs_site_possession(pages, full_text))
    conflicts.extend(check_payment_term_mismatch(pages))
    conflicts.extend(check_retention_mismatch(pages))
    conflicts.extend(check_defect_liability_mismatch(pages))
    conflicts.extend(check_arbitration_vs_court_conflict(pages))
    conflicts = [map_conflict_to_practical_category(c) for c in conflicts]
    return conflicts

# -------------------------
# Streamlit UI
# -------------------------
st.set_page_config(page_title="Contract CAD (ChatGPT One-Shot)", layout="wide")
st.title("Automated Contract CAD Generator — ChatGPT (One-Shot, No Chunking)")

with st.sidebar:
    st.header("Model & Controls")
    # model_name = st.text_input("OpenAI model", value=DEFAULT_MODEL, help="e.g., gpt-5.1, gpt-4.1, gpt-4o-mini")
    model_name = st.selectbox(
    "Select OpenAI model",
    ["gpt-4.1-mini", "gpt-4.1-turbo", "gpt-4.1", "gpt-5.1"],
    index=0
    )
    uploaded_file = st.file_uploader("Upload contract PDF (scanned or digital)", type=["pdf"])
    regen = st.button("Extract / Rebuild context")
    clear_all = st.button("Clear session")

if clear_all:
    for k in list(st.session_state.keys()):
        del st.session_state[k]
    st.experimental_rerun()

# Extract & store
if uploaded_file and (regen or not st.session_state.get("graph_built", False)):
    try:
        with st.spinner("Extracting text (pdfplumber preferred; OCR fallback)..."):
            client = build_openai_client()
            pdf_bytes = uploaded_file.read()
            pages, full_text = extract_text_prefer_pdfplumber(pdf_bytes, dpi=200)
            st.session_state["pages"] = pages
            st.session_state["raw_text"] = full_text  # includes --- PAGE N --- markers
            st.session_state["graph_built"] = True
        st.success("PDF processed.")
    except Exception as e:
        st.exception(e)
        st.error("Processing failed.")

# Layout
col1, col2 = st.columns([2, 1])

# Chat Q&A (one-shot on full text)
with col1:
    st.subheader("Chat with the Contract (One-Shot)")
    if not st.session_state.get("graph_built", False):
        st.info("Upload & Extract a PDF to enable chat.")
    else:
        if "conversation" not in st.session_state:
            st.session_state["conversation"] = []
        for turn in st.session_state["conversation"]:
            st.chat_message(turn["role"]).write(turn["text"])

        user_q = st.chat_input("Ask about the contract (e.g., 'What is the payment clause?')")
        if user_q:
            st.session_state["conversation"].append({"role": "user", "text": user_q})
            client = build_openai_client()
            uprompt = QA_USER_TMPL.format(contract_text=st.session_state["raw_text"], question=user_q)
            try:
                resp = client.chat.completions.create(
                    model=model_name,
                    temperature=0,
                    max_tokens=900,
                    messages=[
                        {"role": "system", "content": QA_SYSTEM},
                        {"role": "user", "content": uprompt},
                    ],
                )
                ans = (resp.choices[0].message.content or "").strip()
            except Exception as e:
                ans = f"Error: {e}"
            st.session_state["conversation"].append({"role": "assistant", "text": ans})
            st.chat_message("assistant").write(ans)

# CAD / Compliance / Conflicts
with col2:
    st.subheader("CAD Generator")
    if not st.session_state.get("graph_built", False):
        st.info("Upload & Extract a PDF to enable CAD generation.")
    else:
        if st.button("Generate CAD (JSON + DOCX + PDF)"):
            try:
                client = build_openai_client()
                full_text = st.session_state["raw_text"]
                pages = st.session_state["pages"]
                with st.spinner("Generating one-shot CAD via ChatGPT..."):
                    _, json_path, docx_path, pdf_path = node_generate_cad_json_docx_pdf_openai(
                        client, model_name, full_text, pages, output_basename="generated_CAD"
                    )
                st.session_state["cad_json"] = json_path
                st.session_state["cad_docx"] = docx_path
                st.session_state["cad_pdf"] = pdf_path
                st.success("CAD generated.")
            except Exception as e:
                st.exception(e)
                st.error("CAD generation failed.")

        if st.session_state.get("cad_json"):
            try:
                with open(st.session_state["cad_json"], "rb") as f:
                    st.download_button("Download CAD JSON", f.read(), file_name=Path(st.session_state['cad_json']).name)
            except Exception as e:
                st.error(f"Failed to open CAD JSON: {e}")
        if st.session_state.get("cad_docx"):
            try:
                with open(st.session_state["cad_docx"], "rb") as f:
                    st.download_button("Download CAD DOCX", f.read(), file_name=Path(st.session_state['cad_docx']).name)
            except Exception as e:
                st.error(f"Failed to open CAD DOCX: {e}")
        if st.session_state.get("cad_pdf"):
            try:
                with open(st.session_state["cad_pdf"], "rb") as f:
                    st.download_button("Download CAD PDF", f.read(), file_name=Path(st.session_state['cad_pdf']).name)
            except Exception as e:
                st.error(f"Failed to open CAD PDF: {e}")

    st.markdown("---")
    st.subheader("📑 Compliance Check (One-Shot per rule)")
    if not st.session_state.get("graph_built", False):
        st.info("Upload a contract first.")
    else:
        default_rules = "Does the contract specify payment terms?\nIs there a termination clause?\nWhat is the governing law?"
        rules_input = st.text_area("Enter compliance rules (one per line)", value=default_rules)
        if st.button("Run Compliance Check"):
            try:
                rules = [r.strip() for r in rules_input.splitlines() if r.strip()]
                client = build_openai_client()
                full_text = st.session_state["raw_text"]
                with st.spinner("Checking compliance..."):
                    results = compliance_check_json_openai(client, model_name, full_text, rules)
                st.subheader("Compliance Report")
                df_rows = []
                for r in results:
                    st.markdown(
                        f"**Rule:** {r.get('rule')}  \n"
                        f"**Present:** {r.get('present')}  \n"
                        f"**Summary:** {r.get('summary')}  \n"
                        f"**Quote:** {r.get('quote')}  \n"
                        f"**Sources:** {r.get('sources')}  \n"
                        f"**Confidence:** {r.get('confidence')}"
                    )
                    df_rows.append([
                        r.get("rule"), r.get("present"), r.get("summary"),
                        r.get("quote"), ", ".join(r.get("sources", [])), r.get("confidence")
                    ])
                if df_rows:
                    st.dataframe(pd.DataFrame(df_rows, columns=["Rule","Present","Summary","Quote","Sources","Confidence"]))
            except Exception as e:
                st.exception(e)
                st.error("Compliance check failed.")

    st.markdown("---")
    st.subheader("⚠️ Conflict Detection")
    if not st.session_state.get("graph_built", False):
        st.info("Upload a contract first.")
    else:
        if st.button("Run Conflict Detection"):
            try:
                pages = st.session_state.get("pages")
                full_text = st.session_state.get("raw_text")
                with st.spinner("Detecting conflicts..."):
                    conflicts = run_conflict_detection(pages, full_text)
                st.session_state["conflicts"] = conflicts
                if not conflicts:
                    st.success("No conflicts detected by automated checks.")
                else:
                    st.warning(f"{len(conflicts)} potential conflicts detected.")
                    for idx, c in enumerate(conflicts, start=1):
                        st.markdown(
                            f"### {idx}. Category: **{c.get('category')}**  \n"
                            f"**Type:** {c['type']}  \n**Severity:** {c['severity']}  \n"
                            f"**Message:** {c['message']}"
                        )
                        st.markdown(f"**Suggested quick fix:** {c.get('resolution_hint')}")
                        st.json(c.get("evidence"))
            except Exception as e:
                st.exception(e)
                st.error("Conflict detection failed.")

    st.markdown("---")
    st.subheader("Raw text & pages preview")
    if st.session_state.get("raw_text"):
        if st.checkbox("Show extracted text (first 15000 chars)"):
            st.text_area("Extracted Text", value=st.session_state["raw_text"][:15000], height=400)
    else:
        st.info("No extracted text yet. Upload a PDF and click Extract.")

st.caption("ChatGPT One-Shot CAD • Large-context model • JSON-enforced responses")
# End of file
