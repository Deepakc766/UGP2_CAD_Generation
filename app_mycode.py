#embedding + retrieval in action: query → vector → similarity search → relevant chunk → LLM answer.
#we are doing embedding through sentence tranformer model and searching and through Google flt 5 laarge moel 
import os
import io
import tempfile
from pathlib import Path
from typing import List, Dict, Any
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r"/opt/homebrew/bin/tesseract"
import shutil
import pytesseract

tesseract_path = shutil.which("tesseract")
if tesseract_path:
    pytesseract.pytesseract.tesseract_cmd = tesseract_path
else:
    raise RuntimeError("Tesseract not found. Please install with `brew install tesseract`")
from typing import List
import streamlit as st
from pdf2image import convert_from_bytes
import pytesseract
from PIL import Image
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings.huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores.chroma import Chroma
from langchain.chains import ConversationalRetrievalChain
from langchain_huggingface import HuggingFacePipeline
from langchain.memory import ConversationBufferMemory
from langchain.llms import HuggingFaceHub 
from langchain.schema import Document
from langchain.text_splitter import TokenTextSplitter
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM

from docx import Document as DocxDocument
from docx.shared import Pt

# Embedding model (use a compact sentence-transformer for speed)
EMBEDDING_MODEL = "sentence-transformers/all-mpnet-base-v2"

# Generative model for RAG and summarization. Change based on resources:
# - light: "google/flan-t5-small" (fast, CPU-friendly)
# - medium: "google/flan-t5-large"
# - stronger: "bigscience/mt0-small" or local Llama/Falcon (if available)
GEN_MODEL = "google/flan-t5-large"   # change as needed
GEN_DEVICE = 0  # -1 for CPU, >=0 for GPU index

# Chroma persistent directory
CHROMA_PERSIST_DIR = "./chroma_store"

# Text chunking
CHUNK_SIZE = 500
CHUNK_OVERLAP = 120

def pdf_bytes_to_images(pdf_bytes: bytes, dpi=300) -> List[Image.Image]:
    """
    Convert PDF bytes to list of PIL images (one per page).
    Requires poppler (pdf2image).
    """
    images = convert_from_bytes(pdf_bytes, dpi=dpi)
    return images

def ocr_image_to_text(image: Image.Image) -> str:
    """
    Run pytesseract OCR to convert a PIL image to text.
    You can use other OCRs (EasyOCR / PaddleOCR) by swapping this function.
    """
    text = pytesseract.image_to_string(image)
    return text

def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    pages = pdf_bytes_to_images(pdf_bytes)
    texts = []
    for i, img in enumerate(pages):
        page_text = ocr_image_to_text(img)
        header = f"\n\n--- PAGE {i+1} ---\n\n"
        texts.append(header + page_text)
    return "\n".join(texts)

# LangGraph

class Node:
    def __init__(self, name: str, fn):
        self.name = name
        self.fn = fn
        self.outputs = None

    def run(self, *args, **kwargs):
        self.outputs = self.fn(*args, **kwargs)
        return self.outputs

class SimpleGraph:
    def __init__(self):
        self.nodes = []

    def add_node(self, node: Node):
        self.nodes.append(node)

    def run(self, start_node_index: int = 0, input=None, **kwargs):
        # For simplicity, run nodes sequentially passing the last output to the next
        data = None
        for idx, node in enumerate(self.nodes[start_node_index:], start=start_node_index):
            if idx == start_node_index:
                data = node.run(input, **kwargs)
            else:
                data = node.run(data)
        return data
def node_ocr_extract(pdf_bytes: bytes) -> str:
    st.info("Running OCR on uploaded PDF (this may take a while)...")
    text = extract_text_from_pdf_bytes(pdf_bytes)
    return text

def node_text_chunking(raw_text: str) -> List[Document]:
    st.info("Chunking text for embeddings...")
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP
    )
    texts = splitter.split_text(raw_text)
    docs = [Document(page_content=t) for t in texts]
    return docs

def node_embed_and_index(docs: List[Document], persist_dir: str = CHROMA_PERSIST_DIR) -> Chroma:
    st.info("Embedding text and building vector store (Chroma)...")
    embedder = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
    # initialize or load chroma
    vectordb = Chroma.from_documents(documents=docs, embedding=embedder, persist_directory=persist_dir)
    vectordb.persist()
    return vectordb

def build_generation_pipeline():
    """
    Build a HF seq2seq pipeline and wrap into LangChain's HuggingFacePipeline.
    We use transformers pipeline for text2text-generation so we can pass prompts + context.
    """
    st.info("Loading generative model (may take time)...")
    tokenizer = AutoTokenizer.from_pretrained(GEN_MODEL)
    model = AutoModelForSeq2SeqLM.from_pretrained(GEN_MODEL)
    # device_map or to(DEVICE) omitted for simplicity; HF pipelines can auto-use CUDA if available
    pipe = pipeline("text2text-generation", model=model, tokenizer=tokenizer, device=GEN_DEVICE, max_length=512)
    llm = HuggingFacePipeline(pipeline=pipe)
    return llm

def node_build_rag_chain(vectordb: Chroma):
    """
    Build a ConversationalRetrievalChain using the vectorstore and a generation LLM.
    """
    llm = build_generation_pipeline()
    # retrieval
    retriever = vectordb.as_retriever(search_type="mmr", search_kwargs={"k": 5})
    memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)
    chain = ConversationalRetrievalChain.from_llm(llm=llm, retriever=retriever, memory=memory, verbose=False)
    return chain

# def node_generate_cad(docs: List[Document], output_name: str = "generated_CAD.docx") -> str:
#     """
#     Generate a DOCX CAD from the parsed chunks (docs).
#     Approach: simple templating + selected summaries for key sections.
#     For production you might instead run a fine-tuned seq2seq LLM to produce each section.
#     """
#     st.info("Generating Contract Appreciation Document (.docx)...")
#     docx = DocxDocument()
#     style = docx.styles['Normal']
#     font = style.font
#     font.name = 'Arial'
#     font.size = Pt(11)

#     docx.add_heading("Contract Appreciation Document (Generated)", level=1)

#     # Naive approach: create salient features from first N chunks and tables
#     # In production: use LLM to produce structured sections
#     docx.add_heading("1. Salient Features of the Contract", level=2)
#     # use first few chunks as the source
#     for i, d in enumerate(docs[:6]):
#         # simplify: add first line of each chunk as bullet
#         text = d.page_content.strip().split("\n")[0:3]
#         bullet = " ".join([t.strip() for t in text if t.strip()])
#         if bullet:
#             docx.add_paragraph(bullet, style='List Bullet')

#     docx.add_heading("2. Important Submittals", level=2)
#     docx.add_paragraph("Automatically extracted submittals (please edit as needed):")
#     # find lines with keywords quickly
#     found = []
#     for d in docs:
#         if any(w in d.page_content.lower() for w in ["submittal", "submit", "documents to be submitted", "performance security"]):
#             lines = d.page_content.split("\n")
#             for l in lines[:5]:
#                 if len(l.strip()) > 10 and l.strip() not in found:
#                     docx.add_paragraph(l.strip(), style='List Bullet')
#                     found.append(l.strip())
#     # Add placeholders for other sections
#     docx.add_heading("3. Notice / Information Clauses", level=2)
#     docx.add_paragraph("Please review and edit as necessary.")

#     docx.add_heading("4. Important Clauses: EOT, Escalation, Variation, Suspension", level=2)
#     docx.add_paragraph("Please review and edit as necessary.")

#     docx.add_heading("5. Payment Clause", level=2)
#     docx.add_paragraph("Please review and edit as necessary.")

#     docx.add_heading("6. Risk Identification Matrix", level=2)
#     docx.add_paragraph("Please review and edit as necessary.")

#     docx.add_heading("7. Claims, Disputes and Arbitration", level=2)
#     docx.add_paragraph("Please review and edit as necessary.")

#     # Save docx
#     out_path = Path(tempfile.gettempdir()) / output_name
#     docx.save(out_path)
#     return str(out_path)

# # --------------------------
# # Streamlit App
# # --------------------------
# st.set_page_config(page_title="Contract CAD Chatbot (LangGraph-style)", layout="wide")
# st.title("Contract CAD Chatbot — OCR → RAG → CAD generation")

# # Sidebar: controls & status
# with st.sidebar:
#     st.header("Pipeline Controls")
#     uploaded_file = st.file_uploader("Upload scanned contract PDF", type=["pdf"])
#     if uploaded_file:
#         st.write(f"Uploaded: {uploaded_file.name}")
#     regenerate = st.button("Regenerate Index / CAD")
#     clear_index = st.button("Clear Vector DB")

#     st.markdown("### Model settings")
#     st.write(f"Embedding model: `{EMBEDDING_MODEL}`")
#     st.write(f"Generation model: `{GEN_MODEL}`")

# # In-memory handles (session state)
# if "vectordb" not in st.session_state:
#     st.session_state["vectordb"] = None
# if "rag_chain" not in st.session_state:
#     st.session_state["rag_chain"] = None
# if "raw_text" not in st.session_state:
#     st.session_state["raw_text"] = ""
# if "docs" not in st.session_state:
#     st.session_state["docs"] = []
# if "cad_path" not in st.session_state:
#     st.session_state["cad_path"] = None

# # Clear DB
# if clear_index:
#     if Path(CHROMA_PERSIST_DIR).exists():
#         import shutil
#         shutil.rmtree(CHROMA_PERSIST_DIR)
#     st.session_state["vectordb"] = None
#     st.success("Cleared vector DB.")

# # Process uploaded file (or regenerate)
# if uploaded_file and (regenerate or st.session_state["vectordb"] is None):
#     pdf_bytes = uploaded_file.read()
#     # Build graph and run
#     graph = SimpleGraph()
#     graph.add_node(Node("OCR", node_ocr_extract))
#     graph.add_node(Node("Chunking", node_text_chunking))
#     graph.add_node(Node("EmbedIndex", node_embed_and_index))
#     graph.add_node(Node("BuildRAG", node_build_rag_chain))
#     # run graph: pass pdf_bytes to OCR node; each subsequent node receives previous output
#     vectordb = vectordb = graph.run(input=pdf_bytes, start_node_index=0)
#     # node_embed_and_index returns vectordb, and node_build_rag_chain returns chain
#     # The SimpleGraph returns last node output (RAG chain) - but we want to capture intermediate
#     # For simplicity we run nodes separately and capture docs too:
#     raw_text = node_ocr_extract(pdf_bytes)
#     docs = node_text_chunking(raw_text)
#     vectordb = node_embed_and_index(docs)
#     rag_chain = node_build_rag_chain(vectordb)

#     # store in session
#     st.session_state["raw_text"] = raw_text
#     st.session_state["docs"] = docs
#     st.session_state["vectordb"] = vectordb
#     st.session_state["rag_chain"] = rag_chain
#     st.success("Indexed uploaded contract!")

# # If vector DB exists but user didn't re-upload, ensure rag_chain is available
# if st.session_state.get("vectordb") and not st.session_state.get("rag_chain"):
#     st.session_state["rag_chain"] = node_build_rag_chain(st.session_state["vectordb"])

# # Left column: Chat
# col1, col2 = st.columns([2, 1])
# with col1:
#     st.subheader("Chat with the contract (RAG)")
#     if st.session_state.get("rag_chain") is None:
#         st.info("Upload a contract PDF to begin. The app will OCR, index, and enable chat.")
#     else:
#         # Display conversation
#         if "conversation" not in st.session_state:
#             st.session_state["conversation"] = []
#         for i, turn in enumerate(st.session_state["conversation"]):
#             if turn["role"] == "user":
#                 st.chat_message("user").write(turn["text"])
#             else:
#                 st.chat_message("assistant").write(turn["text"])

#         user_input = st.chat_input("Ask about the contract (e.g., 'What is the payment clause?')")
#         if user_input:
#             st.session_state["conversation"].append({"role": "user", "text": user_input})
#             with st.spinner("Retrieving answer..."):
#                 chain = st.session_state["rag_chain"]
#                 res = chain({"question": user_input})
#                 answer = res["answer"]
#             st.session_state["conversation"].append({"role": "assistant", "text": answer})
#             st.experimental_rerun()

# with col2:
#     st.subheader("CAD Generator / Editor")
#     if st.session_state.get("docs"):
#         if st.button("Generate CAD (DOCX)"):
#             cad_path = node_generate_cad(st.session_state["docs"])
#             st.session_state["cad_path"] = cad_path
#             st.success(f"CAD generated: {cad_path}")

#         if st.session_state.get("cad_path"):
#             # allow download
#             with open(st.session_state["cad_path"], "rb") as f:
#                 st.download_button("Download CAD (.docx)", f.read(), file_name="generated_CAD.docx")

#             # Convert to PDF on request
#             if st.button("Convert CAD -> PDF (WeasyPrint)"):
#                 try:
#                     from weasyprint import HTML
#                     import mammoth  # optional: docx to HTML converter for good rendering
#                     # We'll convert docx->html (via mammoth) then html->pdf
#                     with open(st.session_state["cad_path"], "rb") as docx_file:
#                         result = mammoth.convert_to_html(docx_file)
#                         html = result.value  # the generated HTML
#                         out_pdf = Path(tempfile.gettempdir()) / "generated_CAD.pdf"
#                         HTML(string=html).write_pdf(str(out_pdf))
#                         with open(out_pdf, "rb") as pdf_f:
#                             st.download_button("Download CAD (.pdf)", pdf_f.read(), file_name="generated_CAD.pdf")
#                 except Exception as e:
#                     st.error("PDF conversion failed. Please install WeasyPrint and mammoth or use LibreOffice CLI.")
#     else:
#         st.info("No indexed contract. Upload a PDF to produce CAD.")

# # Button to show a quick preview of raw OCR text
# if st.session_state.get("raw_text"):
#     if st.checkbox("Show OCR raw text"):
#         st.text_area("OCR Text", value=st.session_state["raw_text"][:15000], height=400)

# st.markdown("---")
# st.caption("This app is a reference implementation. Swap OCR, embeddings, and generation models as desired.")


from docx import Document as DocxDocument
from docx.shared import Pt
from pathlib import Path
import tempfile
import streamlit as st
import re

from docx import Document as DocxDocument
from docx.shared import Pt
from pathlib import Path
import tempfile
import streamlit as st
import re
import pandas as pd

def _ask(rag_chain, q: str) -> str:
    try:
        out = rag_chain({"question": q})
        return (out.get("answer") or "").strip()
    except Exception as e:
        return f"[Error] {e}"

def _split_rows(raw: str, min_cols=2, max_cols=4):
    """Split LLM lines like 'A | B | C | D' into rows; fallback to bullets."""
    rows = []
    for line in raw.split("\n"):
        if not line.strip():
            continue
        parts = [p.strip(" -–•\t") for p in re.split(r"\s*\|\s*|\s+–\s+|\s+-\s+", line) if p.strip()]
        if len(parts) >= min_cols:
            rows.append(parts[:max_cols])
    return rows

def node_generate_cad(rag_chain, output_name: str = "generated_CAD.docx"):
    """
    Advanced CAD generator:
    - Uses RAG questions per section
    - Tables for: Salient, Submittals, Payment, Risk, Claims/Disputes (when possible)
    - Compliance checklist
    - Returns (docx_path, preview_dict) where preview_dict has dataframes for Streamlit preview
    """
    st.info("Generating Contract Appreciation Document (.docx)...")
    docx = DocxDocument()
    style = docx.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    docx.add_heading("Contract Appreciation Document (Generated)", level=1)

    previews = {}  # section_name -> pandas.DataFrame for UI preview


    title = "1. Salient Features of the Contract"
    docx.add_heading(title, level=2)
    q1 = ("Extract key contract details as Key:Value pairs. Include: "
          "Project Name, Location, Contract Ref No., Client/Employer, Contractor, "
          "Stakeholders, Contract Type, Contract Value, Price Escalation (Yes/No & basis), "
          "Award Date, Commencement Date, Completion Date, Defect Liability Period, Scope Overview. "
          "Output each on new line as 'Key | Value'.")
    a1 = _ask(rag_chain, q1)
    rows1 = _split_rows(a1, min_cols=2, max_cols=2)
    if rows1:
        table = docx.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Key"
        table.rows[0].cells[1].text = "Value"
        for k, v in rows1:
            r = table.add_row().cells
            r[0].text, r[1].text = k, v
        previews[title] = pd.DataFrame(rows1, columns=["Key", "Value"])
    else:
        docx.add_paragraph(a1 or "Not found in contract.")


    title = "2. Important Submittals"
    docx.add_heading(title, level=2)
    q2 = ("List required submittals grouped by stage. Output each row as "
          "'Stage (Before/During/After) | Document | Due/Timeline'.")
    a2 = _ask(rag_chain, q2)
    rows2 = _split_rows(a2, min_cols=3, max_cols=3)
    if rows2:
        table = docx.add_table(rows=1, cols=3)
        table.rows[0].cells[0].text = "Stage"
        table.rows[0].cells[1].text = "Document"
        table.rows[0].cells[2].text = "Due / Timeline"
        for s, d, t in rows2:
            r = table.add_row().cells
            r[0].text, r[1].text, r[2].text = s, d, t
        previews[title] = pd.DataFrame(rows2, columns=["Stage", "Document", "Due / Timeline"])
    else:
        docx.add_paragraph(a2 or "Not found in contract.")


    title = "3. Notice / Information Clauses"
    docx.add_heading(title, level=2)
    q3 = ("Summarize notice provisions. For each notice, output: "
          "'Event/Trigger | Party to Notify | Timeline | Method/Channel'.")
    a3 = _ask(rag_chain, q3)
    rows3 = _split_rows(a3, min_cols=3, max_cols=4)
    if rows3:
        # pad missing 4th col
        rows3 = [r + [""]*(4-len(r)) for r in rows3]
        table = docx.add_table(rows=1, cols=4)
        hdr = ["Event/Trigger", "Party to Notify", "Timeline", "Method/Channel"]
        for i, h in enumerate(hdr): table.rows[0].cells[i].text = h
        for e, p, t, m in rows3:
            r = table.add_row().cells
            r[0].text, r[1].text, r[2].text, r[3].text = e, p, t, m
        previews[title] = pd.DataFrame(rows3, columns=hdr)
    else:
        docx.add_paragraph(a3 or "Not found in contract.")


    title = "4. Important Clauses Pertaining to Project Progress (EOT, Escalation, Variation, Suspension)"
    docx.add_heading(title, level=2)
    q4 = ("Identify clauses related to EOT, price escalation, variations, suspension. "
          "Summarize as clear bullet points.")
    a4 = _ask(rag_chain, q4)
    if a4:
        for line in a4.split("\n"):
            if line.strip():
                docx.add_paragraph(line.strip(" -–•"), style="List Bullet")
    else:
        docx.add_paragraph("Not found in contract.")


    title = "5. Payment Clause"
    docx.add_heading(title, level=2)
    q5 = ("Summarize payment terms in rows: 'Payment Type | Frequency | Conditions'. "
          "Include advances (mobilization/secured), interim payments, final settlement, "
          "retention money (rate & release), escalation payments, taxes/levies, penalties/incentives.")
    a5 = _ask(rag_chain, q5)
    rows5 = _split_rows(a5, min_cols=3, max_cols=3)
    if rows5:
        table = docx.add_table(rows=1, cols=3)
        table.rows[0].cells[0].text = "Payment Type"
        table.rows[0].cells[1].text = "Frequency"
        table.rows[0].cells[2].text = "Conditions"
        for ptype, freq, cond in rows5:
            r = table.add_row().cells
            r[0].text, r[1].text, r[2].text = ptype, freq, cond
        previews[title] = pd.DataFrame(rows5, columns=["Payment Type", "Frequency", "Conditions"])
    else:
        docx.add_paragraph(a5 or "Not found in contract.")


    title = "6. Risk Matrix and Risk Allocation"
    docx.add_heading(title, level=2)
    q6 = ("Identify risks categorized across pre-construction, approval, design, commercial, execution. "
          "For each risk, output 'Risk | Severity | Responsibility'.")
    a6 = _ask(rag_chain, q6)
    rows6 = _split_rows(a6, min_cols=3, max_cols=3)
    if rows6:
        table = docx.add_table(rows=1, cols=3)
        table.rows[0].cells[0].text = "Risk"
        table.rows[0].cells[1].text = "Severity"
        table.rows[0].cells[2].text = "Responsibility"
        for risk, sev, resp in rows6:
            r = table.add_row().cells
            r[0].text, r[1].text, r[2].text = risk, sev, resp
        previews[title] = pd.DataFrame(rows6, columns=["Risk", "Severity", "Responsibility"])
    else:
        docx.add_paragraph(a6 or "Not found in contract.")


    title = "7. Claims, Disputes and Arbitration"
    docx.add_heading(title, level=2)
    q7 = ("Summarize claims, disputes, arbitration in rows: "
          "'Clause/Topic | Process/Forum | Jurisdiction | Timeline/Notes'. "
          "Include claim procedures, dispute steps, arbitration rules, venue, governing courts, and timelines.")
    a7 = _ask(rag_chain, q7)
    rows7 = _split_rows(a7, min_cols=2, max_cols=4)
    if rows7:
        rows7 = [r + [""]*(4-len(r)) for r in rows7]
        table = docx.add_table(rows=1, cols=4)
        hdr = ["Clause/Topic", "Process/Forum", "Jurisdiction", "Timeline/Notes"]
        for i, h in enumerate(hdr): table.rows[0].cells[i].text = h
        for c, p, j, t in rows7:
            r = table.add_row().cells
            r[0].text, r[1].text, r[2].text, r[3].text = c, p, j, t
        previews[title] = pd.DataFrame(rows7, columns=hdr)
    else:
        docx.add_paragraph(a7 or "Not found in contract.")
    # 8) Expanded Compliance Checklist

    title = "8. Compliance Checklist (Auto)"
    docx.add_heading(title, level=2)
    checks = {
        # General Clauses
        "Arbitration clause present?": "Is there an arbitration clause? Answer Yes/No and cite the clause number if any.",
        "Governing law specified?": "Does the contract define governing law (e.g., Indian Law, English Law)? Answer Yes/No with reference.",
        "Jurisdiction / venue of arbitration defined?": "Is the jurisdiction or arbitration venue defined? Answer Yes/No with details.",
        "Force Majeure defined?": "Is force majeure defined including examples? Answer Yes/No with reference.",

        # Financial Security & Guarantees
        "Performance security / bank guarantee required?": "Is performance security required? State type, amount and validity. Answer Yes/No.",
        "Advance payment guarantee required?": "Does the contract require an advance payment guarantee? Answer Yes/No with details.",
        "Retention money defined?": "Does the contract define retention money (rate and release conditions)? Answer Yes/No with details.",
        "Liquidated damages clause present?": "Does the contract define liquidated damages (LDs)? State rate, maximum cap. Answer Yes/No.",
        "Bonus / incentives for early completion?": "Are there bonuses/incentives for early completion? Answer Yes/No with details.",

        # Payments & Commercial
        "Payment frequency clearly defined?": "Is interim payment frequency specified (e.g., monthly)? Answer Yes/No with reference.",
        "Escalation/payment adjustment terms defined?": "Does the contract specify price escalation/adjustment terms? Answer Yes/No.",
        "Taxes and levies responsibility defined?": "Does the contract define responsibility for taxes/levies (client or contractor)? Answer Yes/No.",

        # Insurance & Risk
        "Contractor’s All Risk (CAR) insurance required?": "Does the contract require CAR insurance? Answer Yes/No with details.",
        "Third-party liability insurance defined?": "Is third-party liability insurance required? Answer Yes/No with details.",
        "Worker’s compensation insurance required?": "Is worker’s compensation or similar insurance required? Answer Yes/No with reference.",

        # Execution & Technical
        "Defect liability period (DLP) specified?": "Is a defect liability period specified? Answer Yes/No with duration.",
        "Testing & commissioning obligations defined?": "Are testing and commissioning requirements defined? Answer Yes/No.",
        "Quality submittals / shop drawings approval process defined?": "Does the contract specify submittal and approval of drawings/documents? Answer Yes/No.",
        "Change order / variation process specified?": "Is the variation/change order process clearly defined? Answer Yes/No.",
        "Suspension & termination clauses defined?": "Are suspension/termination clauses defined? Answer Yes/No with details."
    }
    for label, q in checks.items():
        ans = _ask(rag_chain, q)
        para = docx.add_paragraph()
        run = para.add_run(f"• {label} ")
        run.bold = True
        docx.add_paragraph(ans or "No information found.")


    # Save DOCX
    out_path = Path(tempfile.gettempdir()) / output_name
    docx.save(out_path)

    return str(out_path), previews


# Streamlit App

st.set_page_config(page_title="Contract CAD Chatbot (LangGraph-style)", layout="wide")
st.title("Contract CAD Chatbot — OCR → RAG → CAD generation")

# Sidebar: controls & status
with st.sidebar:
    st.header("Pipeline Controls")
    uploaded_file = st.file_uploader("Upload scanned contract PDF", type=["pdf"])
    if uploaded_file:
        st.write(f"Uploaded: {uploaded_file.name}")
    regenerate = st.button("Regenerate Index / CAD")
    clear_index = st.button("Clear Vector DB")

    st.markdown("### Model settings")
    st.write(f"Embedding model: `{EMBEDDING_MODEL}`")
    st.write(f"Generation model: `{GEN_MODEL}`")

# In-memory handles (session state)
if "vectordb" not in st.session_state:
    st.session_state["vectordb"] = None
if "rag_chain" not in st.session_state:
    st.session_state["rag_chain"] = None
if "raw_text" not in st.session_state:
    st.session_state["raw_text"] = ""
if "docs" not in st.session_state:
    st.session_state["docs"] = []
if "cad_path" not in st.session_state:
    st.session_state["cad_path"] = None

# Clear DB
if clear_index:
    if Path(CHROMA_PERSIST_DIR).exists():
        import shutil
        shutil.rmtree(CHROMA_PERSIST_DIR)
    st.session_state["vectordb"] = None
    st.success("Cleared vector DB.")

# Process uploaded file (or regenerate)
if uploaded_file and (regenerate or st.session_state["vectordb"] is None):
    pdf_bytes = uploaded_file.read()
    # Build graph and run
    graph = SimpleGraph()
    graph.add_node(Node("OCR", node_ocr_extract))
    graph.add_node(Node("Chunking", node_text_chunking))
    graph.add_node(Node("EmbedIndex", node_embed_and_index))
    graph.add_node(Node("BuildRAG", node_build_rag_chain))
    # run graph: pass pdf_bytes to OCR node; each subsequent node receives previous output
    vectordb = vectordb = graph.run(input=pdf_bytes, start_node_index=0)
    # node_embed_and_index returns vectordb, and node_build_rag_chain returns chain
    # The SimpleGraph returns last node output (RAG chain) - but we want to capture intermediate
    # For simplicity we run nodes separately and capture docs too:
    raw_text = node_ocr_extract(pdf_bytes)
    docs = node_text_chunking(raw_text)
    vectordb = node_embed_and_index(docs)
    rag_chain = node_build_rag_chain(vectordb)

    # store in session
    st.session_state["raw_text"] = raw_text
    st.session_state["docs"] = docs
    st.session_state["vectordb"] = vectordb
    st.session_state["rag_chain"] = rag_chain
    st.success("Indexed uploaded contract!")

# If vector DB exists but user didn't re-upload, ensure rag_chain is available
if st.session_state.get("vectordb") and not st.session_state.get("rag_chain"):
    st.session_state["rag_chain"] = node_build_rag_chain(st.session_state["vectordb"])

# Left column: Chat
col1, col2 = st.columns([2, 1])

with col1:
    st.subheader("Chat with the contract (RAG)")

    if st.session_state.get("rag_chain") is None:
        st.info("Upload a contract PDF to begin.")
    else:
        # Ensure conversation history exists
        if "conversation" not in st.session_state:
            st.session_state["conversation"] = []

        # Render chat history
        for turn in st.session_state["conversation"]:
            if turn["role"] == "user":
                st.chat_message("user").write(turn["text"])
            else:
                st.chat_message("assistant").write(turn["text"])

        # Input box at the bottom
        user_input = st.chat_input("Ask about the contract (e.g., 'What is the payment clause?')")

        if user_input:
            # Append user message
            st.session_state["conversation"].append({"role": "user", "text": user_input})

            # Generate assistant response
            with st.spinner("Retrieving answer..."):
                chain = st.session_state["rag_chain"]
                res = chain({"question": user_input})
                answer = res["answer"]

            # Append assistant message
            st.session_state["conversation"].append({"role": "assistant", "text": answer})

            # Immediately render the assistant’s reply without rerun
            st.chat_message("assistant").write(answer)


# with col2:
#     st.subheader("CAD Generator / Editor")
#     if st.session_state.get("rag_chain"):
#         if st.button("Generate CAD (DOCX)"):
#             cad_path = node_generate_cad(st.session_state["rag_chain"])
#             st.session_state["cad_path"] = cad_path
#             st.success(f"CAD generated: {cad_path}")

#         if st.session_state.get("cad_path"):
#             # allow download
#             with open(st.session_state["cad_path"], "rb") as f:
#                 st.download_button("Download CAD (.docx)", f.read(), file_name="generated_CAD.docx")

#             # Convert to PDF on request
#             if st.button("Convert CAD -> PDF (WeasyPrint)"):
#                 try:
#                     from weasyprint import HTML
#                     import mammoth  # optional: docx to HTML converter for good rendering
#                     # We'll convert docx->html (via mammoth) then html->pdf
#                     with open(st.session_state["cad_path"], "rb") as docx_file:
#                         result = mammoth.convert_to_html(docx_file)
#                         html = result.value  # the generated HTML
#                         out_pdf = Path(tempfile.gettempdir()) / "generated_CAD.pdf"
#                         HTML(string=html).write_pdf(str(out_pdf))
#                         with open(out_pdf, "rb") as pdf_f:
#                             st.download_button("Download CAD (.pdf)", pdf_f.read(), file_name="generated_CAD.pdf")
#                 except Exception as e:
#                     st.error("PDF conversion failed. Please install WeasyPrint and mammoth or use LibreOffice CLI.")
#     else:
#         st.info("No contract. Upload a PDF to produce CAD.")
with col2:
    st.subheader("CAD Generator / Editor")
    if st.session_state.get("rag_chain"):
        # Generate CAD
        if st.button("Generate CAD (DOCX)"):
            try:
                with st.spinner("Generating CAD — this may take a minute..."):
                    # node_generate_cad returns (docx_path, previews)
                    result = node_generate_cad(st.session_state["rag_chain"])
                    # Safety: unpack if tuple; if older version returned string, handle that too
                    if isinstance(result, tuple) and len(result) >= 1:
                        cad_path, previews = result[0], result[1] if len(result) > 1 else None
                    else:
                        cad_path, previews = result, None

                    # Normalize to str and ensure file exists
                    cad_path = str(cad_path)
                    if Path(cad_path).exists():
                        st.session_state["cad_path"] = cad_path
                        st.session_state["cad_previews"] = previews
                        st.success(f"CAD generated: {cad_path}")
                    else:
                        st.error("CAD generation finished but output file not found.")
            except Exception as e:
                st.exception(e)
                st.error("CAD generation failed. Check logs for details.")

        # If present, allow download and conversion
        if st.session_state.get("cad_path"):
            cad_path = st.session_state["cad_path"]
            try:
                with open(cad_path, "rb") as f:
                    st.download_button("Download CAD (.docx)", f.read(), file_name=Path(cad_path).name)
            except Exception as e:
                st.error(f"Failed to open generated CAD file: {e}")

            # Convert to PDF on request (WeasyPrint/mammoth)
            if st.button("Convert CAD -> PDF (WeasyPrint)"):
                try:
                    from weasyprint import HTML
                    import mammoth
                    with open(cad_path, "rb") as docx_file:
                        result = mammoth.convert_to_html(docx_file)
                        html = result.value
                        out_pdf = Path(tempfile.gettempdir()) / "generated_CAD.pdf"
                        HTML(string=html).write_pdf(str(out_pdf))
                        with open(out_pdf, "rb") as pdf_f:
                            st.download_button("Download CAD (.pdf)", pdf_f.read(), file_name="generated_CAD.pdf")
                except Exception as e:
                    st.error("PDF conversion failed. Please install WeasyPrint and mammoth, or use LibreOffice CLI.")
    else:
        st.info("No contract. Upload a PDF to produce CAD.")

# Button to show a quick preview of raw OCR text
if st.session_state.get("raw_text"):
    if st.checkbox("Show OCR raw text"):
        st.text_area("OCR Text", value=st.session_state["raw_text"][:15000], height=400)

st.markdown("---")
st.caption("This app is fir CAD Generation")
