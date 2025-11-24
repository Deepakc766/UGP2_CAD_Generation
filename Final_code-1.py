# app_streamlit_cad_chatgpt.py
import os
import io
import json
import tempfile
from pathlib import Path
from typing import List, Dict, Any, Optional
import logging
import re
import streamlit as st
import pdfplumber
from reportlab.lib.pagesizes import A4
from dateutil import parser as dateparser

from pdf2image import convert_from_bytes
from PIL import Image
import pytesseract
import shutil
from docx import Document as DocxDocument
from docx.shared import Pt
import pandas as pd

# --- ReportLab Platypus (table PDF) ---
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# ============ OpenAI ChatGPT client ============
# Requires: pip install openai==1.*
from openai import OpenAI

# -------------------------
# Config
# -------------------------
# Prefer your large-context model in prod; keeping options in UI.
DEFAULT_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4.1")
MAX_OUTPUT_TOKENS = int(os.environ.get("OPENAI_MAX_OUTPUT_TOKENS", "12000"))
TEMPERATURE = 0.0

# Ensure tesseract path if present
tesseract_path = shutil.which("tesseract")
if tesseract_path:
    pytesseract.pytesseract.tesseract_cmd = tesseract_path

# -------------------------
# OpenAI helpers
# -------------------------
@st.cache_resource(show_spinner=False)
def build_openai_client():
    api_key = st.secrets.get("OPENAI_API_KEY", os.environ.get("OPENAI_API_KEY"))
    if not api_key:
        st.error("Missing OpenAI API key. Set in Streamlit secrets as OPENAI_API_KEY or env var.")
        st.stop()
    client = OpenAI(api_key=api_key)
    return client

def chat_json(
    client: OpenAI,
    model: str,
    system_prompt: str,
    user_prompt: str,
    max_tokens: int = MAX_OUTPUT_TOKENS,
    temperature: float = TEMPERATURE,
) -> str:
    """
    Calls Chat Completions with response_format=json_object to force strict JSON.
    Returns raw JSON string from the assistant.
    """
    resp = client.chat.completions.create(
        model=model,
        temperature=temperature,
        max_tokens=max_tokens,
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": system_prompt.strip()},
            {"role": "user", "content": user_prompt.strip()},
        ],
    )
    return (resp.choices[0].message.content or "").strip()

def safe_json_loads(txt: str) -> Optional[dict]:
    if not txt:
        return None
    try:
        return json.loads(txt)
    except Exception:
        # last-resort: try to extract {...}
        s = txt.find("{"); e = txt.rfind("}")
        if s != -1 and e != -1:
            cand = txt[s:e+1]
            cand = re.sub(r",\s*}", "}", cand)
            cand = re.sub(r",\s*]", "]", cand)
            try:
                return json.loads(cand)
            except Exception:
                return None
        return None

# -------------------------
# PDF → Text extraction
# -------------------------
def pdf_bytes_to_images(pdf_bytes: bytes, dpi: int = 200) -> List[Image.Image]:
    return convert_from_bytes(pdf_bytes, dpi=dpi)

def ocr_image_to_text(image: Image.Image) -> str:
    return pytesseract.image_to_string(image)

def extract_text_from_pdf_bytes(pdf_bytes: bytes, dpi: int = 200) -> str:
    pages = pdf_bytes_to_images(pdf_bytes, dpi=dpi)
    texts = []
    for i, img in enumerate(pages):
        page_text = ocr_image_to_text(img)
        header = f"\n\n--- PAGE {i+1} ---\n\n"
        texts.append(header + page_text)
    return "\n".join(texts)

def extract_text_prefer_pdfplumber(pdf_bytes: bytes, dpi: int = 200):
    """
    Use pdfplumber to extract selectable text per page; fallback to OCR for pages with no text.
    Returns: (pages_list, full_text)
      pages_list: [{'page': n, 'text': str, 'is_ocr': bool}, ...]
      full_text: string with --- PAGE n --- markers
    """
    pages = []
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                try:
                    txt = page.extract_text() or ""
                except Exception:
                    txt = ""
                if txt.strip():
                    pages.append({"page": i, "text": txt, "is_ocr": False})
                else:
                    try:
                        pil_img = page.to_image(resolution=dpi).original
                        ocr_txt = pytesseract.image_to_string(pil_img)
                    except Exception:
                        ocr_txt = ""
                    pages.append({"page": i, "text": ocr_txt, "is_ocr": True})
    except Exception as e:
        logging.warning("pdfplumber failed: %s. Falling back to OCR for whole PDF.", e)
        ocr_full = extract_text_from_pdf_bytes(pdf_bytes, dpi=dpi)
        page_texts = [p for p in ocr_full.split("\f")] if "\f" in ocr_full else [ocr_full]
        for idx, p_text in enumerate(page_texts, start=1):
            pages.append({"page": idx, "text": p_text, "is_ocr": True})

    full = []
    for p in pages:
        full.append(f"\n\n--- PAGE {p['page']} ---\n{p['text']}")
    return pages, "\n".join(full)

# -------------------------
# Prompts (ONE-SHOT, full contract)
# -------------------------
CAD_SYSTEM = """
You are a senior contracts analyst.
Output STRICT JSON ONLY (no markdown, no comments, no prose outside the JSON).
Never hallucinate; if not explicitly supported by contract text, use null.
Prefer concise phrasing. Normalize dates to YYYY-MM-DD when explicit.
For every factual extraction, attach a sources array with page/ clause markers (e.g., "page 7", "clause 37.4").
"""

CAD_STYLE_GUIDE = """
Sections and order (must match exactly):
1. Salient Features of the Contract
2. List of Important Submittals
3. Notice / Information Clauses
4. Important clauses pertaining to project progress - EOT, Escalation, Variation, Suspension, etc.
5. Payment Clause
6. Identify Risk, Risk Matrix, and Risk Allocation
7. Claim, Dispute, and Arbitration clause
"""

CAD_USER_TMPL = """Use ONLY the CONTRACT TEXT to produce ONE JSON object exactly matching the schema.
Respond with strict JSON (start with '{{' and end with '}}'). No extra text.

CONTRACT TEXT (page markers included):
{contract_text}

STYLE GUIDE:
{style_guide}

SCHEMA:
{{
  "salient_features": [
    {{"description":"string","clause_or_ref":"string|null","details":"string|null","sources":["page X"]}}
  ],
  "important_submittals": [
    {{"stage":"before|during|after|null","document":"string","due":"string|null","notes":"string|null","clause_or_ref":"string|null","sources":["page X"]}}
  ],
  "notice_information_clauses": [
    {{"description":"string","clause_or_ref":"string|null","details":"string","sources":["page X"]}}
  ],
  "project_progress_clauses": [
    {{"topic":"string","clause_or_ref":"string|null","summary":"string","sources":["page X"]}}
  ],
  "payment": {{
    "type_of_contract":"string|null",
    "mode_of_payment":"string|null",
    "measurement_and_evaluation":"string|null",
    "billing_timing":"string|null",
    "other_terms":"string|null",
    "contract_sum": {{"amount": number|null, "currency":"string|null","sources":["page X"]}},
    "mobilization_advance": {{"applicable": true|false, "amount_pct": number|null, "interest":"string|null", "security":"string|null", "release_conditions":"string|null","sources":["page X"]}},
    "retention": {{"percent": number|null,"release_condition":"string|null","sources":["page X"]}},
    "interim_payments": [{{"frequency":"string|null","certifier":"string|null","payment_days_after_certification": number|null,"sources":["page X"]}}],
    "price_escalation": {{"applicable": true|false, "basis":"string|null", "formula_or_reference":"string|null","sources":["page X"]}},
    "final_payment": {{"required_docs":["string"],"release_days_after_submission": number|null,"sources":["page X"]}},
    "sources":["page X"]
  }},
  "risk_matrix": {{
    "severity_buckets":[{{"label":"string","criteria":"string"}}],
    "probability_buckets":[{{"label":"string","criteria":"string"}}],
    "matrix_legend": "string|null"
  }},
  "risks_and_allocation": [
    {{"id":"string","category":"string","risk_element":"string","probability":"Most Likely|Likely|Occasional|Unlikely|Remote|null","severity":"Very Severe|Major|Significant|Minor|Insignificant|null","rating":"Low|Medium|High|Critical|null","owner":"Employer|Contractor|Shared|null","mitigation":"string|null","sources":["page X"]}}
  ],
  "claims_disputes_arbitration": {{
    "arbitration": {{"applicable": true|false,"forum_or_rules":"string|null","notes":"string|null","sources":["page X"]}},
    "court_jurisdiction":"string|null",
    "claims_summary":[{{"topic":"string","payable_or_not":"Payable|Not Payable|Unclear","notes":"string","sources":["page X"]}}],
    "dispute_areas":[{{"description":"string","sources":["page X"]}}],
    "excusable_delays":[{{"item":"string","sources":["page X"]}}],
    "delay_compensation_clause":"string|null",
    "sources":["page X"]
  }}
}}

CRITICAL RULES:
- Fill fields ONLY if supported by the contract text. If unknown, use null.
- Every extracted fact MUST include "sources" with page and, if present, clause references.
- Keep text concise and specific (no generic boilerplate). Use lists instead of long paragraphs where helpful.
- Do NOT invent clause numbers; if absent, keep clause_or_ref null and rely on page numbers in "sources".
"""

# COMPLIANCE_SYSTEM = """You are a contracts compliance checker. Always return STRICT JSON (no markdown)."""

# COMPLIANCE_USER_TMPL = """Use ONLY the CONTRACT TEXT to evaluate the rule. Return ONE strict JSON object.

# CONTRACT TEXT:
# {contract_text}

# RULE:
# {rule}

# Return JSON in one of these shapes:
# {{"rule":"{rule}","present": true,"summary":"short","quote":"<=200 chars verbatim","sources":["page X","clause Y"],"confidence":0.9}}
# or
# {{"rule":"{rule}","present": false,"summary": null,"quote": null,"sources": [],"confidence": 0.99}}
# """


COMPLIANCE_SYSTEM = """
You are a senior contracts compliance auditor.
You MUST return STRICT JSON ONLY (no markdown, no prose outside JSON).
Work ONLY from the contract text provided. If unknown, mark as unknown and give confidence accordingly.
Always cite pages (and clause numbers when visible) in a sources array like ["page 7","clause 14.3"].
"""

# COMPLIANCE_USER_TMPL = """You are given the FULL CONTRACT TEXT and a LIST OF RULES.
# Audit compliance for ALL rules in one pass (to ensure consistent judgments).
# Return STRICT JSON that matches the schema below.

# CONTRACT TEXT (page markers included):
# {contract_text}

# RULES (one per line):
# {rules_text}

# SCHEMA:
# {{
#   "results": [
#     {{
#       "rule": "string",                      // the exact rule text from input
#       "status": "Compliant|Non-Compliant|Ambiguous|Not Found",
#       "summary": "short reasoning (<=300 chars)",
#       "evidence": [                          // short verbatim quotes (<=200 chars each)
#         {{"quote":"string","sources":["page X","clause Y"]}}
#       ],
#       "confidence": 0.0                      // 0-1
#     }}
#   ],
#   "overall_compliance_score": 0,             // 0-100 integer (%), weighted by severity if inferable
#   "notes": "string|null"                     // any general remarks, null if none
# }}
# CRITICAL:
# - Use null or empty arrays when appropriate; never invent facts.
# - Prefer page + clause when available; otherwise page only.
# - Keep quotes short and precise (<=200 chars).
# """

COMPLIANCE_USER_TMPL = (
    "Use ONLY the CONTRACT TEXT to evaluate the rule. Return ONE strict JSON object.\n\n"
    "CONTRACT TEXT:\n{contract_text}\n\n"
    "RULE:\n{rule}\n\n"
    "Return JSON in one of these shapes:\n"
    '{{"rule":"{rule}","present": true,"summary":"short","quote":"<=200 chars verbatim","sources":["page X","clause Y"],"confidence":0.9}}\n'
    "or\n"
    '{{"rule":"{rule}","present": false,"summary": null,"quote": null,"sources": [],"confidence": 0.99}}'
)


QA_SYSTEM = """You are a contracts Q&A assistant. Answer ONLY from the contract text."""

QA_USER_TMPL = """CONTRACT TEXT:
{contract_text}

QUESTION:
{question}

If unknown, reply exactly: "Not found in contract."
"""

# -------------------------
# Robust rendering helpers
# -------------------------
def _get(d, key, default=""):
    """Safe dict getter that never returns None."""
    if not isinstance(d, dict):
        return default
    val = d.get(key, default)
    return default if val is None else val

def _to_text(x, default=""):
    """
    Coerce any value to a safe string for DOCX/PDF.
    - None -> default
    - list/tuple -> comma-joined strings
    - dict -> compact JSON string
    - other -> str(x)
    """
    if x is None:
        return default
    if isinstance(x, (list, tuple)):
        return ", ".join(_to_text(y, default="") for y in x)
    if isinstance(x, dict):
        return json.dumps(x, ensure_ascii=False)
    return str(x)

def _as_list_of_dicts(seq, primary_key):
    """
    Normalize a sequence to a list of dicts.
    - dict -> keep
    - str  -> wrap as {primary_key: str}
    - else -> skip
    """
    out = []
    if not isinstance(seq, (list, tuple)):
        return out
    for it in seq:
        if isinstance(it, dict):
            out.append(it)
        elif isinstance(it, str):
            out.append({primary_key: it})
    return out

# -------------------------
# One-shot CAD generation
# -------------------------
def node_generate_cad_json_docx_pdf_openai(
    client: OpenAI,
    model: str,
    full_text_with_pages: str,
    pages,
    output_basename: str = "generated_CAD"
):
    # 0) Trim obvious whitespace noise to save tokens (but keep page markers)
    compact_text = re.sub(r"[ \t]+", " ", full_text_with_pages)
    compact_text = re.sub(r"\n{3,}", "\n\n", compact_text)

    # 1) Call LLM once for the whole contract
    user_prompt = CAD_USER_TMPL.format(
        contract_text=compact_text,
        style_guide=CAD_STYLE_GUIDE.strip()
    )
    raw = chat_json(client, model, CAD_SYSTEM, user_prompt, max_tokens=MAX_OUTPUT_TOKENS)

    # 2) Parse/repair JSON once; then a stricter repair if needed
    data = safe_json_loads(raw)
    if not data or not isinstance(data, dict):
        repair_prompt = (
            "Convert this content into ONE strict JSON object that matches the previously provided schema. "
            "Do not add comments or prose—JSON only.\n\n" + (raw or "")
        )
        repaired = chat_json(client, model, CAD_SYSTEM, repair_prompt, max_tokens=MAX_OUTPUT_TOKENS)
        data = safe_json_loads(repaired)
        if not data:
            # Last resort: re-extract from source with schema
            reemit_prompt = (
                "Re-extract from the CONTRACT TEXT using the exact schema previously given. "
                "Output STRICT JSON only. If a field is unknown, set it to null.\n\n"
                f"CONTRACT TEXT:\n{compact_text}\n\nSTYLE GUIDE:\n{CAD_STYLE_GUIDE.strip()}"
            )
            reemitted = chat_json(client, model, CAD_SYSTEM, reemit_prompt, max_tokens=MAX_OUTPUT_TOKENS)
            data = safe_json_loads(reemitted)
            if not data:
                raise RuntimeError("Failed to obtain valid CAD JSON from LLM.")

    # 3) Save JSON
    json_path = Path(tempfile.gettempdir()) / f"{output_basename}.json"
    with open(json_path, "w", encoding="utf8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    # -------- Normalize all sections defensively --------
    salient_features = _as_list_of_dicts(data.get("salient_features", []), "description")
    important_submittals = _as_list_of_dicts(data.get("important_submittals", []), "document")
    notice_information = _as_list_of_dicts(data.get("notice_information_clauses", []), "description")
    project_progress = _as_list_of_dicts(data.get("project_progress_clauses", []), "topic")

    pay = data.get("payment", {});          pay = pay if isinstance(pay, dict) else {}
    risk_matrix = data.get("risk_matrix", {}); risk_matrix = risk_matrix if isinstance(risk_matrix, dict) else {}
    sev_buckets = _as_list_of_dicts(risk_matrix.get("severity_buckets", []), "label")
    prob_buckets = _as_list_of_dicts(risk_matrix.get("probability_buckets", []), "label")

    risks = _as_list_of_dicts(data.get("risks_and_allocation", []), "risk_element")

    cda = data.get("claims_disputes_arbitration", {}); cda = cda if isinstance(cda, dict) else {}
    arbitration = cda.get("arbitration", {}); arbitration = arbitration if isinstance(arbitration, dict) else {}
    claims = _as_list_of_dicts(cda.get("claims_summary", []), "topic")
    disputes = _as_list_of_dicts(cda.get("dispute_areas", []), "description")
    excusable_delays = _as_list_of_dicts(cda.get("excusable_delays", []), "item")

    # 4) DOCX rendering — mirrors the sample CAD layout
    docx = DocxDocument()
    style = docx.styles["Normal"]; style.font.name = "Arial"; style.font.size = Pt(11)

    # Title
    docx.add_paragraph().add_run().add_break()
    docx.add_heading("Contract Appreciation Document (Generated)", level=1)

    # ---- 1. Salient Features of the Contract ----
    docx.add_heading("1. Salient Features of the Contract", level=2)
    if salient_features:
        table = docx.add_table(rows=1, cols=4)
        hdrs = ["Sl. No.", "Description", "Clause No.", "Details"]
        for i, h in enumerate(hdrs): table.rows[0].cells[i].text = h
        for idx, r in enumerate(salient_features, start=1):
            row = table.add_row().cells
            row[0].text = _to_text(idx)
            row[1].text = _to_text(_get(r, "description"))
            row[2].text = _to_text(_get(r, "clause_or_ref"))
            row[3].text = _to_text(_get(r, "details"))
    else:
        docx.add_paragraph("No salient features found.")

    # ---- 2. List of Important Submittals ----
    docx.add_heading("2. List of Important Submittals", level=2)
    if important_submittals:
        table = docx.add_table(rows=1, cols=6)
        for i, h in enumerate(["Sl. No.", "Stage", "Document", "Due", "Notes", "Clause No."]):
            table.rows[0].cells[i].text = h
        for idx, s in enumerate(important_submittals, start=1):
            c = table.add_row().cells
            c[0].text = _to_text(idx)
            c[1].text = _to_text(_get(s, "stage"))
            c[2].text = _to_text(_get(s, "document"))
            c[3].text = _to_text(_get(s, "due"))
            c[4].text = _to_text(_get(s, "notes"))
            c[5].text = _to_text(_get(s, "clause_or_ref"))
    else:
        docx.add_paragraph("No submittals listed.")

    # ---- 3. Notice / Information Clauses ----
    docx.add_heading("3. Notice / Information Clauses", level=2)
    if notice_information:
        table = docx.add_table(rows=1, cols=4)
        for i, h in enumerate(["Sl. No.", "Description", "Clause No.", "Details"]):
            table.rows[0].cells[i].text = h
        for idx, n in enumerate(notice_information, start=1):
            c = table.add_row().cells
            c[0].text = _to_text(idx)
            c[1].text = _to_text(_get(n, "description"))
            c[2].text = _to_text(_get(n, "clause_or_ref"))
            c[3].text = _to_text(_get(n, "details"))
    else:
        docx.add_paragraph("No notice/information clauses found.")

    # ---- 4. Project Progress (EOT, Escalation, Variation, Suspension, etc.) ----
    docx.add_heading("4. Important clauses pertaining to project progress - EOT, Escalation, Variation, Suspension, etc.", level=2)
    if project_progress:
        table = docx.add_table(rows=1, cols=4)
        for i, h in enumerate(["Sl. No.", "Topic", "Clause No.", "Summary"]):
            table.rows[0].cells[i].text = h
        for idx, ppc in enumerate(project_progress, start=1):
            c = table.add_row().cells
            c[0].text = _to_text(idx)
            c[1].text = _to_text(_get(ppc, "topic"))
            c[2].text = _to_text(_get(ppc, "clause_or_ref"))
            c[3].text = _to_text(_get(ppc, "summary"))
    else:
        docx.add_paragraph("No project progress clauses found.")

    # ---- 5. Payment Clause ----
    docx.add_heading("5. Payment Clause", level=2)
    def _p(label, val): docx.add_paragraph(f"{label}: {_to_text(val) if (val or val==0) else 'N/A'}")
    _p("Type of Contract", _get(pay, "type_of_contract"))
    _p("Mode of Payment", _get(pay, "mode_of_payment"))
    _p("Measurement and Evaluations", _get(pay, "measurement_and_evaluation"))
    _p("Billing Timing", _get(pay, "billing_timing"))
    _p("Other Terms", _get(pay, "other_terms"))

    money_lines = []
    cs = pay.get("contract_sum") if isinstance(pay.get("contract_sum"), dict) else {}
    if isinstance(cs, dict) and (cs.get("amount") is not None):
        money_lines.append(f"Contract Sum: {cs.get('amount')} {cs.get('currency') or ''}".strip())

    ma = pay.get("mobilization_advance") if isinstance(pay.get("mobilization_advance"), dict) else {}
    if isinstance(ma, dict):
        if ma.get("applicable") is True:
            if ma.get("amount_pct") is not None:
                money_lines.append(f"Mobilization Advance: Yes, {ma.get('amount_pct')}% @ {ma.get('interest') or 'N/A'}")
            else:
                money_lines.append("Mobilization Advance: Yes")
        elif ma.get("applicable") is False:
            money_lines.append("Mobilization Advance: No")

    rt = pay.get("retention") if isinstance(pay.get("retention"), dict) else {}
    if isinstance(rt, dict) and (rt.get("percent") is not None):
        money_lines.append(f"Retention: {rt.get('percent')}% ({rt.get('release_condition') or 'release terms N/A'})")

    es = pay.get("price_escalation") if isinstance(pay.get("price_escalation"), dict) else {}
    if isinstance(es, dict):
        if es.get("applicable") is True:
            basis = es.get("basis") or es.get("formula_or_reference")
            money_lines.append(f"Price Escalation: Applicable ({basis or 'basis/formula N/A'})")
        elif es.get("applicable") is False:
            money_lines.append("Price Escalation: Not applicable")

    for ln in money_lines:
        docx.add_paragraph("• " + _to_text(ln))

    ips = pay.get("interim_payments") if isinstance(pay.get("interim_payments"), list) else []
    if ips:
        docx.add_paragraph("Interim Payments:")
        for ip in ips:
            ip = ip if isinstance(ip, dict) else {}
            docx.add_paragraph(
                f"  - {_to_text(_get(ip, 'frequency', 'N/A'))}; "
                f"Certifier: {_to_text(_get(ip, 'certifier', 'N/A'))}; "
                f"DAYS after certification: {_to_text(_get(ip, 'payment_days_after_certification', 'N/A'))}"
            )

    # ---- 6. Risk Matrix & Allocation ----
    docx.add_heading("6. Identify Risk, Risk Matrix, and Risk Allocation", level=2)
    if sev_buckets:
        docx.add_paragraph("Severity Buckets:")
        for s in sev_buckets: docx.add_paragraph(f"  - {_to_text(_get(s,'label'))}: {_to_text(_get(s,'criteria'))}")
    if prob_buckets:
        docx.add_paragraph("Probability Buckets:")
        for p in prob_buckets: docx.add_paragraph(f"  - {_to_text(_get(p,'label'))}: {_to_text(_get(p,'criteria'))}")

    if risks:
        table = docx.add_table(rows=1, cols=9)
        hdrs = ["ID","Category","Risk Element","Probability","Severity","Rating","Owner","Mitigation","Sources"]
        for i, h in enumerate(hdrs): table.rows[0].cells[i].text = h
        for r in risks:
            c = table.add_row().cells
            c[0].text = _to_text(_get(r, "id"))
            c[1].text = _to_text(_get(r, "category"))
            c[2].text = _to_text(_get(r, "risk_element"))
            c[3].text = _to_text(_get(r, "probability"))
            c[4].text = _to_text(_get(r, "severity"))
            c[5].text = _to_text(_get(r, "rating"))
            c[6].text = _to_text(_get(r, "owner"))
            c[7].text = _to_text(_get(r, "mitigation"))
            c[8].text = _to_text(_get(r, "sources"))
    else:
        docx.add_paragraph("No risks listed.")

    # ---- 7. Claims, Dispute, and Arbitration ----
    docx.add_heading("7. Claim, Dispute, and Arbitration clause", level=2)
    docx.add_paragraph(
        f"Arbitration: Applicable: {_to_text('Yes' if _get(arbitration,'applicable') else 'No')}; "
        f"Forum/Rules: {_to_text(_get(arbitration,'forum_or_rules','N/A'))}"
    )
    docx.add_paragraph(f"Court Jurisdiction: {_to_text(_get(cda,'court_jurisdiction','N/A'))}")

    if claims:
        docx.add_paragraph("Claims Summary:")
        for cl in claims:
            topic = _to_text(_get(cl, "topic"))
            payable = _to_text(_get(cl, "payable_or_not", "Unclear"))
            notes = _to_text(_get(cl, "notes"))
            docx.add_paragraph(f"  - {topic}: {payable} — {notes}")

    if disputes:
        docx.add_paragraph("Dispute Areas:")
        for d in disputes:
            docx.add_paragraph(f"  - {_to_text(_get(d, 'description'))}")

    if excusable_delays:
        docx.add_paragraph("Excusable Delays:")
        for e in excusable_delays:
            docx.add_paragraph(f"  - {_to_text(_get(e, 'item'))}")

    delay_comp = _get(cda, "delay_compensation_clause")
    if delay_comp is not None:
        docx.add_paragraph(f"Delay Compensation Clause: {_to_text(delay_comp)}")

    # Save DOCX
    docx_path = Path(tempfile.gettempdir()) / f"{output_basename}.docx"
    docx.save(docx_path)

    # 5) 1-page Summary PDF (tabular, Platypus only — no canvas usage)
    pdf_path = Path(tempfile.gettempdir()) / f"{output_basename}.pdf"
    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        rightMargin=40, leftMargin=40, topMargin=60, bottomMargin=40
    )
    styles = getSampleStyleSheet()
    styles["Normal"].fontName = "Helvetica"
    styles["Normal"].fontSize = 10
    styles["Normal"].leading = 12
    h2 = ParagraphStyle(name="H2", parent=styles["Heading2"], fontName="Helvetica-Bold", spaceAfter=8)
    h3 = ParagraphStyle(name="H3", parent=styles["Heading3"], fontName="Helvetica-Bold", spaceBefore=6, spaceAfter=6)

    elements = []
    elements.append(Paragraph("Contract Appreciation Document (Summary)", h2))
    elements.append(Spacer(1, 6))
    elements.append(Paragraph("Key Salient Features", h3))
    elements.append(Spacer(1, 4))

    # Use normalized salient_features
    table_data = [["Sl. No.", "Description", "Details"]]
    for i, row in enumerate(salient_features[:6], start=1):
        desc = _to_text(_get(row, "description")).strip()
        det  = _to_text(_get(row, "details")).strip()
        table_data.append([_to_text(i), Paragraph(desc, styles["Normal"]), Paragraph(det, styles["Normal"])])

    table = Table(table_data, colWidths=[45, 210, 285], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#F0F0F0")),
        ("GRID", (0,0), (-1,-1), 0.5, colors.black),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("ALIGN", (0,0), (0,-1), "CENTER"),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("LEFTPADDING", (0,0), (-1,-1), 6),
        ("RIGHTPADDING", (0,0), (-1,-1), 6),
        ("TOPPADDING", (0,0), (-1,-1), 4),
        ("BOTTOMPADDING", (0,0), (-1,-1), 4),
    ]))
    elements.append(table)
    doc.build(elements)

    return str(data), str(json_path), str(docx_path), str(pdf_path)

# -------------------------
# Compliance (one-shot per rule) — UNCHANGED
# -------------------------
# def compliance_check_json_openai(
#     client: OpenAI,
#     model: str,
#     full_text_with_pages: str,
#     rules: List[str]
# ) -> List[Dict[str, Any]]:
#     results = []
#     for rule in rules:
#         uprompt = COMPLIANCE_USER_TMPL.format(contract_text=full_text_with_pages, rule=rule)
#         raw = chat_json(client, model, COMPLIANCE_SYSTEM, uprompt, max_tokens=800)
#         obj = safe_json_loads(raw)
#         if not obj:
#             # fallback: mark not present with low confidence
#             obj = {"rule": rule, "present": False, "summary": None, "quote": None, "sources": [], "confidence": 0.50}
#         results.append(obj)
#     return results
def _normalize_sources(x):
    if not x:
        return []
    if isinstance(x, list):
        return [str(i) for i in x]
    # sometimes a single string comes back
    return [str(x)]

def _sanitize_bool(x):
    if isinstance(x, bool):
        return x
    if isinstance(x, str):
        s = x.strip().lower()
        if s in ("true", "yes", "present", "1"):
            return True
        if s in ("false", "no", "absent", "0"):
            return False
    return False

def _sanitize_float(x, default=0.0):
    try:
        return float(x)
    except Exception:
        return default

def compliance_check_json_openai(
    client: OpenAI,
    model: str,
    full_text_with_pages: str,
    rules: List[str]
) -> List[Dict[str, Any]]:
    """
    Returns a list of dicts:
      {
        "rule": str,
        "present": bool,
        "summary": str|None,
        "quote": str|None,
        "sources": [str],
        "confidence": float
      }
    Robust to occasional non-JSON or string responses.
    """
    results: List[Dict[str, Any]] = []

    for rule in rules:
        uprompt = COMPLIANCE_USER_TMPL.format(
            contract_text=full_text_with_pages,
            rule=rule
        )
        raw = ""
        try:
            resp = client.chat.completions.create(
                model=model,
                temperature=0,
                max_tokens=800,
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": COMPLIANCE_SYSTEM},
                    {"role": "user", "content": uprompt},
                ],
            )
            raw = (resp.choices[0].message.content or "").strip()
        except Exception as e:
            # API failure → emit a synthetic "not present" row with error note
            results.append({
                "rule": rule,
                "present": False,
                "summary": f"Error calling model: {e}",
                "quote": None,
                "sources": [],
                "confidence": 0.0
            })
            continue

        # Try strict JSON first
        obj = safe_json_loads(raw)

        # If still not a dict, try a couple more lenient shapes
        if not isinstance(obj, dict):
            # Some models may (rarely) wrap JSON in a list
            if isinstance(obj, list) and obj and isinstance(obj[0], dict):
                obj = obj[0]
            else:
                # last resort default
                obj = {
                    "rule": rule,
                    "present": False,
                    "summary": None,
                    "quote": None,
                    "sources": [],
                    "confidence": 0.50
                }

        # Normalize fields so UI never crashes
        norm = {
            "rule": str(obj.get("rule", rule)),
            "present": _sanitize_bool(obj.get("present")),
            "summary": (obj.get("summary") if obj.get("summary") not in ("", None) else None),
            "quote": (obj.get("quote") if obj.get("quote") not in ("", None) else None),
            "sources": _normalize_sources(obj.get("sources")),
            "confidence": _sanitize_float(obj.get("confidence"), 0.7),
        }
        results.append(norm)

    return results





# ---- FULL-DOC CONFLICT AUDIT (CATEGORIES) ----
CONFLICT_SYSTEM = """
You are a contracts analyst performing a conflict audit on the FULL contract.
Return STRICT JSON ONLY. No markdown, no trailing text.
Detect contradictions/inconsistencies across the ENTIRE text.
Whenever possible, attach short quotes and page/ clause sources.
Classify each conflict into a normalized type and severity.
"""

CONFLICT_USER_TMPL = """Audit the FULL CONTRACT TEXT for internal contradictions and cross-references that do not align.
Return STRICT JSON matching the schema.

CONTRACT TEXT (page markers included):
{contract_text}

CHECK at least the following categories (and add others you truly find):
- commencement_vs_site_possession
- payment_term_mismatch
- retention_mismatch
- defect_liability_mismatch
- dispute_resolution_conflict
- performance_security_mismatch
- liquidated_damages_mismatch_or_missing_cap
- mobilization_advance_inconsistency
- escalation_clause_inconsistency
- warranty_vs_guarantee_overlap
- termination_notice_mismatch
- governing_law_multiple
- force_majeure_missing_or_weak
- tax_responsibility_conflict

SCHEMA:
{{
  "conflicts": [
    {{
      "type": "string",
      "severity": "Critical|High|Medium|Low",
      "message": "short human-readable description",
      "category": "string",                               // practical grouping (e.g., Financial / Payment)
      "resolution_hint": "string",
      "evidence": [                                       // ≤3 snippets
        {{"quote":"<=200 chars","sources":["page X","clause Y"]}}
      ],
      "confidence": 0.0
    }}
  ],
  "summary": "short (<=400 chars) overview of conflict landscape"
}}
RULES:
- Use only what is supported by the text.
- Prefer later/primary clauses when making inferences; if precedence unclear, mark confidence lower.
- Keep evidence quotes concise (<=200 chars) and attach page/ clause sources.
"""

def conflict_audit_openai(
    client: OpenAI,
    model: str,
    full_text_with_pages: str
) -> Dict[str, Any]:
    user_prompt = CONFLICT_USER_TMPL.format(contract_text=full_text_with_pages)
    raw = chat_json(client, model, CONFLICT_SYSTEM, user_prompt, max_tokens=MAX_OUTPUT_TOKENS)
    data = safe_json_loads(raw)
    if not isinstance(data, dict) or "conflicts" not in (data or {}):
        # attempt repair
        repair = chat_json(
            client, model, CONFLICT_SYSTEM,
            f"Reformat to STRICT JSON matching the SCHEMA only (no prose):\n\n{raw}",
            max_tokens=MAX_OUTPUT_TOKENS
        )
        data = safe_json_loads(repair) or {"conflicts": [], "summary": "repair_failed"}
    # Normalize
    data["conflicts"] = data.get("conflicts") or []
    for c in data["conflicts"]:
        c["type"] = c.get("type") or "other"
        c["severity"] = c.get("severity") or "Medium"
        c["message"] = c.get("message") or ""
        c["category"] = c.get("category") or "Other / Administrative"
        c["resolution_hint"] = c.get("resolution_hint") or "Review conflict and clarify governing clause."
        c["confidence"] = float(c.get("confidence") or 0.6)
        # normalize evidence
        ev = []
        for e in c.get("evidence") or []:
            if isinstance(e, dict):
                q = (e.get("quote") or "")[:200]
                src = e.get("sources") if isinstance(e.get("sources"), list) else []
                ev.append({"quote": q, "sources": src})
        c["evidence"] = ev[:3]
    data["summary"] = data.get("summary") or ""
    return data

def check_commencement_vs_site_possession(pages, full_text):
    date_pattern = r'((?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4})|(?:\d{1,2}\s+[A-Za-z]+\s+\d{4})|(?:[A-Za-z]+\s+\d{1,2},\s*\d{4}))'
    site_pos_pattern = re.compile(r'(site possession|possession of the site|hand over (?:the )?site)[^\.\n]{0,250}(' + date_pattern + r')', re.I)
    commence_pattern = re.compile(r'(commenc(?:e|ement) date|start date|date of commencement|date of commencement of works)[^\.\n]{0,250}(' + date_pattern + r')', re.I)

    site_matches, comm_matches = [], []
    for p in pages:
        text = p.get('text') or ""
        for m in site_pos_pattern.finditer(text):
            raw = m.group(2)
            try:
                dt = dateparser.parse(raw, dayfirst=True)
            except Exception:
                dt = None
            site_matches.append({"page": p['page'], "raw": raw, "date": dt, "snippet": m.group(0)})
        for m in commence_pattern.finditer(text):
            raw = m.group(2)
            try:
                dt = dateparser.parse(raw, dayfirst=True)
            except Exception:
                dt = None
            comm_matches.append({"page": p['page'], "raw": raw, "date": dt, "snippet": m.group(0)})

    conflicts = []
    if site_matches and comm_matches:
        comm_dates = [c['date'] for c in comm_matches if c['date']]
        site_dates = [s['date'] for s in site_matches if s['date']]
        if comm_dates and site_dates:
            min_comm = min(comm_dates)
            max_site = max(site_dates)
            if max_site > min_comm:
                conflicts.append({
                    "type": "commencement_vs_site_possession",
                    "severity": "Critical",
                    "message": f"Site possession ({max_site.date()}) occurs after commencement ({min_comm.date()}).",
                    "evidence": {"commencement": comm_matches, "site_possession": site_matches}
                })
    return conflicts

def check_payment_term_mismatch(pages):
    pay_pattern = re.compile(r'pay(?:ment)?\s+(?:within|in)\s+(\d{1,3})\s+days', re.I)
    matches = []
    for p in pages:
        for m in pay_pattern.finditer(p.get('text') or ""):
            try:
                matches.append({"page": p['page'], "days": int(m.group(1)), "snippet": m.group(0)})
            except Exception:
                pass
    conflicts = []
    if matches:
        days_set = sorted(set(m['days'] for m in matches))
        if len(days_set) > 1:
            conflicts.append({
                "type": "payment_term_mismatch",
                "severity": "High",
                "message": f"Different payment terms found: {days_set} days.",
                "evidence": matches
            })
    return conflicts

def check_retention_mismatch(pages):
    ret_pattern = re.compile(r'retention[^\d]{0,20}(\d{1,2})\s*%?', re.I)
    matches = []
    for p in pages:
        for m in ret_pattern.finditer(p.get('text') or ""):
            try:
                pct = int(m.group(1))
                matches.append({"page": p['page'], "pct": pct, "snippet": m.group(0)})
            except Exception:
                continue
    conflicts = []
    if matches:
        pct_set = sorted(set(m['pct'] for m in matches))
        if len(pct_set) > 1:
            conflicts.append({
                "type": "retention_mismatch",
                "severity": "High",
                "message": f"Different retention percentages found: {pct_set}%.",
                "evidence": matches
            })
    return conflicts

def check_defect_liability_mismatch(pages):
    dl_pattern = re.compile(r'(defect liability period|defects liability|warranty period)[^\d]{0,60}(\d{1,3})\s*(months|month|years|year)?', re.I)
    matches = []
    for p in pages:
        for m in dl_pattern.finditer(p.get('text') or ""):
            try:
                val = int(m.group(2))
                unit = (m.group(3) or "months").lower()
                months = val * 12 if "year" in unit else val
                matches.append({"page": p['page'], "value": months, "raw": m.group(0)})
            except Exception:
                continue
    conflicts = []
    if matches:
        vals = sorted(set(m['value'] for m in matches))
        if len(vals) > 1:
            conflicts.append({
                "type": "defect_liability_mismatch",
                "severity": "High",
                "message": f"Different defect liability lengths found (in months): {vals}.",
                "evidence": matches
            })
    return conflicts

def check_arbitration_vs_court_conflict(pages):
    arb_pattern = re.compile(r'\barbitrat', re.I)
    court_pattern = re.compile(r'\bcourt', re.I)
    arbs, courts = [], []
    for p in pages:
        t = p.get('text') or ""
        if arb_pattern.search(t):
            arbs.append({"page": p['page']})
        if court_pattern.search(t):
            courts.append({"page": p['page']})
    conflicts = []
    if arbs and courts:
        conflicts.append({
            "type": "dispute_resolution_conflict",
            "severity": "Critical",
            "message": "Arbitration and court references coexist; possible conflict in dispute resolution.",
            "evidence": {"arbitration_pages": arbs, "court_pages": courts}
        })
    return conflicts

CONFLICT_TO_CATEGORY = {
    "commencement_vs_site_possession": ("Contractual Disputes (Site Possession)", "Deem commencement as possession or grant EOT + cost; see Site Possession."),
    "payment_term_mismatch": ("Contractual Disputes (Payment)", "Harmonize payment timeline; prefer main body or latest-dated doc."),
    "retention_mismatch": ("Financial / Payment", "Clarify a single retention % and update annex."),
    "defect_liability_mismatch": ("Quality / Warranty", "Adopt stricter DLP or clarify governing schedule."),
    "dispute_resolution_conflict": ("Dispute Resolution", "Prefer arbitration with interim relief carve-out; remove ambiguity."),
}

def map_conflict_to_practical_category(conflict):
    ctype = conflict.get("type")
    cat, hint = CONFLICT_TO_CATEGORY.get(ctype, ("Other / Administrative", "Review conflict and map manually."))
    conflict["category"] = cat
    conflict["resolution_hint"] = hint
    return conflict

# def run_conflict_detection(pages, full_text):
#     conflicts = []
#     conflicts.extend(check_commencement_vs_site_possession(pages, full_text))
#     conflicts.extend(check_payment_term_mismatch(pages))
#     conflicts.extend(check_retention_mismatch(pages))
#     conflicts.extend(check_defect_liability_mismatch(pages))
#     conflicts.extend(check_arbitration_vs_court_conflict(pages))
#     conflicts = [map_conflict_to_practical_category(c) for c in conflicts]
#     return conflicts




def _first_int(nums):
    for n in nums:
        try:
            return int(n)
        except Exception:
            continue
    return None

def check_performance_security_mismatch(pages):
    # e.g., Performance Security/Bank Guarantee % inconsistent across pages (5% vs 10%)
    ps_pattern = re.compile(r'(performance (?:security|guarantee|bank guarantee)).{0,80}?(\d{1,2})\s*%?', re.I)
    hits = []
    for p in pages:
        t = p.get('text') or ""
        for m in ps_pattern.finditer(t):
            try:
                pct = int(m.group(2))
                hits.append({"page": p['page'], "pct": pct, "snippet": m.group(0)[:220]})
            except Exception:
                continue
    conflicts = []
    if hits:
        uniq = sorted(set(h['pct'] for h in hits))
        if len(uniq) > 1:
            conflicts.append({
                "type": "performance_security_mismatch",
                "severity": "High",
                "message": f"Different Performance Security percentages found: {uniq}%.",
                "evidence": hits
            })
    return conflicts

def check_liquidated_damages_mismatch_or_missing_cap(pages):
    # Find LD rate and cap; flag if multiple rates or cap missing when LD appears
    ld_rate_pattern = re.compile(r'(liquidated damages|LD).{0,120}?(\d{1,2})\s*%?', re.I)
    cap_pattern = re.compile(r'(cap|maximum|ceiling).{0,60}?(\d{1,3})\s*%?', re.I)
    ld_hits, cap_hits = [], []
    for p in pages:
        t = p.get('text') or ""
        for m in ld_rate_pattern.finditer(t):
            try:
                pct = int(m.group(2))
                ld_hits.append({"page": p['page'], "pct": pct, "snippet": m.group(0)[:220]})
            except Exception:
                continue
        for m in cap_pattern.finditer(t):
            # Very rough: capture any % after "cap/maximum/ceiling"
            nums = re.findall(r'(\d{1,3})\s*%?', m.group(0))
            val = _first_int(nums)
            if val is not None:
                cap_hits.append({"page": p['page'], "pct": val, "snippet": m.group(0)[:220]})
    conflicts = []
    if ld_hits:
        uniq_rates = sorted(set(h['pct'] for h in ld_hits))
        if len(uniq_rates) > 1:
            conflicts.append({
                "type": "liquidated_damages_mismatch",
                "severity": "High",
                "message": f"Different LD rates found: {uniq_rates}%.",
                "evidence": ld_hits
            })
        # If LD exists but no visible cap anywhere, warn
        if not cap_hits:
            conflicts.append({
                "type": "liquidated_damages_missing_cap",
                "severity": "Medium",
                "message": "Liquidated damages mentioned but no explicit cap/maximum found nearby.",
                "evidence": ld_hits[:3]
            })
    return conflicts

def check_mobilization_advance_inconsistency(pages):
    # Detect multiple advance % values
    adv_pattern = re.compile(r'(mobilization|mobilisation).{0,40}advance.{0,80}?(\d{1,2})\s*%?', re.I)
    hits = []
    for p in pages:
        t = p.get('text') or ""
        for m in adv_pattern.finditer(t):
            try:
                pct = int(m.group(2))
                hits.append({"page": p['page'], "pct": pct, "snippet": m.group(0)[:220]})
            except Exception:
                continue
    conflicts = []
    if hits:
        uniq = sorted(set(h['pct'] for h in hits))
        if len(uniq) > 1:
            conflicts.append({
                "type": "mobilization_advance_inconsistency",
                "severity": "High",
                "message": f"Mobilization advance percentages conflict: {uniq}%.",
                "evidence": hits
            })
    return conflicts

def check_escalation_clause_inconsistency(pages):
    # If escalation is "not applicable" somewhere and formula/basis elsewhere
    not_app_pattern = re.compile(r'(price|cost)\s+escalation.{0,40}(not\s+applicable|no\s+escalation)', re.I)
    basis_pattern = re.compile(r'(price|cost)\s+escalation.{0,80}(WPI|CPI|index|formula|basis|escalation shall apply)', re.I)
    na_hits, basis_hits = [], []
    for p in pages:
        t = p.get('text') or ""
        for m in not_app_pattern.finditer(t):
            na_hits.append({"page": p['page'], "snippet": m.group(0)[:220]})
        for m in basis_pattern.finditer(t):
            basis_hits.append({"page": p['page'], "snippet": m.group(0)[:220]})
    conflicts = []
    if na_hits and basis_hits:
        conflicts.append({
            "type": "escalation_clause_inconsistency",
            "severity": "High",
            "message": "Escalation marked 'not applicable' in one place but basis/formula mentioned elsewhere.",
            "evidence": (na_hits[:2] + basis_hits[:2])
        })
    return conflicts

def check_warranty_vs_guarantee_overlap(pages):
    # Different durations between "warranty", "guarantee", "defects liability"
    period_pat = re.compile(r'(warranty|guarantee|defects liability).{0,60}?(\d{1,3})\s*(months?|years?)', re.I)
    vals = []
    for p in pages:
        t = p.get('text') or ""
        for m in period_pat.finditer(t):
            try:
                n = int(m.group(2))
                unit = (m.group(3) or "").lower()
                months = n * 12 if 'year' in unit else n
                vals.append({"page": p['page'], "months": months, "snippet": m.group(0)[:220]})
            except Exception:
                continue
    conflicts = []
    if vals:
        uniq = sorted(set(v['months'] for v in vals))
        if len(uniq) > 1:
            conflicts.append({
                "type": "warranty_vs_guarantee_overlap",
                "severity": "Medium",
                "message": f"Warranty/Guarantee/DLP durations differ: {uniq} months.",
                "evidence": vals
            })
    return conflicts

def check_termination_notice_mismatch(pages):
    # Mismatch in termination notice period
    term_pat = re.compile(r'(termination|terminate).{0,80}?(\d{1,3})\s+days', re.I)
    hits = []
    for p in pages:
        t = p.get('text') or ""
        for m in term_pat.finditer(t):
            try:
                days = int(m.group(2))
                hits.append({"page": p['page'], "days": days, "snippet": m.group(0)[:220]})
            except Exception:
                continue
    conflicts = []
    if hits:
        uniq = sorted(set(h['days'] for h in hits))
        if len(uniq) > 1:
            conflicts.append({
                "type": "termination_notice_mismatch",
                "severity": "High",
                "message": f"Termination notice periods conflict: {uniq} days.",
                "evidence": hits
            })
    return conflicts

def check_governing_law_multiple(pages):
    # Multiple governing law jurisdictions
    gov_pat = re.compile(r'(governed by|governing law).{0,60}?laws? of\s+([A-Za-z ,.&-]+)', re.I)
    hits = []
    for p in pages:
        t = p.get('text') or ""
        for m in gov_pat.finditer(t):
            juris = (m.group(2) or "").strip().rstrip(".;,")
            if juris:
                hits.append({"page": p['page'], "jurisdiction": juris, "snippet": m.group(0)[:220]})
    conflicts = []
    if hits:
        uniq = sorted(set(h['jurisdiction'] for h in hits))
        if len(uniq) > 1:
            conflicts.append({
                "type": "governing_law_multiple",
                "severity": "Critical",
                "message": f"Multiple governing law jurisdictions found: {uniq}.",
                "evidence": hits
            })
    return conflicts

def check_force_majeure_missing_or_weak(pages):
    # "Force majeure" mentioned? If not at all, flag missing (Low/Medium). If present but too vague (no examples), flag weak.
    fm_pat = re.compile(r'force\s+majeure', re.I)
    examples_pat = re.compile(r'(act of god|war|riot|epidemic|pandemic|natural disaster|earthquake|flood|strike)', re.I)
    fm_pages, example_pages = set(), set()
    fm_hits, ex_hits = [], []
    for p in pages:
        t = p.get('text') or ""
        if fm_pat.search(t):
            fm_pages.add(p['page'])
            fm_hits.append({"page": p['page'], "snippet": (fm_pat.search(t).string[:220])})
        if examples_pat.search(t):
            example_pages.add(p['page'])
            ex_hits.append({"page": p['page'], "snippet": (examples_pat.search(t).string[:220])})
    conflicts = []
    if not fm_pages:
        conflicts.append({
            "type": "force_majeure_missing",
            "severity": "Medium",
            "message": "No force majeure clause detected.",
            "evidence": []
        })
    elif fm_pages and not example_pages:
        conflicts.append({
            "type": "force_majeure_weak",
            "severity": "Low",
            "message": "Force majeure mentioned but lacks illustrative events/examples.",
            "evidence": fm_hits[:3]
        })
    return conflicts

def check_tax_responsibility_conflict(pages):
    # Conflicting responsibility for GST/TDS/taxes (Employer vs Contractor)
    tax_pat = re.compile(r'(GST|TDS|tax(?:es)?)', re.I)
    resp_pat = re.compile(r'(employer|contractor|owner)\s+(?:shall|will|to)\s+(?:bear|pay|be responsible)', re.I)
    hits = []
    for p in pages:
        t = p.get('text') or ""
        if tax_pat.search(t):
            # pull small snippets around 'tax' lines
            for line in t.splitlines():
                if tax_pat.search(line):
                    snip = line.strip()
                    # crude attribution search
                    subj = "Unclear"
                    subj_m = resp_pat.search(line)
                    if subj_m:
                        subj = subj_m.group(1).capitalize()
                    hits.append({"page": p['page'], "subject": subj, "snippet": snip[:220]})
    conflicts = []
    if hits:
        subjects = sorted(set(h['subject'] for h in hits if h['subject'] != "Unclear"))
        if len(subjects) > 1:
            conflicts.append({
                "type": "tax_responsibility_conflict",
                "severity": "Medium",
                "message": f"Different parties appear responsible for taxes: {subjects}.",
                "evidence": hits[:6]
            })
    return conflicts

#CATEGORY MAP (EXTEND) 

CONFLICT_TO_CATEGORY.update({
    "performance_security_mismatch": ("Financial / Security", "Confirm one Performance Security % and amend all references."),
    "liquidated_damages_mismatch": ("Delay / Damages", "Harmonize LD rate; prefer latest governing clause."),
    "liquidated_damages_missing_cap": ("Delay / Damages", "Add an explicit LD cap to limit exposure."),
    "mobilization_advance_inconsistency": ("Financial / Advance", "Fix single advance % and recovery terms."),
    "escalation_clause_inconsistency": ("Price Adjustment", "Clarify whether escalation applies; remove contradictions."),
    "warranty_vs_guarantee_overlap": ("Quality / Warranty", "Align warranty/guarantee/DLP durations in one schedule."),
    "termination_notice_mismatch": ("Termination", "Adopt a single notice period and update cross-references."),
    "governing_law_multiple": ("Legal / Governing Law", "Choose one governing law; delete conflicting mentions."),
    "force_majeure_missing": ("Force Majeure", "Insert standard force majeure clause with examples."),
    "force_majeure_weak": ("Force Majeure", "Add non-exhaustive examples and procedure for notice/mitigation."),
    "tax_responsibility_conflict": ("Taxes / Commercial", "Clarify which party bears which taxes; align everywhere."),
})

#  (EXTEND one )

def run_conflict_detection(pages, full_text):
    conflicts = []
    # existing detectors
    conflicts.extend(check_commencement_vs_site_possession(pages, full_text))
    conflicts.extend(check_payment_term_mismatch(pages))
    conflicts.extend(check_retention_mismatch(pages))
    conflicts.extend(check_defect_liability_mismatch(pages))
    conflicts.extend(check_arbitration_vs_court_conflict(pages))
    # new detectors
    conflicts.extend(check_performance_security_mismatch(pages))
    conflicts.extend(check_liquidated_damages_mismatch_or_missing_cap(pages))
    conflicts.extend(check_mobilization_advance_inconsistency(pages))
    conflicts.extend(check_escalation_clause_inconsistency(pages))
    conflicts.extend(check_warranty_vs_guarantee_overlap(pages))
    conflicts.extend(check_termination_notice_mismatch(pages))
    conflicts.extend(check_governing_law_multiple(pages))
    conflicts.extend(check_force_majeure_missing_or_weak(pages))
    conflicts.extend(check_tax_responsibility_conflict(pages))

    # map to category & add hints
    conflicts = [map_conflict_to_practical_category(c) for c in conflicts]
    return conflicts
# ---- Merge LLM + regex conflicts (upgraded) ----
_SEVERITY_RANK = {"Critical": 3, "High": 2, "Medium": 1, "Low": 0}

def _norm_text(x: Any) -> str:
    if not isinstance(x, str):
        return ""
    # light normalization to improve dedupe keys
    t = re.sub(r"\s+", " ", x).strip()
    return t[:300]  # cap key length

def _best_severity(a: str, b: str) -> str:
    ra, rb = _SEVERITY_RANK.get(a or "", -1), _SEVERITY_RANK.get(b or "", -1)
    return a if ra >= rb else b

def _best_confidence(a: Any, b: Any) -> float:
    try: ca = float(a)
    except: ca = 0.0
    try: cb = float(b)
    except: cb = 0.0
    return max(ca, cb)

def _ensure_evidence_list(ev):
    out = []
    if isinstance(ev, list):
        for e in ev:
            if isinstance(e, dict):
                quote = _norm_text(e.get("quote") or "")[:200]
                srcs = e.get("sources") if isinstance(e.get("sources"), list) else []
                out.append({"quote": quote, "sources": srcs})
    return out

def _merge_evidence(ev1, ev2, cap=3):
    # dedupe by (quote, tuple(sources))
    seen = set()
    merged = []
    for e in (ev1 or []) + (ev2 or []):
        q = _norm_text(e.get("quote") or "")
        s = tuple(e.get("sources") or [])
        key = (q, s)
        if q and key not in seen:
            merged.append({"quote": q[:200], "sources": list(s)[:4]})
            seen.add(key)
        if len(merged) >= cap:
            break
    return merged

def _apply_category_hint(conflict: Dict[str, Any]) -> Dict[str, Any]:
    if not conflict.get("category") or not conflict.get("resolution_hint"):
        t = conflict.get("type")
        if t in CONFLICT_TO_CATEGORY:
            cat, hint = CONFLICT_TO_CATEGORY[t]
            conflict["category"] = conflict.get("category") or cat
            conflict["resolution_hint"] = conflict.get("resolution_hint") or hint
    return conflict

def _canonicalize(c: Dict[str, Any]) -> Dict[str, Any]:
    # Ensure required fields exist and are shaped consistently
    c = dict(c or {})
    c["type"] = _norm_text(c.get("type") or "other") or "other"
    c["severity"] = c.get("severity") or "Medium"
    c["message"] = _norm_text(c.get("message") or "")
    c["category"] = c.get("category") or ""
    c["resolution_hint"] = c.get("resolution_hint") or ""
    c["confidence"] = _best_confidence(c.get("confidence"), 0.0)
    c["evidence"] = _ensure_evidence_list(c.get("evidence"))
    return _apply_category_hint(c)

def _pages_signature(evidence) -> str:
    # build a coarse signature like "p3|p7|p12" to help dedupe near-duplicates
    pages = []
    for e in evidence or []:
        for s in (e.get("sources") or []):
            m = re.search(r'page\s*(\d+)', s, re.I)
            if m:
                pages.append(int(m.group(1)))
    pages = sorted(set(pages))[:6]
    return "|".join(f"p{p}" for p in pages)

def merge_conflicts(llm_data: Dict[str, Any], regex_conflicts: List[Dict[str, Any]]) -> Dict[str, Any]:
    out = {"conflicts": [], "summary": ""}
    buckets: Dict[tuple, Dict[str, Any]] = {}

    # 1) ingest LLM conflicts first (preferred)
    for c in (llm_data.get("conflicts") or []):
        c = _canonicalize(c)
        key = (c["type"], c["message"])  # primary dedupe key
        # add a secondary “near-dup” key based on pages if message varies slightly
        near = (c["type"], _pages_signature(c["evidence"]))
        if key not in buckets and near in buckets:
            key = near
        if key not in buckets:
            buckets[key] = c
        else:
            # merge into existing
            base = buckets[key]
            base["severity"] = _best_severity(base.get("severity"), c.get("severity"))
            base["confidence"] = _best_confidence(base.get("confidence"), c.get("confidence"))
            base["evidence"] = _merge_evidence(base.get("evidence"), c.get("evidence"))
            # prefer longer message if base is short
            if len(c.get("message") or "") > len(base.get("message") or ""):
                base["message"] = c.get("message")

    # 2) append regex conflicts (mapped) as secondary source
    for r in (regex_conflicts or []):
        mapped = {
            "type": r.get("type"),
            "severity": r.get("severity"),
            "message": r.get("message"),
            "category": r.get("category"),
            "resolution_hint": r.get("resolution_hint"),
            "evidence": [
                {"quote": (ev.get("snippet") or "")[:200], "sources": [f"page {ev.get('page')}"]}
                for ev in (r.get("evidence") or []) if isinstance(ev, dict)
            ],
            "confidence": 0.65
        }
        mapped = _canonicalize(mapped)
        key = (mapped["type"], mapped["message"])
        near = (mapped["type"], _pages_signature(mapped["evidence"]))
        if key not in buckets and near in buckets:
            key = near
        if key not in buckets:
            buckets[key] = mapped
        else:
            base = buckets[key]
            base["severity"] = _best_severity(base.get("severity"), mapped.get("severity"))
            base["confidence"] = _best_confidence(base.get("confidence"), mapped.get("confidence"))
            base["evidence"] = _merge_evidence(base.get("evidence"), mapped.get("evidence"))
            if len(mapped.get("message") or "") > len(base.get("message") or ""):
                base["message"] = mapped.get("message")

    # 3) sort for display: severity desc → confidence desc → type
    merged = list(buckets.values())
    merged.sort(key=lambda c: (_SEVERITY_RANK.get(c.get("severity"), -1), float(c.get("confidence") or 0.0), c.get("type")), reverse=True)

    out["conflicts"] = merged
    out["summary"] = (llm_data.get("summary") or "").strip()
    return out
def merge_conflicts(llm_data: Dict[str, Any], regex_conflicts: List[Dict[str, Any]]) -> Dict[str, Any]:
    out = {"conflicts": [], "summary": ""}
    seen = set()

    def _key(c):
        return (c.get("type",""), c.get("message",""))

    # 1) Take LLM conflicts first
    for c in (llm_data.get("conflicts") or []):
        if not isinstance(c, dict): 
            continue
        k = _key(c)
        if k not in seen:
            out["conflicts"].append(c)
            seen.add(k)

    # 2) Map regex conflicts to LLM-like shape and add
    for r in (regex_conflicts or []):
        if not isinstance(r, dict): 
            continue

        # Flatten evidence to common shape
        evs = []
        evidence = r.get("evidence")
        if isinstance(evidence, list):
            # list of items with page/snippet
            for ev in evidence:
                if isinstance(ev, dict):
                    quote = (ev.get("snippet") or ev.get("raw") or "")[:200]
                    src = [f"page {ev.get('page')}"] if ev.get("page") is not None else []
                    evs.append({"quote": quote, "sources": src})
        elif isinstance(evidence, dict):
            # dict of lists: merge top-3 snippets across keys
            for lst in evidence.values():
                if isinstance(lst, list):
                    for ev in lst:
                        if isinstance(ev, dict):
                            quote = (ev.get("snippet") or ev.get("raw") or "")[:200]
                            src = [f"page {ev.get('page')}"] if ev.get("page") is not None else []
                            evs.append({"quote": quote, "sources": src})
            evs = evs[:3]

        mapped = {
            "type": r.get("type") or "",
            "severity": r.get("severity") or "",
            "category": r.get("category") or "",
            "message": r.get("message") or "",
            "resolution_hint": r.get("resolution_hint") or "",
            "evidence": evs,
            "confidence": 0.65,
        }
        k = _key(mapped)
        if k not in seen:
            out["conflicts"].append(mapped)
            seen.add(k)

    out["summary"] = (llm_data.get("summary") or "").strip()
    return out


st.set_page_config(page_title="Contract CAD (ChatGPT One-Shot)", layout="wide")
st.title("Automated Contract CAD Generator")

with st.sidebar:
    st.header("Model & Controls")
    model_name = st.selectbox(
        "Select OpenAI model",
        ["gpt-4.1-mini", "gpt-4.1-turbo", "gpt-4.1", "gpt-5.1"],
        index=2
    )
    st.caption("Tip: use gpt-5.1 for very large contracts (~150k tokens).")

    with st.expander("Advanced generation options", expanded=False):
        st.session_state["opt_temperature"] = st.slider("Temperature", 0.0, 1.0, 0.0, 0.1)
        st.session_state["opt_max_output_tokens"] = st.number_input(
            "Max output tokens (CAD JSON)", min_value=2000, max_value=64000, value=8000, step=500
        )
        st.session_state["opt_summary_rows"] = st.number_input(
            "Summary PDF: no. of salient features rows", min_value=3, max_value=20, value=6, step=1
        )

    uploaded_file = st.file_uploader("Upload contract PDF (scanned or digital)", type=["pdf"])
    regen = st.button("Extract / Rebuild context")
    clear_all = st.button("Clear session")

if clear_all:
    for k in list(st.session_state.keys()):
        del st.session_state[k]
    st.experimental_rerun()

# Extract & store
if uploaded_file and (regen or not st.session_state.get("graph_built", False)):
    try:
        with st.spinner("Extracting text (pdfplumber preferred; OCR fallback)..."):
            client = build_openai_client()
            pdf_bytes = uploaded_file.read()
            pages, full_text = extract_text_prefer_pdfplumber(pdf_bytes, dpi=200)
            st.session_state["pages"] = pages
            st.session_state["raw_text"] = full_text  # includes --- PAGE N --- markers
            st.session_state["graph_built"] = True
            # clear derived artifacts when rebuilding
            for k in ["cad_json", "cad_docx", "cad_pdf", "llm_conflicts", "merged_conflicts"]:
                st.session_state.pop(k, None)
        st.success("PDF processed.")
    except Exception as e:
        st.exception(e)
        st.error("Processing failed.")

# Layout
col1, col2 = st.columns([2, 1])

# Chat Q&A (one-shot on full text)
with col1:
    st.subheader("Chat with the Contract (One-Shot)")
    if not st.session_state.get("graph_built", False):
        st.info("Upload & Extract a PDF to enable chat.")
    else:
        if "conversation" not in st.session_state:
            st.session_state["conversation"] = []
        for turn in st.session_state["conversation"]:
            st.chat_message(turn["role"]).write(turn["text"])

        user_q = st.chat_input("Ask about the contract (e.g., 'What is the payment clause?')")
        if user_q:
            st.session_state["conversation"].append({"role": "user", "text": user_q})
            client = build_openai_client()
            uprompt = QA_USER_TMPL.format(contract_text=st.session_state["raw_text"], question=user_q)
            try:
                resp = client.chat.completions.create(
                    model=model_name,
                    temperature=0,
                    max_tokens=900,
                    messages=[
                        {"role": "system", "content": QA_SYSTEM},
                        {"role": "user", "content": uprompt},
                    ],
                )
                ans = (resp.choices[0].message.content or "").strip()
            except Exception as e:
                ans = f"Error: {e}"
            st.session_state["conversation"].append({"role": "assistant", "text": ans})
            st.chat_message("assistant").write(ans)

# CAD / Compliance / Conflicts
with col2:
    st.subheader("CAD Generator")
    if not st.session_state.get("graph_built", False):
        st.info("Upload & Extract a PDF to enable CAD generation.")
    else:
        if st.button("Generate CAD (JSON + DOCX + PDF)"):
            try:
                client = build_openai_client()
                full_text = st.session_state["raw_text"]
                pages = st.session_state["pages"]

                # --- Local overrides; avoid 'global' mutation ---
                local_max_output_tokens = int(st.session_state.get("opt_max_output_tokens", MAX_OUTPUT_TOKENS))
                local_temperature = float(st.session_state.get("opt_temperature", TEMPERATURE))

                with st.spinner("Generating one-shot CAD via ChatGPT..."):
                    # Pass overrides through your chat_json calls by temporarily shadowing globals if needed
                    # (If your node_generate_... uses the globals directly, you can pass them in as kwargs instead.)
                    _ = node_generate_cad_json_docx_pdf_openai(
                        client=client,
                        model=model_name,
                        full_text_with_pages=full_text,
                        pages=pages,
                        output_basename="generated_CAD"
                    )
                    # Note: if you want the overrides to affect CAD LLM calls,
                    # plumb them down into node_generate_cad_json_docx_pdf_openai
                    # so it uses local_max_output_tokens/local_temperature.

                # node_generate_* already returns paths; capture them again explicitly
                _, json_path, docx_path, pdf_path = node_generate_cad_json_docx_pdf_openai(
                    client, model_name, full_text, pages, output_basename="generated_CAD"
                )
                st.session_state["cad_json"] = json_path
                st.session_state["cad_docx"] = docx_path
                st.session_state["cad_pdf"] = pdf_path
                st.success("CAD generated.")
            except Exception as e:
                st.exception(e)
                st.error("CAD generation failed.")

        # Downloads
        if st.session_state.get("cad_json"):
            try:
                with open(st.session_state["cad_json"], "rb") as f:
                    st.download_button("Download CAD JSON", f.read(), file_name=Path(st.session_state['cad_json']).name)
            except Exception as e:
                st.error(f"Failed to open CAD JSON: {e}")
        if st.session_state.get("cad_docx"):
            try:
                with open(st.session_state["cad_docx"], "rb") as f:
                    st.download_button("Download CAD DOCX", f.read(), file_name=Path(st.session_state['cad_docx']).name)
            except Exception as e:
                st.error(f"Failed to open CAD DOCX: {e}")
        if st.session_state.get("cad_pdf"):
            try:
                with open(st.session_state["cad_pdf"], "rb") as f:
                    st.download_button("Download CAD PDF", f.read(), file_name=Path(st.session_state['cad_pdf']).name)
            except Exception as e:
                st.error(f"Failed to open CAD PDF: {e}")

    st.markdown("---")
    st.subheader("📑 Compliance Check (One-Shot per rule)")
    if not st.session_state.get("graph_built", False):
        st.info("Upload a contract first.")
    else:
        default_rules = "Does the contract specify payment terms?\nIs there a termination clause?\nWhat is the governing law?"
        rules_input = st.text_area("Enter compliance rules (one per line)", value=default_rules)
        colA, colB = st.columns(2)
        with colA:
            run_comp = st.button("Run Compliance Check")
        with colB:
            export_comp = st.button("Export Compliance CSV")

        if run_comp or export_comp:
            try:
                import json, csv, tempfile
                rules = [r.strip() for r in rules_input.splitlines() if r.strip()]
                if not rules:
                    st.warning("Please enter at least one rule.")
                else:
                    client = build_openai_client()
                    full_text = st.session_state["raw_text"]

                    with st.spinner("Checking compliance..."):
                        results = compliance_check_json_openai(client, model_name, full_text, rules)

                    # --- Robust normalization before rendering ---
                    if isinstance(results, str):
                        try:
                            results = json.loads(results)
                        except Exception:
                            results = []

                    st.subheader("Compliance Report")

                    show_debug = st.checkbox("Show raw compliance debug", value=False)
                    if show_debug:
                        st.write("Type of results:", type(results))
                        st.write("Length:", len(results) if isinstance(results, list) else "n/a")
                        st.json(results)

                    if not results:
                        st.info("No compliance results returned. Check rules input and raw text extraction.")
                    else:
                        df_rows = []
                        for item in results:
                            # Ensure dict per row
                            if not isinstance(item, dict):
                                try:
                                    item = json.loads(item) if isinstance(item, str) else {}
                                except Exception:
                                    if show_debug:
                                        st.warning(f"Skipping malformed compliance item: {str(item)[:150]}")
                                    continue

                            src = item.get("sources") or []
                            src_txt = ", ".join(str(s) for s in src) if isinstance(src, list) else str(src)

                            st.markdown(
                                f"**Rule:** {item.get('rule','')}  \n"
                                f"**Present:** {item.get('present','')}  \n"
                                f"**Summary:** {item.get('summary','')}  \n"
                                f"**Quote:** {item.get('quote','')}  \n"
                                f"**Sources:** {src_txt}  \n"
                                f"**Confidence:** {item.get('confidence','')}"
                            )

                            df_rows.append([
                                item.get("rule",""),
                                item.get("present",""),
                                item.get("summary",""),
                                item.get("quote",""),
                                src_txt,
                                item.get("confidence",""),
                            ])

                        if df_rows:
                            st.dataframe(pd.DataFrame(
                                df_rows, columns=["Rule","Present","Summary","Quote","Sources","Confidence"]
                            ))

                        if export_comp and df_rows:
                            tmp = Path(tempfile.gettempdir()) / "compliance_report.csv"
                            with open(tmp, "w", newline="", encoding="utf-8") as f:
                                w = csv.writer(f)
                                w.writerow(["Rule","Present","Summary","Quote","Sources","Confidence"])
                                for row in df_rows:
                                    w.writerow(row)
                            with open(tmp, "rb") as f:
                                st.download_button("Download Compliance CSV", f.read(), file_name="compliance_report.csv")

            except Exception as e:
                st.exception(e)
                st.error("Compliance check failed.")

    st.markdown("---")
    st.subheader("⚠️ Conflict Detection")
    if not st.session_state.get("graph_built", False):
        st.info("Upload a contract first.")
    else:
        use_llm_merge = st.checkbox("Include LLM conflict miner and merge with regex (recommended)", value=True,
                                    help="Runs an LLM to extract conflicts from full text and merges with regex detectors.")
        colX, colY, colZ = st.columns([1,1,1])
        with colX:
            run_conf_btn = st.button("Run Conflict Detection")
        with colY:
            export_conf_btn = st.button("Export Conflicts JSON")
        with colZ:
            export_conf_csv_btn = st.button("Export Conflicts CSV")

        if run_conf_btn or export_conf_btn or export_conf_csv_btn:
            try:
                import json, csv, tempfile
                pages = st.session_state.get("pages")
                full_text = st.session_state.get("raw_text")

                with st.spinner("Detecting conflicts (regex)..."):
                    regex_conflicts = run_conflict_detection(pages, full_text)

                merged_payload = {"conflicts": [], "summary": ""}
                if use_llm_merge:
                    llm_conf = {"conflicts": [], "summary": ""}
                    if "run_llm_conflict_miner" in globals():
                        with st.spinner("Mining conflicts via LLM..."):
                            client = build_openai_client()
                            llm_conf = run_llm_conflict_miner(client, model_name, full_text) or {"conflicts": [], "summary": ""}
                    else:
                        st.info("LLM conflict miner not found in this file. Only regex conflicts will be shown.")
                    merged_payload = merge_conflicts(llm_conf, regex_conflicts)
                else:
                    merged_payload = merge_conflicts({"conflicts": [], "summary": ""}, regex_conflicts)

                st.session_state["merged_conflicts"] = merged_payload

                # Render results
                conflicts = merged_payload.get("conflicts") or []
                if not conflicts:
                    st.success("No conflicts detected by automated checks.")
                else:
                    st.warning(f"{len(conflicts)} potential conflicts detected.")
                    for idx, c in enumerate(conflicts, start=1):
                        st.markdown(
                            f"### {idx}. Category: **{c.get('category','')}**  \n"
                            f"**Type:** {c.get('type','')}  \n**Severity:** {c.get('severity','')}  \n"
                            f"**Confidence:** {c.get('confidence','')}  \n"
                            f"**Message:** {c.get('message','')}"
                        )
                        if c.get("resolution_hint"):
                            st.markdown(f"**Suggested quick fix:** {c.get('resolution_hint')}")
                        if c.get("evidence"):
                            st.json(c.get("evidence"))

                # Exports
                if export_conf_btn and st.session_state.get("merged_conflicts"):
                    tmp = Path(tempfile.gettempdir()) / "conflicts_merged.json"
                    with open(tmp, "w", encoding="utf-8") as f:
                        json.dump(st.session_state["merged_conflicts"], f, ensure_ascii=False, indent=2)
                    with open(tmp, "rb") as f:
                        st.download_button("Download Conflicts JSON", f.read(), file_name="conflicts_merged.json")

                if export_conf_csv_btn and st.session_state.get("merged_conflicts"):
                    tmpc = Path(tempfile.gettempdir()) / "conflicts_merged.csv"
                    with open(tmpc, "w", newline="", encoding="utf-8") as f:
                        w = csv.writer(f)
                        w.writerow(["Type","Category","Severity","Confidence","Message","Resolution Hint","Evidence (sources)","Evidence (quotes)"])
                        for c in st.session_state["merged_conflicts"].get("conflicts", []):
                            srcs, quotes = [], []
                            for e in c.get("evidence") or []:
                                srcs.append("; ".join(e.get("sources") or []))
                                quotes.append(e.get("quote") or "")
                            w.writerow([
                                c.get("type",""), c.get("category",""), c.get("severity",""),
                                c.get("confidence",""), c.get("message",""),
                                c.get("resolution_hint",""),
                                " | ".join(srcs), " | ".join(quotes)
                            ])
                    with open(tmpc, "rb") as f:
                        st.download_button("Download Conflicts CSV", f.read(), file_name="conflicts_merged.csv")

            except Exception as e:
                st.exception(e)
                st.error("Conflict detection failed.")

    st.markdown("---")
    st.subheader("Raw text & pages preview")
    if st.session_state.get("raw_text"):
        if st.checkbox("Show extracted text (first 15000 chars)"):
            st.text_area("Extracted Text", value=st.session_state["raw_text"][:15000], height=400)
    else:
        st.info("No extracted text yet. Upload a PDF and click Extract.")

# def merge_conflicts(llm_data: Dict[str, Any], regex_conflicts: List[Dict[str, Any]]) -> Dict[str, Any]:
#     out = {"conflicts": [], "summary": ""}
#     seen = set()
#     # take LLM results first
#     for c in (llm_data.get("conflicts") or []):
#         key = (c.get("type"), c.get("message"))
#         if key not in seen:
#             out["conflicts"].append(c); seen.add(key)
#     # append regex (mapped to similar shape)
#     for r in (regex_conflicts or []):
#         # map your existing fields to the LLM schema
#         c = {
#             "type": r.get("type"),
#             "severity": r.get("severity"),
#             "message": r.get("message"),
#             "category": r.get("category"),
#             "resolution_hint": r.get("resolution_hint"),
#             "evidence": [
#                 {"quote": (ev.get("snippet") or "")[:200], "sources": [f"page {ev.get('page')}"]}
#                 for ev in (r.get("evidence") or []) if isinstance(ev, dict)
#             ][:3],
#             "confidence": 0.65
#         }
#         key = (c.get("type"), c.get("message"))
#         if key not in seen:
#             out["conflicts"].append(c); seen.add(key)
#     out["summary"] = (llm_data.get("summary") or "").strip()
#     return out

# st.caption("ChatGPT One-Shot CAD • Large-context model • JSON-enforced responses")
# # End of file



